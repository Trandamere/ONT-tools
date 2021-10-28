#!/usr/bin/env python

import argparse
import pandas as pd
import os
import docx
import logging
import re
import math
from multiprocessing import Pool
from docx.shared import RGBColor
from docx.shared import Pt
from docx import shared
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docxtpl import DocxTemplate, RichText
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
# 打印出运行的时间
time1 = '运行时间：' + str(datetime.datetime.now())
print(time1)

# 设定监控日志输出文件名和内容形式
logging.basicConfig(format='%(asctime)s - %(message)s', filename='/mnt/c/Users/luping/Desktop/报告流程/RD/运行信息.txt', filemode='a', level=logging.INFO)

# 参数的导入与处理
parser = argparse.ArgumentParser()
parser.add_argument('-i', "--result_excel", required=True, help="the excel file with the result selected")
parser.add_argument('-b', "--database", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/RD/思可愈数据库-TNP-Seq病原菌测序项目2021.05.18.xlsx',help="database provided by the Ministry of Medicine")
parser.add_argument('-w', "--word_template_folder", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/RD/',help="folder where all word report templates are located")
parser.add_argument('-e', "--excel_template_folder", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/RD/',help="folder where all excel report templates are located")
parser.add_argument('-n', "--processes_number", type=int, default=5,help="并行进程数目")
parser.add_argument('-c', "--complex_excel", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/RD/mycobacterium_tuberculosis_complex.xlsx',help="结核分支杆菌复合群包含微生物表格")
parser.add_argument('-s', "--summary_excel", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/RD/OUTPUT/',help="summary documents before processing")
parser.add_argument('-o', "--output_dir", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/RD/OUTPUT/',help="supplement sample result")
parser.add_argument('-m', "--name_excel", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/RD/药品中英文对照表.xlsx',help="检测药品中英文对照表")
parser.add_argument('-B', "--barcode_picutre", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/RD/barcode/',help="条形码图片所在")
args = parser.parse_args()


info_client = pd.read_excel(args.result_excel).fillna('NA')
barcode_picture_path = args.barcode_picutre
Interpretation = pd.read_excel(args.result_excel, sheet_name='species_report').fillna('NA')
sheet=pd.read_excel(args.result_excel,sheet_name=None)
if 'resistance_report' in list(sheet.keys()):
    drug_resistance_df = pd.read_excel(args.result_excel, sheet_name='resistance_report').fillna('NA')
    if drug_resistance_df.shape[0] == 0:
        drug_resistance_df = 0
else:
    drug_resistance_df = 0
picture_dtat_df = pd.read_excel(args.result_excel, sheet_name='length_report').fillna('NA')
AMR_stat_df = pd.read_excel(args.result_excel, sheet_name='depth_report').fillna('NA')
medical_DB = pd.read_excel(args.database).fillna('NA')
result_file_name = args.result_excel.split("/")[-1].lower()
complex_df = pd.read_excel(args.complex_excel)

####################################################
# 定义函数
# 标准化输入微生物名称所用函数
# 标准化输入内容名称所用函数(变小写)
# 判断是否为结核分枝杆菌复合群
def change_bacteria_list(bacteria_list: list):
    new_bacteria_list = []
    for bacteria in bacteria_list:
        compare_bacteria: str = Nor(bacteria)
        find_complex_df = complex_df[complex_df['name'] == compare_bacteria]
        if find_complex_df.shape[0] == 0:
            new_bacteria_list.append(bacteria)
        else:
            new_bacteria_list.append('Mycobacterium tuberculosis complex')
    return new_bacteria_list
def Nor(x: str
    ) -> str:
    first: str= x.strip()
    standardized_string: str = " ".join(first.split())
    standardized_string: str = standardized_string.lower()
    return standardized_string


# 分类正式报告结果
def microbial_classification(bacteria_list: list,
    medical_DB: pd.DataFrame) -> list:
    result1_list = []
    result2_list = []
    result3_list = []
    result4_list = []
    result5_list = []
    result6_list = []
    for microbial in bacteria_list:
        #print('microbial',microbial)
        # compare_bac_name = Nor(micro)
        compare_bac_name: str = Nor(microbial)
        # print(medical_DB['英文名称'])
        # print(compare_bac_name)
        # print(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '分类'].iloc[0])
        #print(compare_bac_name)
        try:
            kingdom = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '分类'].iloc[0]
            genus = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '属'].iloc[0]
            # print(kingdom)
            # print(genus)
        except IndexError:
            kingdom = 'NA'
            genus = 'NA'
            # print(kingdom)
            # print(genus)
        if genus == '分枝杆菌属':
            result1_list.append(microbial)
        elif kingdom == '细菌':
            result2_list.append(microbial)
        elif kingdom == '真菌':
            result3_list.append(microbial)
        elif kingdom == '病毒' or kingdom == 'DNA病毒' or kingdom == 'RNA病毒':
            result4_list.append(microbial)
        elif kingdom == '古菌' or kingdom == '其他病原':
            result5_list.append(microbial)
        elif kingdom == '寄生虫':
            result6_list.append(microbial)
        else:
            logging.info(f"{microbial}没有在数据库中找到，导致报告中未显示！")
    return [result1_list,result2_list,result3_list,result4_list,result5_list,result6_list]

# 标准化输入列名称所用函数
def Nor_col(x: str
    ) -> str:
    first: str= str(x).strip()
    standardized_string: str = " ".join(first.split())
    return standardized_string
def table_1_make(sample_code: str,
    medical_DB: pd.DataFrame,
    complex_df: pd.DataFrame,
    all_bac: list,
    mic_dict: dict,
    result_list: list,
    Interpretation: pd.DataFrame
    ) -> list:
    sample_result_list = []
    bac_list = []
    column_R = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_R_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
    for bac_name in result_list:
        if len(bac_name) != 0:
            compare_bac_name = Nor(bac_name)
            dic_bac: dict ={}
            find_complex_df = complex_df[complex_df['name'] == compare_bac_name]
            if find_complex_df.shape[0] == 0:
                dic_bac['分类'] = all_bac[1]
            else:
                dic_bac['分类'] = all_bac[0]
            dic_bac['检测结果'] = '阳性'
            try:
                dic_bac['中文名'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]
                dic_bac['微生物'] = bac_name
                dic_bac['分类顺序'] = mic_dict[dic_bac['分类']]
            except IndexError:
                logging.info(f"{bac_name}在数据库中没有找到")
            try:
                dic_bac['序列数'] = Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0]
            except KeyError:
                logging.info(f"{bac_name}没有找到相应的序列数")
            bac_list.append(dic_bac['分类'])
            sample_result_list.append(dic_bac)
    bac_list = list(set(bac_list))
    for non_bac in all_bac:
        if non_bac not in bac_list:
            dic_bac = {}
            dic_bac['分类'] = non_bac
            dic_bac['分类顺序'] = mic_dict[dic_bac['分类']]
            dic_bac['检测结果'] = '阴性'
            dic_bac['中文名'] = '/'
            dic_bac['微生物'] = '/'
            dic_bac['序列数'] = '/'
            sample_result_list.append(dic_bac)
    sample_result_list: list = sorted(sample_result_list, key = lambda x:x['分类顺序'])
    return sample_result_list

def table_2_3_make(sample_code: str,
    complex_df: pd.DataFrame,
    AMR_df: pd.DataFrame,
    AMR_stat_colname: str,
    id_dict: dict
    ) -> list:
    AMR_stat_colname1 = AMR_stat_colname.replace('count','depth').replace('_RD','')
    number = 0
    for depth in AMR_df[AMR_stat_colname1].tolist():
        if str(depth).startswith(r'*'):
            number += 1
            break
    if number != 0:
        handle_df = AMR_df[AMR_df[AMR_stat_colname1].str.contains('\*', na=False)]
        table2_number = 0
        table3_number = 0
        table2_list = []
        table3_list = []
        handle_df.apply(table_2_3_make_dict, complex_df=complex_df,AMR_stat_colname=AMR_stat_colname,id_dict=id_dict,table2_list=table2_list,table3_list=table3_list, axis=1)
        for table_line in table2_list:
            table2_number += 1
            table_line['序号'] = table2_number
        for table_line in table3_list:
            table3_number += 1
            table_line['序号'] = table3_number 
        if len(table2_list) == 0:
            table2_list = [{'突变率': '--', '基因': '--', '突变描述': '--', '药品': '--', '氨基酸突变': '--', '序号': '--'}]
        if len(table3_list) == 0:
            table3_list = [{'突变率': '--', '基因': '--', '突变描述': '--', '药品': '--', '氨基酸突变': '--', '序号': '--'}]
    else:
        table2_list = [{'突变率': '--', '基因': '--', '突变描述': '--', '药品': '--', '氨基酸突变': '--', '序号': '--'}]
        table3_list = [{'突变率': '--', '基因': '--', '突变描述': '--', '药品': '--', '氨基酸突变': '--', '序号': '--'}]
    return [table2_list,table3_list]


def table_4_make(sample_code: str,
    medical_DB: pd.DataFrame,
    result_list: list
    ) -> list:
    sample_result_list = []
    bac_list = []
    for bac_name in result_list:
        # print(bac_name)
        if len(bac_name) != 0:
            compare_bac_name = Nor(bac_name)
            dic_bac: dict ={}
            try:
                dic_bac['名称'] = str(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]) + '(' + bac_name + ')'
                dic_bac['临床意义'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '临床意义'].iloc[0]
                dic_bac['推荐用药'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '常用药物'].iloc[0]
            except IndexError:
                logging.info(f"{bac_name}在数据库中没有找到")
            dic_bac['行名1'] = '检 出 菌 种'
            dic_bac['行名2'] = '临 床 意 义'
            dic_bac['行名3'] = '推 荐 用 药'
            sample_result_list.append(dic_bac)
    if len(sample_result_list) == 0:
        dic_bac: dict ={}
        dic_bac['行名1'] = '检 出 菌 种'
        dic_bac['行名2'] = '临 床 意 义'
        dic_bac['行名3'] = '推 荐 用 药'
        dic_bac['名称'] = '-'
        dic_bac['临床意义'] = '-'
        dic_bac['推荐用药'] = '-'
        sample_result_list.append(dic_bac)
    return sample_result_list


def table_2_3_make_dict(AMR_df: pd.Series,
    complex_df: pd.DataFrame,
    AMR_stat_colname: str,
    id_dict: dict,
    table2_list: list,
    table3_list: list,
    ) -> list:
    new_dict = {}
    compare_bac_name = Nor(id_dict[AMR_df['chr']])
    find_complex_df = complex_df[complex_df['name'] == compare_bac_name]
    AMR_stat_colname2 = AMR_stat_colname.replace('count','vaf')
    new_dict['突变率'] = AMR_df[AMR_stat_colname2]
    new_dict['基因'] = AMR_df['gene_name']
    new_dict['突变描述'] = AMR_df['hgvs_c']
    new_dict['氨基酸突变'] = AMR_df['hgvs_p']
    if 'NA' in new_dict['氨基酸突变']:
        new_dict['氨基酸突变']= '-'
    medicine_list = AMR_df['drug'].split(';')
    medicine_name_list = []
    for medicine in medicine_list:
        medicine_name = name_df.loc[name_df['药品名称'] == medicine, '中文名'].iloc[0]
        medicine_name_list.append(medicine_name)
    if len(medicine_name_list) == 1:
        new_dict['药品'] = medicine_name_list[0]
    else:
        new_dict['药品'] = ';'.join(medicine_name_list)
    if find_complex_df.shape[0] == 0:
        table3_list.append(new_dict)
    else:
        table2_list.append(new_dict)
    if len(table2_list) == 0:
        table2_list = [{'突变率': '--', '基因': '--', '突变描述': '--', '药品': '--', '氨基酸突变': '--', '序号': '--'}]
    if len(table3_list) == 0:
        table3_list = [{'突变率': '--', '基因': '--', '突变描述': '--', '药品': '--', '氨基酸突变': '--', '序号': '--'}]


def table_context_TB(sample_code: str,
    info_client: pd.DataFrame,
    medical_DB: pd.DataFrame,
    complex_df: pd.DataFrame,
    id_dict: dict,
    AMR_df: pd.DataFrame,
    AMR_stat_colname: str,
    Interpretation: pd.DataFrame,
    formal_report,
    supplementary_report
    ) -> list:
    handle_df = info_client.query('样本编号 == @sample_code').iloc[0,:]
    all_bac: list = ['结核分枝杆菌复合群（MTBC）','非结核分枝杆菌（NTM）']
    mic_dict: dict = {'结核分枝杆菌复合群（MTBC）':1,'非结核分枝杆菌（NTM）':2}
    # result_list = handle_df['正式报告结果'].split(',')
    result_list = formal_report.split(',')
    result_list = table_2_3_make(sample_code=sample_code,complex_df=complex_df,AMR_df=AMR_df,AMR_stat_colname=AMR_stat_colname,id_dict=id_dict)
    table2_list = result_list[0]
    table3_list = result_list[1]
    return [table2_list,table3_list]


def make_picture_length(length_colname: str,
    picture_dtat_df: pd.DataFrame,
    ) -> None:
    R_out = length_colname + '.r'
    rscript=f'''#! /path/to/Rscript
library(openxlsx)
library(ggplot2)
data<- read.xlsx('{args.result_excel}', sheet='length_report')
names(data)[names(data) == '{length_colname}'] <- 'Frequency'
data <- data[c(1:8),]
row.names(data) <- as.list(data)$length
read_length_hist <- ggplot(data, mapping=aes(x=rownames(data), y=Frequency)) +
geom_bar(stat="identity", fill="#1F4E79", colour="#1F4E79") +
scale_x_discrete(limits=factor(rownames(data))) +
labs(x="length(bp)", y="ratio(%)") +
theme(panel.grid=element_blank(), panel.background=element_rect(color="black", fill="transparent")) + 
theme(axis.text =element_text(size=4))
ggsave(file="{length_colname}_length.png",read_length_hist, width = 3.4, height = 3.4)
ggsave(file="{length_colname}_length_single.png",read_length_hist, width = 6, height = 3)
    '''
    out = open(R_out,'w')
    out.write(rscript)
    out.close()
    cmd = "Rscript " + R_out
    os.system(cmd)
    os.remove(R_out)


# 将表格添加到特定的字后面
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def move_picture_length(doc,
    png_name: str,
    png_name2: str,
    png_name_depth: str
    ) -> None:
    table = doc.add_table(rows=2, cols=3)
    cell_length_title_paragraph = table.cell(0,1).paragraphs[0]#'测序长度统计'
    cell_length_title_paragraph.text = '测序长度统计'
    cell_length_title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell_length = table.cell(1, 0)
    ph_length   = cell_length.paragraphs[0]
    run_length  = ph_length.add_run()
    cell = table.cell(1, 1)
    ph =cell.paragraphs[0]
    run = ph.add_run()
    run.add_picture(png_name2)
    target = None
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith('图一'):
            # print(paragraph_text)
            target = paragraph
            break
    move_table_after(table, target)
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith('图一'):
            delete_paragraph(paragraph)
    os.remove(png_name)
    os.remove(png_name2)
    os.remove(png_name_depth)
    return doc


def move_picture_length_depth(doc,
    png_name_length: str,
    png_name_length2: str,
    png_name_depth: str
    ) -> None:
    table = doc.add_table(rows=2, cols=2)
    #cell_length_title = table.cell(0,0)
    cell_length_title_paragraph = table.cell(0,0).paragraphs[0]#'测序长度统计'
    cell_length_title_paragraph.text = '测序长度统计'
    cell_length_title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell_depth_title_paragraph  = table.cell(0,1).paragraphs[0]
    cell_depth_title_paragraph.text  = '测序深度统计'
    cell_depth_title_paragraph.alignment  = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell_length = table.cell(1, 0)
    cell_depth  = table.cell(1, 1) 
    #添加标题
    ph_length   = cell_length.paragraphs[0]
    ph_depth    = cell_depth.paragraphs[0]
    run_length  = ph_length.add_run()
    run_depth   = ph_depth.add_run()
    run_length.add_picture(png_name_length)
    run_depth.add_picture(png_name_depth)
    target = None
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith('图一'):
            # print(paragraph_text)
            target = paragraph
            break
    move_table_after(table, target)
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith('图一'):
            delete_paragraph(paragraph)
    os.remove(png_name_length)
    os.remove(png_name_length2)
    os.remove(png_name_depth)
    return doc

# def move_picture_barcode(doc,
#     png_name: str,
#     barcode_picture_path
#     ) -> None:
#     for paragraph in doc.paragraphs:
#         paragraph_text = paragraph.text
#         if paragraph_text.endswith('条形码'):
#             # paragraph.text=paragraph.text.replace('条形码','')
#             paragraph.clear()
#             print(paragraph_text,type(paragraph_text))
#             # paragraph_text =''
#             run = paragraph.add_run(paragraph_text)
#             run.add_picture(barcode_picture_path+png_name+'.png',width=Cm(2),height=Cm(2))
#             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
#             run.font.size = Pt(1)
#             print(barcode_picture_path+png_name+'.png')
#             break
#     return doc
def move_picture_barcode(doc,
    png_name: str,
    barcode_picture_path):
    table = doc.tables[0]
    png_name=png_name.replace('-s','')
    paragraph=table.cell(1,0).paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(barcode_picture_path+png_name+'.png',width=Cm(2.1),height=Cm(2.1))
    print(barcode_picture_path+png_name+'.png')
    return doc




#删除段落
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None



def make_picture_depth(AMR_stat_colname: str,
    AMR_stat_df: pd.DataFrame,
    file_name: str
    ) -> None:
    R_out = file_name + '.r'
    rscript=f'''#! /path/to/Rscript
library(openxlsx)
library(ggplot2)
data<- read.xlsx('{AMR_stat_df}', sheet='depth_report')
names(data)[names(data) == '{file_name}'] <- 'Frequency'
data <- data[c(1:11),]
row.names(data) <- as.list(data)$depth
read_length_hist <- ggplot(data, mapping=aes(x=rownames(data), y=Frequency)) +
geom_bar(stat="identity", fill="#1F4E79", colour="#1F4E79") +
scale_x_discrete(limits=factor(rownames(data))) +
labs(x="Depth", y="Count") +
theme(panel.grid=element_blank(), panel.background=element_rect(color="black", fill="transparent")) + 
theme(axis.text =element_text(size=4))
ggsave(file="{file_name}.png",read_length_hist, width = 3.4, height = 3.4)
    '''
    out = open(R_out,'w')
    out.write(rscript)
    out.close()
    cmd = "Rscript " + R_out
    os.system(cmd)
    os.remove(R_out)


def move_picture_depth(doc,
    png_name: str
    ) -> None:
    table = doc.add_table(rows=2, cols=2)
    cell = table.cell(1, 1)
    ph =cell.paragraphs[0]
    run = ph.add_run()
    run.add_picture(png_name)
    target = None
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith('图一'):
            # print(paragraph_text)
            target = paragraph
            delete_paragraph(paragraph)
            break
    move_table_after(table, target)
    os.remove(png_name)
    return doc


# 查询检测项目的简称
def project_shorthand(sample_code: str) -> str:
    project_name = info_client.loc[info_client['样本编号'] == sample_code, '检测项目'].iloc[0].rstrip()
    hand = 0
    if '呼吸' in project_name:
        hand: str = 'HX'
    else:
        name = info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]
        logging.info(f'{name}的检测项目填写错误导致报告生成失败！')
    # hand: str = 'RD'
    return hand


def appendix_drugs(table2_list: list,
    table3_list: list,
    first_list: list,
    second_list: list,
    NTM_list: list
    ) -> list:
    table2_drug_list = []
    for table_line in table2_list:
        new_drug_list = table_line['药品'].split(';')
        table2_drug_list.extend(new_drug_list)
    first_find_list = []
    first_dont_find_list = []
    for first_drug in first_list:
        if first_drug in table2_drug_list:
            first_find_list.append(first_drug)
        else:
            first_dont_find_list.append(first_drug)
    if len(first_find_list) == 0:
        first_find_list = ['/']
    if len(first_dont_find_list) == 0:
        first_dont_find_list = ['/']
    first_find = '\n'.join(first_find_list)
    first_dont_find = '\n'.join(first_dont_find_list)
    rt1 = RichText('')
    rt2 = RichText('')
    rt1.add(first_find)
    rt2.add(first_dont_find)
    second_find_list = []
    second_dont_find_list = []
    for second_drug in second_list:
        if second_drug in table2_drug_list:
            second_find_list.append(second_drug)
        else:
            second_dont_find_list.append(second_drug)
    if len(second_find_list) == 0:
        second_find_list = ['/']
    second_find = '\n'.join(second_find_list)
    number = math.floor(len(second_dont_find_list)/2)
    second_dont_find_one_list = second_dont_find_list[0:number]
    second_dont_find_two_list = second_dont_find_list[number:]
    if len(second_dont_find_one_list) == 0:
        second_dont_find_one_list = ['/']
    if len(second_dont_find_two_list) == 0:
        second_dont_find_two_list = ['/']
    second_dont_find_one = '\n'.join(second_dont_find_one_list)
    second_dont_find_two = '\n'.join(second_dont_find_two_list)
    rt3 = RichText('')
    rt4 = RichText('')
    rt5 = RichText('')
    rt3.add(second_find)
    rt4.add(second_dont_find_one)
    rt5.add(second_dont_find_two)
    table3_drug_list = []
    for table_line in table3_list:
        new_drug_list = table_line['药品'].split(';')
        table3_drug_list.extend(new_drug_list)
    third_find_list = []
    third_dont_find_list = []
    for third_drug in NTM_list:
        if third_drug in table3_drug_list:
            third_find_list.append(third_drug)
        else:
            third_dont_find_list.append(third_drug)
    if len(third_find_list) == 0:
        third_find_list = ['/']
    if len(third_dont_find_list) == 0:
        third_dont_find_list = ['/']
    third_find = '\n'.join(third_find_list)
    third_dont_find = '\n'.join(third_dont_find_list)
    rt6 = RichText('')
    rt7 = RichText('')
    rt6.add(third_find)
    rt7.add(third_dont_find)
    return [rt1,rt2,rt3,rt4,rt5,rt6,rt7]
def remove_item(old_list: list,
    remove_element: str):
    new_list = []
    for element in old_list:
        if element != remove_element:
            new_list.append(element)
    return new_list

def add_medical_interpretation(AMR_stat_colname: str,
    sample_type: str,
    all_result_list: list,
    MTB_result_list: list,
    NTM_result_list: list,
    id_dict: dict,
    MTB_gene_dict: dict,
    NTM_gene_dict: dict,
    AMR_df: pd.DataFrame):
    MTB_gene_found_list = []
    MTB_mutation_gene_found_list = []
    MTB_gene_unfound_list = []
    NTM_gene_found_list = []
    NTM_mutation_gene_found_list = []
    NTM_gene_unfound_list = []
    MTB_drugs_found_list = []
    MTB_mutation_drugs_found_list = []
    MTB_drugs_unfound_list = []
    NTM_drugs_found_list = []
    NTM_mutation_drugs_found_list = []
    NTM_drugs_unfound_list = []
    rt = RichText('')
    rt.add(f'该{sample_type}中,')
    unfound_info = ''
    handle_df = AMR_df[AMR_df[AMR_stat_colname] != 0]
    number = 0
    for result in all_result_list:
        if result['微生物'] == 'Mycobacterium tuberculosis':
            number += 1
            MTB_df = handle_df[handle_df['chr'] == 'NC_000962.3']
            MTB_genes = MTB_df['gene_name'].tolist()
            for key, value in MTB_gene_dict.items():
                if key in MTB_genes:
                    MTB_gene_found_list.append(key)
                    MTB_drugs_found_list.extend(value)
                else:
                    MTB_gene_unfound_list.append(key)
                    MTB_drugs_unfound_list.extend(value)
            MTB_drugs_found_list = list(set(MTB_drugs_found_list))
            MTB_drugs_unfound_list = list(set(MTB_drugs_unfound_list))
            for line in MTB_result_list:
                if line['序号'] != '--':
                    MTB_mutation_gene_found_list.append(line['基因'])
                    MTB_mutation_drugs_found_list.append(line['药品'])
            new_MTB_drugs_unfound_list = MTB_drugs_unfound_list
            new_MTB_gene_unfound_list = MTB_gene_unfound_list
            for key, value in MTB_gene_dict.items():
                for drug in value:
                    if drug in MTB_drugs_unfound_list and drug in MTB_mutation_drugs_found_list:
                        new_MTB_drugs_unfound_list = remove_item(MTB_drugs_unfound_list,drug)
                        if len(value) == 1:
                            new_MTB_gene_unfound_list = remove_item(MTB_gene_unfound_list,key)
            if len(MTB_mutation_gene_found_list) != 0:
                mutation_gene_info = '、'.join(MTB_mutation_gene_found_list)
                mutation_drugs_info = '、'.join(MTB_mutation_drugs_found_list)
                rt.add(f'结核分枝杆菌检出{mutation_gene_info}基因耐药位点突变，表示可能会对{mutation_drugs_info}药物耐药。请结合突变率及分子DST与表型DST符合率来综合判断。\n')
            else:
                found_gene_info = '、'.join(MTB_gene_found_list)
                found_drugs_info = '、'.join(MTB_drugs_found_list)
                rt.add(f'结核分枝杆菌检出{found_gene_info}基因耐药位点未发生突变，表示对{found_drugs_info}药物敏感。请结合突变率及分子DST与表型DST符合率来综合判断。\n')
            if len(new_MTB_gene_unfound_list) != 0:
                unfound_gene_info = '、'.join(new_MTB_gene_unfound_list)
                unfound_drugs_info = '、'.join(new_MTB_drugs_unfound_list)
                unfound_info = unfound_info + f'由于该样本中结核分枝杆菌含量较低，{unfound_gene_info}基因未能检出，无法判断其是否突变，从而不能判断对{unfound_drugs_info}敏感/耐药。\n'
        elif result['微生物'] != '/' and result['分类'] == '非结核分枝杆菌（NTM）':
            chinese_name = result['中文名']
            number += 1
            NTM_id_list = []
            for key, value in NTM_result_list.items():
                if value == result['微生物']:
                    NTM_id_list.append(key)
            NTM_df = handle_df[handle_df['chr'].isin(NTM_id_list)]
            NTM_genes = NTM_df['gene_name'].tolist()
            tmp_NTM_gene_found_list = []
            tmp_NTM_mutation_gene_found_list = []
            tmp_NTM_gene_unfound_list = []
            tmp_NTM_drugs_found_list = []
            tmp_NTM_mutation_drugs_found_list = []
            tmp_NTM_drugs_unfound_list = []
            for key, value in NTM_gene_dict.items():
                if key in NTM_genes:
                    tmp_NTM_gene_found_list.append(key)
                    tmp_NTM_drugs_found_list.extend(value)
                else:
                    tmp_NTM_gene_unfound_list.append(key)
                    tmp_NTM_drugs_unfound_list.extend(value)
            for line in NTM_result_list:
                if line['序号'] != '--':
                    tmp_NTM_mutation_gene_found_list.append(line['基因'])
                    tmp_NTM_mutation_drugs_found_list.append(line['药品'])
            new_tmp_NTM_drugs_unfound_list = tmp_NTM_drugs_unfound_list
            new_tmp_NTM_gene_unfound_list = tmp_NTM_gene_unfound_list
            for key, value in NTM_gene_dict.items():
                for drug in value:
                    if drug in tmp_NTM_drugs_unfound_list and drug in tmp_NTM_mutation_drugs_found_list:
                        new_tmp_NTM_drugs_unfound_list = remove_item(tmp_NTM_drugs_unfound_list,drug)
                        if len(value) == 1:
                            new_tmp_NTM_gene_unfound_list = remove_item(tmp_NTM_gene_unfound_list,key)
            if len(tmp_NTM_mutation_gene_found_list) != 0:
                mutation_gene_info = '、'.join(tmp_NTM_mutation_gene_found_list)
                mutation_drugs_info = '、'.join(tmp_NTM_mutation_drugs_found_list)
                rt.add(f'{chinese_name}检出{mutation_gene_info}基因耐药位点突变，表示可能会对{mutation_drugs_info}药物耐药。请结合突变率及分子DST与表型DST符合率来综合判断。\n')
            else:
                found_gene_info = '、'.join(tmp_NTM_gene_found_list)
                found_drugs_info = '、'.join(tmp_NTM_drugs_found_list)
                rt.add(f'{chinese_name}检出{found_gene_info}基因耐药位点未发生突变，表示对{found_drugs_info}药物敏感。请结合突变率及分子DST与表型DST符合率来综合判断。\n')
            if len(new_tmp_NTM_gene_unfound_list) != 0:
                unfound_gene_info = '、'.join(new_tmp_NTM_gene_unfound_list)
                unfound_drugs_info = '、'.join(new_tmp_NTM_drugs_unfound_list)
                unfound_info = unfound_info + f'由于该样本中{chinese_name}含量较低，{unfound_gene_info}基因未能检出，无法判断其是否突变，从而不能判断对{unfound_drugs_info}敏感/耐药。\n'
            NTM_gene_found_list.extend(tmp_NTM_gene_found_list)
            NTM_mutation_gene_found_list.extend(tmp_NTM_mutation_gene_found_list)
            NTM_gene_unfound_list.extend(tmp_NTM_gene_unfound_list)
            NTM_drugs_found_list.extend(tmp_NTM_drugs_found_list)
            NTM_mutation_drugs_found_list.extend(tmp_NTM_mutation_drugs_found_list)
            NTM_drugs_unfound_list.extend(tmp_NTM_drugs_unfound_list)
    if number == 0:
        rt.add(f'未检出分枝杆菌，请结合临床综合判断是否为其他的感染。\n')
    else:
        rt.add(unfound_info)
    return [rt,MTB_gene_found_list,NTM_gene_found_list]
    

def appendix_color_change(doc,
    table_list: list,
    number: int
    ) -> None:
    if number == 1:
        table_number_list = [16,17]#需要确定
    else:
        table_number_list = [18]
    for table_number in table_number_list:
        table = doc.tables[table_number]
        rownums = len(table.rows)
        for table_line in table_list:
            drug_list = table_line['药品'].split(';')
            for drug in drug_list:
                for x in range(rownums):
                    drug_name = table.cell(x,1).text
                    if drug_name == drug:
                        if table_line['基因'] == table.cell(x,0).text:
                            run = table.cell(x,2).paragraphs[0]
                            run.text = "检出"
                            content = run.text
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(11)
                            run.font.name = 'Arial'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def color_change(doc,
    table_number
    ) -> None:
    table = doc.tables[table_number]
    rownums = len(table.rows)
    for x in range(rownums):
        if table.cell(x,0).text == '结核分枝杆菌复合群（MTBC）' or table.cell(x,0).text == '非结核分枝杆菌（NTM）' :
            run = table.cell(x,table_number-3).paragraphs[0]
            content = run.text
            run.text = ''
            run = run.add_run(content)
            if content == '阳性':
                run.font.color.rgb = RGBColor(255,0,0)
            else:
                run.font.color.rgb = RGBColor(0,176,80)
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
def color_change3(doc,
    ) -> None:
    number = 7
    while number < 9:
        table_info_name = f'表{number}信息'
        # print(table_info_name)
        table_info = dic_client[sample_code][table_info_name]
        # print(table_info)
        change_color_table = doc.tables[2+number]
        if table_info[0]['基因'] != '--':
            for i,line in enumerate(table_info):
                row = i+2
                col = len(change_color_table.columns)
                #print(change_color_table(row,col).text)
                for col_number in range(col):
                    run = change_color_table.cell(row,col_number).paragraphs[0]
                    content = run.text
                    run.text = ''
                    run = run.add_run(content)
                    run.font.color.rgb = RGBColor(255,0,0)
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    if number==7:
                        if col_number == 0:
                            run.italic = True
                    else:
                        if col_number == 1:
                            run.italic = True
        else:
            run = change_color_table.cell(2,1).paragraphs[0]
            run.text = ''
            run = run.add_run('--')
            run.font.color.rgb = RGBColor(0,0,0)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')						
        number += 1

def color_change2(doc,
    sample_code: str,
    info_client: pd.DataFrame,
    medical_DB: pd.DataFrame,
    complex_df: pd.DataFrame,
    appendix_species: list,
    formal_report,
    supplementary_report
    ) -> None:
    handle_df = info_client.query('样本编号 == @sample_code').iloc[0,:]
    # result_list = handle_df['正式报告结果'].split(',')
    result_list = formal_report.split(',')
    table = doc.tables[20]
    rownums = len(table.rows)
    find_list = []
    #print('result_list',result_list)
    for bac_name in result_list:
        #print('bac_name',bac_name)
        if len(bac_name) != 0:
            compare_bac_name = Nor(bac_name)
            chinese_name = 0
            #print('compare_bac_name',compare_bac_name)
            try:
                chinese_name = Nor_col(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0])
                #print('chinese_name',chinese_name)
            except IndexError:
                logging.info(f"{handle_df['患者姓名']}的结果{bac_name}未找到中文名")
            # print(chinese_name)
            if type(chinese_name) == str:
                if chinese_name in appendix_species:
                    for x in range(rownums):
                        #print('table.cell(x,1).text',table.cell(x,1).text)
                        if table.cell(x,1).text == chinese_name:
                            #print('table.cell(x,1).text',table.cell(x,1).text)
                            run = table.cell(x,3).paragraphs[0]
                            run.text = '检出'
                            content = run.text
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(9)
                            run.font.name = 'Arial'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            run = table.cell(x,2).paragraphs[0]
                            content = run.text
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(9)
                            run.font.name = 'Arial'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            run = table.cell(x,1).paragraphs[0]
                            content = run.text
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(9)
                            run.font.name = 'Arial'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体') 
                    find_list.append(chinese_name)                           
                    #result_list.remove(bac_name)
    #print('find_list',find_list)           
    if len(result_list) != 0:
        for bac_name in result_list:
            if len(bac_name) != 0:
                compare_bac_name = Nor_col(bac_name)
                find_complex_df = complex_df[complex_df['name'] == compare_bac_name]
                if find_complex_df.shape[0] == 0:
                    classification_name = '非结核分枝杆菌'
                else:
                    classification_name = '结核分枝杆菌复合群'
                chinese_name = 0
                try:
                    chinese_name = Nor_col(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0])
                except IndexError:
                    logging.info(f"{handle_df['患者姓名']}的结果{bac_name}未找到中文名")
                if type(chinese_name) == str:
                    for x in range(rownums):
                        if table.cell(x,1).text != '结核分枝杆菌' and table.cell(x,1).text not in find_list and table.cell(x,1).text == classification_name:
                            table.cell(x,1).text = chinese_name
                            run = table.cell(x,3).paragraphs[0]
                            run.text = '检出'
                            content = run.text
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(10)
                            run.font.name = 'Arial'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    #result_list.remove(bac_name)
                    find_list.append(chinese_name)

# 表格信息的生成2
def table2_context(sample_code: str,
    info_client: pd.DataFrame,
    Interpretation: pd.DataFrame,
    formal_report,
    supplementary_report,
    medical_DB: pd.DataFrame) -> list:
    # print(sample_code)
    all_bac_list: list = [['细菌'],['细菌'],['真菌'],['病毒'],['其他病原'],['寄生虫']]
    mic_dict_list: dict = [{'细菌':1}, {'细菌':1}, {'真菌':1}, {'病毒':1, 'DNA病毒':1, 'RNA病毒':1}, {'其他病原':1},{'寄生虫':1}]
    bacteria_list: list = []
    pat_name = handle_df['患者姓名'].iloc[0]
    pat_sample_id_list=[]
    if '-s' in handle_df['样本编号'].iloc[0]:
        pat_sample_id_list.append(handle_df['样本编号'].iloc[0]).replace('-s','')
    else:
        pat_sample_id_list.append(handle_df['样本编号'].iloc[0])
    pat_sample_id = pat_sample_id_list[0]    
    info_client_info = info_client[(info_client['患者姓名'] == pat_name) & (info_client['样本编号'] == pat_sample_id) & (info_client['备注'].str.contains('通用引物'))]
    
    #print(info_client_info)
    #print('sample_code',sample_code)
    #print ('info_client[样本编号]',info_client['样本编号'])
    if formal_report != 'NA':
        bacteria_list: list = formal_report.split(',')
    #print(bacteria_list)
    # bacteria_list = change_bacteria_list(bacteria_list)
    # print(bacteria_list)
    all_kinds_list = microbial_classification(bacteria_list=bacteria_list,medical_DB=medical_DB)
    #print(all_kinds_list)
    supplementary_results = 0
    table1_list = []
    for index, result_list in enumerate(all_kinds_list):
        # print(result_list)
        table1_list.append(find_info(result_list=result_list, sample_code=sample_code, Interpretation=Interpretation, info_client=info_client, medical_DB=medical_DB, all_bac=all_bac_list[index], mic_dict=mic_dict_list[index], supplementary_results=supplementary_results))
        # print(table1_list)
    table2_list = []
    #print(table1_list)
    #print(table2_list)
    if supplementary_report != 'NA':
        supplementary_results += 1
        bacteria_list: list = supplementary_report.split(',')
        # bacteria_list = change_bacteria_list(bacteria_list)
        table2_list: list = find_info(result_list=bacteria_list, sample_code=sample_code, Interpretation=Interpretation, info_client=info_client, medical_DB=medical_DB, all_bac=all_bac_list[0], mic_dict=mic_dict_list[0], supplementary_results=supplementary_results)
    else:
        table2_list = [{'中文名': '--', '分类': '--', '分类顺序': 1, '相对丰度': '--', '序列数': '--', '微生物': '--', '备注': '--'}]
    return [table1_list,table2_list]


# 表格信息的生成2
def table7_make(sample_code: str,
    info_client: pd.DataFrame,
    drug_resistance_df: pd.DataFrame
    ) -> list:
    handle_df = info_client[info_client['样本编号'] == sample_code]
    pat_name = handle_df['患者姓名'].iloc[0]
    pat_sample_id_list=[]
    if '-s' in handle_df['样本编号'].iloc[0]:
        pat_sample_id_list.append(handle_df['样本编号'].iloc[0])
    else:
        pat_sample_id_list.append(handle_df['样本编号'].iloc[0]+'-s')
    pat_sample_id = pat_sample_id_list[0]    
    #drug_resistance_info = info_client[(info_client['患者姓名'] == pat_name) & (info_client['备注'].str.contains('普通耐药'))]
    #drug_resistance_info = info_client[(info_client['患者姓名'] == pat_name) & (info_client['样本编号'] == pat_sample_id) &(info_client['备注'].str.contains('结核耐药'))]
    drug_resistance_info = info_client[(info_client['患者姓名'] == pat_name) & (info_client['样本编号'] == pat_sample_id) &(info_client['备注'].str.contains('普通耐药'))]
    # print(drug_resistance_info)
    table_list = []
    if drug_resistance_info.shape[0] != 0:
        compare_sample_code = drug_resistance_info['样本编号'].iloc[0]
        # print(compare_sample_code)
        drug_resistance_colname = drug_resistance_info.loc[drug_resistance_info['样本编号'] == compare_sample_code, 'barcode'].iloc[0] + '_' + project_shorthand(compare_sample_code) + '_' + str(drug_resistance_info.loc[drug_resistance_info['样本编号'] == compare_sample_code, '患者姓名'].iloc[0]) + '_' + compare_sample_code + '_num'
        # print(drug_resistance_colname)
        # print(drug_resistance_df)
        try:
            drug_resistance_result = drug_resistance_df[drug_resistance_colname]
        except KeyError:
            drug_resistance_colname = drug_resistance_info.loc[drug_resistance_info['样本编号'] == compare_sample_code, 'barcode'].iloc[0] + '_' + str(drug_resistance_info.loc[drug_resistance_info['样本编号'] == compare_sample_code, '患者姓名'].iloc[0]) + '_' + compare_sample_code + '_num'
            drug_resistance_result = drug_resistance_df[drug_resistance_colname]
        gene_list = drug_resistance_df['gene'].tolist()
        for gene in gene_list:
            number = drug_resistance_df.query('gene == @gene').iloc[0,:][drug_resistance_colname]
            if int(number) != 0:
                dic_bac ={}
                dic_bac['基因'] = gene
                dic_bac['药物'] = drug_resistance_df.query('gene == @gene').iloc[0,:]['drug']
                table_list.append(dic_bac)
    if len(table_list) == 0:
        dic_bac ={}
        dic_bac['基因'] = '--'
        dic_bac['药物'] = '--'
        table_list.append(dic_bac)
    return table_list

# 临床解析的生成
def clinical(sample_code: str,
    info_client: pd.DataFrame,
    medical_DB: pd.DataFrame,
    formal_report,
    supplementary_report,
    manufacturer: str):
    rt = RichText('')
    if  manufacturer == 'beagle':
        bacteria_list: list =  formal_report.split(',')
    elif manufacturer == 'seegene' or 'boruilin' or 'beijing':
        if supplementary_report != 'NA':
            bacteria_list: list =  supplementary_report.split(',')
        else:
            return rt
    # bacteria_list = change_bacteria_list(bacteria_list)
    dic_number: dict = {1:'一', 2:'二', 3:'三', 4:'四', 5:'五', 6:'六', 7:'七', 8:'八', 9:'九', 10:'十', 11:'十一', 12:'十二', 13:'十三', 14:'十四', 15:'十五', 16:'十六', 17:'十七', 18:'十八'}
    drug_indication_value: int = 0
    for index, bac_name in enumerate(bacteria_list):
        compare_bac_name: str = Nor(bac_name)
        try:
            new_name = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]
            if str(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]) != 'NA':
                chinese_name: str = str(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0])
            else:
                chinese_name = ''
            rt.add(dic_number[index+1]+' '+chinese_name, bold=True)
            rt.add('(', bold=True)
            report_name: str = bac_name.replace("[","")
            report_name: str = report_name.replace("]","")
            rt.add(report_name, italic=True, bold=True)
            rt.add(')\n', bold=True)
            if  manufacturer == 'beagle':
                try:
                    if str(medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '常用药物'].iloc[0]) != 'NA':
                        rt.add('    '+'1 临床意义', bold=True)
                        rt.add('\n'+'    '+medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '临床意义'].iloc[0] )
                        rt.add('\n'+'    '+'2 常用药物', bold=True)
                        rt.add('\n'+'    '+str(medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '常用药物'].iloc[0])+'\n'  )
                        drug_indication_value += 1
                    elif str(medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '临床意义'].iloc[0]) != 'NA':
                        rt.add('    '+'临床意义', bold=True)
                        rt.add('\n'+'    '+medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '临床意义'].iloc[0]+'\n' )
                except IndexError:
                    logging.info(f"正式结果中的{bac_name}在数据库中的检测项目信息有问题，影响临床解析的生成")
            elif manufacturer == 'seegene' or 'boruilin' or 'beijing':
                try:
                    rt.add('    '+medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '临床意义'].iloc[0]+'\n')
                except IndexError:
                    logging.info(f"正式结果中的{bac_name}在数据库中的检测项目信息有问题，影响临床解析的生成")
        except IndexError:
            logging.info(f"正式结果中的{bac_name}在数据库中没有找到，影响临床解析的生成")
    if drug_indication_value != 0 and manufacturer == 'beagle':
        rt.add('\n注：常用药物为临床常规药物，且无法覆盖药敏结果，具体用药请结合临床药敏结果或医院耐药监测数据酌情用药。', bold=True)
    return rt


# 在数据库中搜索表格信息
def find_info(
    result_list: list,
    sample_code: str,
    Interpretation: pd.DataFrame,
    info_client: pd.DataFrame,
    medical_DB: pd.DataFrame,
    all_bac: list,
    mic_dict: dict,
    supplementary_results:int
) -> list:
    bac_list: list = []
    sample_result_list: list = []
    column_P = info_client.loc[info_client['样本编号'] == str(sample_code), 'barcode'].iloc[0] + '_P_' + str(project_shorthand(sample_code)) + '_' + str(info_client.loc[info_client['样本编号'] == str(sample_code), '患者姓名'].iloc[0]) + '_' + str(sample_code)
    column_R = info_client.loc[info_client['样本编号'] == str(sample_code), 'barcode'].iloc[0] + '_R_' + str(project_shorthand(sample_code)) + '_' + str(info_client.loc[info_client['样本编号'] == str(sample_code), '患者姓名'].iloc[0]) + '_' + str(sample_code)
    for bac_name in result_list:
        # print(bac_name)
        # print(column_P)
        compare_bac_name: str = Nor(bac_name)
        dic_bac: dict ={}
        try:
            dic_bac['中文名'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]
        except IndexError:
            dic_bac['中文名'] = 'NA'
        try:
            dic_bac['分类'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '分类'].iloc[0]
            try:
                dic_bac['分类顺序'] = mic_dict[dic_bac['分类']]
            except KeyError:
                dic_bac['分类顺序'] = 0   
        except IndexError:
            dic_bac['分类'] = 'NA'
            dic_bac['分类顺序'] = 0
        if dic_bac['分类'] not in ['DNA病毒','RNA病毒','病毒','真菌','细菌']:
            dic_bac['分类'] = '其他病原'
            try:
                dic_bac['分类顺序'] = mic_dict[dic_bac['分类']]
            except KeyError:
                dic_bac['分类顺序'] = 0   
        try:
            dic_bac['相对丰度'] = '%.2f'%float(Interpretation.loc[Interpretation['name'] == compare_bac_name, column_P].iloc[0])
        except KeyError:
            column_P = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_P_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
            dic_bac['相对丰度'] = '%.2f'%float(Interpretation.loc[Interpretation['name'] == compare_bac_name, column_P].iloc[0])
        except IndexError:
            dic_bac['相对丰度'] = 'NA'
        try:
            try:
                dic_bac['序列数'] = int(Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0].replace('*',''))
            except:
                dic_bac['序列数'] = int(Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0])
        except KeyError:
            try:
                column_R = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_R_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
                dic_bac['序列数'] = int(Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0].replace('*',''))
            except:
                column_R = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_R_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
                dic_bac['序列数'] = int(Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0])
        except IndexError:
            dic_bac['序列数'] = 'NA'
        report_name = bac_name.replace("[","")
        dic_bac['微生物'] = report_name.replace("]","")
        try:
            dic_bac['备注'] = medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '备注'].iloc[0]
        except IndexError:
            dic_bac['备注'] = 'NA'
        try:
            bac_list.append(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '分类'].iloc[0])
        except IndexError:
            pass
        sample_result_list.append(dic_bac)
    if supplementary_results == 1:
        return sample_result_list
    bac_list = list(set(bac_list))
    if 'DNA病毒' in bac_list or 'RNA病毒' in bac_list:
        bac_list.append('病毒')
    for non_bac in all_bac:
        if non_bac not in bac_list:
            dic_bac = {}
            dic_bac['分类'] = non_bac
            dic_bac['分类顺序'] = mic_dict[non_bac]
            dic_bac['相对丰度'] = '--'
            dic_bac['序列数'] = '--'
            dic_bac['微生物'] = '--'
            dic_bac['备注'] = '--'
            dic_bac['中文名'] = '--'
            sample_result_list.append(dic_bac)
    sample_result_list: list = sorted(sample_result_list, key = lambda x:x['分类顺序'])
    return sample_result_list


def make_word_report(sample_code: str,
    dic_client: dict,
    report_dic: dict,
    doc_dic: dict,
    medical_DB: pd.DataFrame,
    value:str,
    complex_df:pd.DataFrame,
    appendix_species:pd.DataFrame,
    filename_date:str,
    formal_report,
    supplementary_report
    ) -> None:
    # print(sample_code)
    # print(str(dic_client[sample_code]['姓名']))
    project_shortname = project_shorthand(sample_code)
    #print(project_shortname)
    suffix = dic_client[sample_code]['检测项目'].split('测')[-1]
    #print('suffix',suffix)
    result_report_name = filename_date + '_'+ sample_code + '_' + str(dic_client[sample_code]['姓名'])+'_RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位'+ suffix + '.docx'
    manufacturer = dic_client[sample_code]['代理商']
    sample_result_name_list = Nor(dic_client[sample_code]['检测微生物']).split(",")
    project = info_client.loc[info_client['样本编号'] == sample_code, '检测项目'].iloc[0]
    save_path = os.path.join(args.output_dir, manufacturer,filename_date)
    if save_path and not os.path.exists(save_path):
        os.makedirs(save_path)
    if dic_client[sample_code]['代理商'] == 'seegene' or 'boruilin' or 'beijing':
        if 'DNA' in suffix and 'RNA' in suffix:
            df_name = report_dic[project_shortname][1] + '.xlsx'
            open_name = doc_dic[project_shortname][1] + '.docx'
            # print(dic_client[sample_code]['姓名'])
            # print(df_name)
            # print(open_name)
        elif 'RNA' in suffix:
            df_name = report_dic[project_shortname][2] + '.xlsx'
            open_name = doc_dic[project_shortname][2] + '.docx'
        else:
            df_name = report_dic[project_shortname][0] + '.xlsx'
            open_name = doc_dic[project_shortname][0] + '.docx'

    # print(df_name)
    open_docx_path = os.path.join(args.word_template_folder,manufacturer,open_name)
    if open_docx_path and os.path.exists(open_docx_path):
        doc = DocxTemplate(open_docx_path)
    #doc=DocxTemplate('/home/MicroReports/test/RD-seq/piplines/报告模板/word/seegene/RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位'+suffix+'.docx')
    doc.render(value)
    png_name_length = dic_client[sample_code]['测序长度图片文件名']
    png_name_length2 = dic_client[sample_code]['测序长度图片文件名单独']
    png_name_depth  = dic_client[sample_code]['测序深度图片文件名']
    handle_df = info_client.query('样本编号 == @sample_code').iloc[0,:]
    Target = formal_report
    # print('Target',Target)
    if Target == 'NA':
        pass
    elif Target == ',' or Target == '，':
        if os.path.exists(dic_client[sample_code]['测序长度图片文件名']):
                    try:
                            # print(dic_client[key]['测序长度图片文件名单独'])
                        doc = move_picture_length(doc=doc,png_name=png_name_length, png_name2 =png_name_length2, png_name_depth=png_name_depth)
                    except IndexError:
                        print('err2')
    else:
        #print(info_client['检测项目'])
        #if '耐药基因' in info_client['检测项目']:
        if os.path.exists(dic_client[sample_code]['测序长度图片文件名']) and os.path.exists(dic_client[sample_code]['测序深度图片文件名']):
            # print(dic_client[key]['测序长度图片文件名'])
            # print(dic_client[key]['测序深度图片文件名'])
            try:
                doc = move_picture_length_depth(doc=doc,png_name_length=png_name_length, png_name_depth=png_name_depth, png_name_length2 =png_name_length2)
            except IndexError:
                print('err1')


    if '耐药基因' in project:
            table_number = 4
            # table_number = 5
            #result_report_name = "{}RDseq-呼吸感染症候群基因检测+耐药基因鉴定.docx".format(value['姓名'])
            #result_report_name = filename_date + '_'+ sample_code + '_' + str(dic_client[sample_code]['姓名'])+'_RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位'+ suffix + '.docx'
            logging.info(f"{result_report_name}开始生成")
            number = 1
            appendix_color_change(doc=doc,table_list=dic_client[sample_code]['表8信息'],number=number)
            number += 1
            appendix_color_change(doc=doc,table_list=dic_client[sample_code]['表9信息'],number=number)
            color_change(doc=doc,table_number=table_number)
            color_change2(doc=doc,sample_code=sample_code,info_client=info_client,medical_DB=medical_DB,complex_df=complex_df,appendix_species=appendix_species, formal_report=formal_report, supplementary_report=supplementary_report)
            color_change3(doc=doc)
    #print(sample_result_name_list)
    if Nor(dic_client[sample_code]['补充微生物']) != 'na':
        supply_result_name_list = Nor(dic_client[sample_code]['补充微生物']).split(",")
        sample_result_name_list.extend(supply_result_name_list)
    # sample_result_name_list = change_bacteria_list(sample_result_name_list)
    if 'na' in sample_result_name_list:
        sample_result_name_list.remove('na')
    excel_reader = pd.ExcelFile(os.path.join(args.excel_template_folder,df_name))
    sheet_name_list = excel_reader.sheet_names
    # print(sheet_name_list)
    df1 = excel_reader.parse(sheet_name=sheet_name_list[0])
    try:
        df2 = excel_reader.parse(sheet_name=sheet_name_list[1])
    except IndexError:
        df2 = []
    try:
        df3 = excel_reader.parse(sheet_name=sheet_name_list[2])
    except IndexError:
        df3 = []   
    try:
        df4 = excel_reader.parse(sheet_name=sheet_name_list[3])
    except IndexError:
        df4 = []
    logging.info(f"{result_report_name}开始生成")
    for result_name in sample_result_name_list:
        try:
            chinese_name = medical_DB.loc[(medical_DB['英文名称']== Nor(result_name)) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '种'].iloc[0]
            if chinese_name != 'NA':
                genus_name = medical_DB.loc[(medical_DB['英文名称']== Nor(result_name)) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '属'].iloc[0]
                if genus_name != 'NA' or '病毒' in chinese_name:
                    micro_type = medical_DB.loc[(medical_DB['英文名称']== Nor(result_name)) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '类型'].iloc[0]
                    classification = medical_DB.loc[(medical_DB['英文名称']== Nor(result_name)) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '分类'].iloc[0]
                    pathogenicity_info = medical_DB.loc[(medical_DB['英文名称']== Nor(result_name)) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '备注'].iloc[0]
                    if pathogenicity_info == '人体共生菌':
                        col_numbers = 3
                        df4 = add_micro(df4, chinese_name, genus_name, micro_type, pathogenicity_info, col_numbers, classification)
                    elif classification == sheet_name_list[0]:
                        col_numbers = 4
                        df1 = add_micro(df1, chinese_name, genus_name, micro_type, pathogenicity_info, col_numbers, classification)
                    elif classification == sheet_name_list[1]:
                        col_numbers = 4
                        df2 = add_micro(df2, chinese_name, genus_name, micro_type, pathogenicity_info, col_numbers, classification)              
                    else:
                        col_numbers = 4
                        df3 = add_micro(df3, chinese_name, genus_name, micro_type, pathogenicity_info, col_numbers, classification)
                else:
                    logging.info(f"生成{result_report_name}时发现，数据库中{chinese_name}没有属名,请手动添加")
            else:
                logging.info(f"生成{result_report_name}时发现，数据库中{result_name}没有中文名,请手动添加")
        except IndexError:
            if result_name != 'NA':
                logging.info(f"生成{result_report_name}时发现，数据库中不存在{result_name},或数据库中{result_name}检测项目信息有问题")
    df1['致病性'] = df1['致病性'].map(lambda x: x.replace("条件致病菌","条件致病"))
    try:#去除细菌中已经报告的分枝杆菌信息
        # print('df1',df1['细菌'])
        df1=df1.drop(df1[df1['细菌'].str.contains('分枝杆菌')].index)
        # print('df1',df1['细菌'])
    except:
        pass
    use_col_width_dic_list = col_width_dic_list
    if type(df2) != list:
        df2['致病性'] = df2['致病性'].map(lambda x: x.replace("条件致病菌","条件致病")) 
    if type(df3) != list:
        df3['致病性'] = df3['致病性'].map(lambda x: x.replace("条件致病菌","条件致病"))
    if type(df4) == list:
        if type(df3) == list:
            if type(df2) == list:
                appendix_list = [df1]
            else:
                appendix_list = [df1,df2]
        else:
            appendix_list = [df1,df2,df3]
    else:
        appendix_list = [df1,df2,df3,df4]
    # # print('appendix_list[0]',type(appendix_list),appendix_list[0]['细菌'])
    # print('df1',df1)
    # print('df1',df1['细菌'])
    use_expect_text_list_dict = expect_text_list_dict
    if 'RNA' in suffix and 'DNA' not in suffix:
        if manufacturer=='seegene':
            use_expect_text_list_dict = {'seegene':['RNA常见病毒筛查范围']}
            use_col_width_dic_list = [{0: Cm(2.0), 1: Cm(3.5), 2: Cm(2), 3: Cm(1.4), 4: Cm(2.0), 5: Cm(3.5), 6: Cm(2), 7: Cm(1.4)}]
        if manufacturer=='boruilin':
            use_expect_text_list_dict = {'boruilin':['RNA常见病毒筛查范围']}
            use_col_width_dic_list = [{0: Cm(2.0), 1: Cm(3.5), 2: Cm(2), 3: Cm(1.4), 4: Cm(2.0), 5: Cm(3.5), 6: Cm(2), 7: Cm(1.4)}]
        if manufacturer=='beijing':
            use_expect_text_list_dict = {'beijing':['RNA常见病毒筛查范围']}
            use_col_width_dic_list = [{0: Cm(2.0), 1: Cm(3.5), 2: Cm(2), 3: Cm(1.4), 4: Cm(2.0), 5: Cm(3.5), 6: Cm(2), 7: Cm(1.4)}]
    for index,df in enumerate(appendix_list):
        # print(index,df)
        if (df.shape[0]) % 2 != 0 or df.iloc[int((df.shape[0])/2),1][0].isspace():
            try:
                df.loc[df.shape[0]]=['--'] * df.shape[1]
            except:
                pass
        row_num = int((df.shape[0])/2) + 1
        col_num = 8
        if index == 3:
            col_num = 6
        table = doc.add_table(rows=row_num, cols=col_num)

        tabBgColor(table, col_num, '#1F4E79')
        col_width_dic = use_col_width_dic_list[index]
        table = change_type(df,table,col_width_dic)
        # for row in range(row_num):
        #     for col in range(col_num):
        #         print('table',table.cell(row,col).paragraphs[0].text)
        for i in range(col_num):#将第一列的边框底部修改#第一列字体进行白色加粗
            #Set_cell_border(table.cell(row, col), bottom={"sz": 24, "val": "single", "color": "#FFD965"})
            Set_cell_border(table.cell(0, i), bottom={"sz": 24, "val": "single", "color": "#FFD965"})
            #run = table.cell(0,i).paragraphs[0]
            run = table.cell(0,i).paragraphs[0]
            content = run.text
            run.text = ''
            run = run.add_run(content)
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.bold = True
            run = table.cell(0,i).paragraphs[0]
            # print(run.text)
        # print(result_report_name)
        # print(table)
        target = None
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            # print(paragraph_text)
            # print(use_expect_text_list_dict[manufacturer][index])
            if paragraph_text.endswith(use_expect_text_list_dict[manufacturer][index]):
                #print(paragraph.text)
                target = paragraph
                break
        move_table_after(table, target)
        #
    doc.render(dic_client[sample_code])
    result_file = os.path.join(save_path, result_report_name)
    doc.save(result_file)
    doc = DocxTemplate(result_file)
    form_modification(doc=doc,dic_client=dic_client,sample_code=sample_code)
    if dic_client[sample_code]['检测微生物'] != ',' and dic_client[sample_code]['检测微生物'] != '，' and dic_client[sample_code]['检测微生物'] != 'NA':
        Interpretation_list = Interpretation_addition(sample_code=sample_code,dic_client=dic_client,doc=doc,insert_info=dic_client[sample_code]['注释'])
        parse_table_add(sample_code=sample_code,dic_client=dic_client,doc=doc,medical_DB=medical_DB,Interpretation_list=Interpretation_list)
    try:
        doc = move_picture_barcode(doc=doc, png_name=sample_code,barcode_picture_path=barcode_picture_path)
    except:
        logging.info(f"{handle_df['患者姓名']}的结果条形码图片未找到，请核对是否放入数据库中")
        pass
    doc.save(result_file)
    if result_file and os.path.exists(result_file):
        logging.info(f"{result_report_name}生成成功！")


# 检出微生物添加到附录中
def add_micro(df: pd.DataFrame,
    chinese_name: str,
    genus_name: str,
    micro_type: str,
    pathogenicity_info: str,
    col_numbers: int,
    classification: str
) -> pd.DataFrame:
    appendix_list = df.iloc[:, 1].apply(Nor).tolist()
    # print(appendix_list)
    # print(Nor(chinese_name))
    # print(Nor(genus_name))
    if (Nor(chinese_name) in appendix_list) and (Nor(genus_name) in appendix_list):
        micro_row = df[df.iloc[:, 1].apply(Nor)==chinese_name].index.tolist()[0]
        genus_row = df[df.iloc[:, 1].apply(Nor)==genus_name].index.tolist()[0]
        df.loc[micro_row, '结果'] = '检出'
        df.loc[genus_row, '结果'] = '检出'
    elif (Nor(genus_name) in appendix_list):
        genus_row = df[df.iloc[:, 1].apply(Nor)==genus_name].index.tolist()[0]
        df1 = df.loc[:genus_row]
        df2 = df.loc[(genus_row+1):]          
        if col_numbers == 3:
            df3 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:['  '+chinese_name],df.columns.tolist()[2]:['检出']})
        else:
            df3 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:['  '+chinese_name],df.columns.tolist()[2]:[pathogenicity_info],df.columns.tolist()[3]:['检出']}) 
        df = df1.append(df3, ignore_index = True).append(df2, ignore_index = True)
        df.loc[genus_row, '结果'] = '检出'
    elif '病毒' not in classification:
        if col_numbers == 3:
            df4 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:[genus_name],df.columns.tolist()[2]:['检出']})
            df5 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:['  '+chinese_name],df.columns.tolist()[2]:['检出']})            
        else:
            df4 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:[genus_name],df.columns.tolist()[2]:[pathogenicity_info],df.columns.tolist()[3]:['检出']})
            df5 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:['  '+chinese_name],df.columns.tolist()[2]:[pathogenicity_info],df.columns.tolist()[3]:['检出']})
        df = df.append(df4, ignore_index = True).append(df5, ignore_index = True)
    if '病毒' in classification:
        if (Nor(chinese_name) in appendix_list):
            micro_row = df[df.iloc[:, 1].apply(Nor)==chinese_name].index.tolist()[0]
            df.loc[micro_row, '结果'] = '检出'
        else:
            if col_numbers == 3:
                df6 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:[chinese_name],df.columns.tolist()[2]:['检出']}) 
            else:
                df6 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:[chinese_name],df.columns.tolist()[2]:[pathogenicity_info],df.columns.tolist()[3]:['检出']})
            df = df.append(df6, ignore_index = True)
    return df

# 将excel中的数据框转化为docx中的表格
def change_type(df: pd.DataFrame,
    table,
    col_width_dic: dict
    ):
    for col in list(range(len(table.columns))):
        for row in list(range(len(table.rows))):
            # print(col)
            # print(row)
            table.cell(row, col).width = col_width_dic[col]
            if row == 0:
                try:
                    table.cell(row,col).text = list(df.columns)[col]
                    # print(table.cell(row,col).text)
                    change_format(table,row,col)
                    run = table.cell(row,col).paragraphs[0]
                    run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except IndexError:
                    cell_col_index = col - int(len(table.columns)/2)
                    table.cell(row,col).text = list(df.columns)[cell_col_index]
                    change_format(table,row,col)
                    run = table.cell(row,col).paragraphs[0]
                    run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # print(table.cell(row,col).text)
            
            else:
                cell_col_index = col
                cell_row_index = row - 1
                try:
                    table.cell(row,col).text = df.iloc[cell_row_index, cell_col_index]
                    run = table.cell(row,col).paragraphs[0]
                    # print(table.cell(row,col).text)
                    change_format(table,row,col)
                    if cell_col_index == 1 and run.text != '--' :
                        run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except IndexError:
                    cell_col_index -= int(len(table.columns)/2)
                    cell_row_index = row + len(table.rows) - 2
                    table.cell(row,col).text = df.iloc[cell_row_index, cell_col_index]
                    # print(table.cell(row,col).text)
                    change_format(table,row,col)
                    run = table.cell(row,col).paragraphs[0]
                    if cell_col_index == 1 and run.text != '--' :
                        run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Set_cell_border(table.cell(row,col),
            top={"sz": 12, "val": "single", "color": "#F2F2F2"},
            bottom={"sz": 12, "val": "single", "color": "#F2F2F2"},
            start={"sz": 12, "val": "single", "color": "#F2F2F2"},
            end={"sz": 12, "val": "single", "color": "#F2F2F2"})
    return table

# 添加表格框线
def Set_cell_border(cell: _Cell, **kwargs):#修改单元格边框
    """
    设置单元格边框函数
    使用方法:
    Set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    传入参数有cell, 即单元格；top指上边框；bottom指下边框；start指左边框；end指右边框。
    "sz"指线的粗细程度；"val"指线型，比如单线，虚线等；"color"指颜色，颜色编码可百度；
    "space"指间隔，一般不设置，设置的值大于0会导致线错开；"shadow"指边框阴影
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

# seegene报告修改表格
def form_modification(doc,
    dic_client: dict,
    sample_code: str):
    print(doc.tables[2])
    table = doc.tables[2]
    
    add_dict = {'number_1':[],'number_2':[],'number_3':[],'number_4':[],'number_5':[],'number_6':[]}
    for line_info in dic_client[sample_code]['表10信息']:
        if line_info['分类'] == '细菌':
            add_dict['number_2'].append(line_info)
        elif line_info['分类'] == '真菌':
            add_dict['number_3'].append(line_info)
        elif line_info['分类'] == '病毒' or line_info['分类'] == 'DNA病毒' or line_info['分类'] == 'RNA病毒':
            add_dict['number_4'].append(line_info)
        elif line_info['分类'] == '古菌' or line_info['分类'] == '其他病原':
            add_dict['number_5'].append(line_info)
        elif line_info['分类'] == '寄生虫':
            add_dict['number_6'].append(line_info)
    number = 1
    while number < 7:
        table_info_name = f'表{number}信息'
        table_info = dic_client[sample_code][table_info_name]
        change_color_table = doc.tables[2+number]
        if table_info[0]['微生物'] != '--':
            for i,line in enumerate(table_info):
                row = i+2
                col = len(change_color_table.columns)
                #print(change_color_table(row,col).text)
                for col_number in range(col):
                    run = change_color_table.cell(row,col_number).paragraphs[0]
                    content = run.text
                    run.text = ''
                    run = run.add_run(content)
                    run.font.color.rgb = RGBColor(255,0,0)
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    if col_number == 1:
                        run.italic = True              
        else:
            run = change_color_table.cell(2,1).paragraphs[0]
            run.text = ''
            run = run.add_run('--')
            run.font.color.rgb = RGBColor(0,0,0)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        add_key = f'number_{number}'
        table_info.extend(add_dict[add_key])
        result_list = []
        if len(table_info) != 1 or table_info[0]['微生物'] != '--':
            add_number = 0
            for input_dict in table_info:
                if input_dict['备注'] == '人体共生菌':
                    add_number += 1
                elif input_dict['中文名'] != 'NA' and input_dict['中文名'] != '--':
                    if input_dict in add_dict[add_key]:
                        result_list.append(input_dict['中文名'] + '（补充报告部分）')
                    else:
                        result_list.append(input_dict['中文名'])
                elif input_dict['中文名'] == 'NA' and input_dict['微生物'] != '--':
                    if input_dict in add_dict[add_key]:
                        result_list.append(input_dict['微生物'] + '（补充报告部分）')
                    else:
                        result_list.append(input_dict['微生物'])
            if len(result_list) != 0:
                result_info = "，".join(result_list)
                run = table.cell(number,1).paragraphs[0]
                run.text = ''
                run = run.add_run(result_info)
                run.font.color.rgb = RGBColor(255,0,0)
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            if add_number != 0:
                p = table.cell(number,1).paragraphs[0]
                if p.text != '未检出疑似病原体':
                    run = p.add_run('，疑似微生态菌群')
                else:
                    p.text = ''
                    run = p.add_run('疑似微生态菌群')
                run.font.color.rgb = RGBColor(255,0,0)
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')    
        number += 1
    table_info =  dic_client[sample_code]['表7信息']
    if table_info[0]['基因'] != '--':
        for infp_dict in table_info:
            result_list.append(infp_dict['基因'])
        result_info = ",".join(result_list)
        run = table.cell(number,1).paragraphs[0]
        run.text = ''
        run = run.add_run(result_info)
        run.font.color.rgb = RGBColor(255,0,0)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.italic = True   
    number += 1
    result_list2 = []#需要单独开一个，不然和普通耐药累加
    table_info =  dic_client[sample_code]['表8信息']
    table_info.extend(dic_client[sample_code]['表9信息'])
    # print(table_info)
    if len(table_info) != 2 or table_info[0]['基因'] != '--' or table_info[1]['基因'] != '--':
        for infp_dict in table_info:
            if infp_dict['基因'] != '--' and infp_dict['基因'] not in result_list2:
                result_list2.append(infp_dict['基因'])
        result_info = ",".join(result_list2)
        #print('result_info',result_info)
        run = table.cell(number,1).paragraphs[0]
        run.text = ''
        run = run.add_run(result_info)
        run.font.color.rgb = RGBColor(255,0,0)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.italic = True



# 解读数量确认并添加
def Interpretation_addition(sample_code: str,
    dic_client: dict,
    doc,
    insert_info: str):
    dic_number: dict = {1:'一', 2:'二', 3:'三', 4:'四', 5:'五', 6:'六', 7:'七', 8:'八', 9:'九', 10:'十', 11:'十一', 12:'十二', 13:'十三', 14:'十四', 15:'十五', 16:'十六', 17:'十七', 18:'十八'}
    bac_list: list = dic_client[sample_code]['检测微生物'].split(',')
    Interpretation_list = []
    for i,bac in enumerate(bac_list):
        Interpretation_list.append(f'解读{dic_number[i + 1]}')
    for Interpretation in Interpretation_list:
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            if paragraph_text.endswith(insert_info):
                prior_paragraph = paragraph.insert_paragraph_before(Interpretation)
    return Interpretation_list
    
# word中改变首行背景
def tabBgColor(table,cols,colorStr):
    shading_list = locals()
    for i in range(cols):
        shading_list['shading_elm_'+str(i)] = parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'),bgColor = colorStr))
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_'+str(i)])
# seegene模板添加解析表格
def parse_table_add(sample_code: str,
    dic_client: dict,
    doc,
    medical_DB: pd.DataFrame,
    Interpretation_list: list):
    dic_number: dict = {1:'一', 2:'二', 3:'三', 4:'四', 5:'五', 6:'六', 7:'七', 8:'八', 9:'九', 10:'十', 11:'十一', 12:'十二', 13:'十三', 14:'十四', 15:'十五', 16:'十六', 17:'十七', 18:'十八'}
    bac_list: list = dic_client[sample_code]['检测微生物'].split(',')
    # bac_list = change_bacteria_list(bac_list)
    for bac_name in reversed(bac_list):
        if len(bac_name) != 0:
            number = dic_number[bac_list.index(bac_name) + 1]
            compare_bac_name = Nor(bac_name)
            table = doc.add_table(rows=3, cols=2)
            tabBgColor(table, 2, '#1F4E79')
            Interpretation = Interpretation_list[bac_list.index(bac_name)]
            # print(Interpretation)
            for col in list(range(len(table.columns))):
                for row in list(range(len(table.rows))):
                    if row == 0:
                        if col == 0: 
                            table.cell(row, col).width = Cm(2.8)
                            try:
                                if medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0] != 'NA':
                                    info = number + '、' + medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0].strip() + '(' + bac_name + ')'
                                else:
                                    info = number + '、' + bac_name
                            except IndexError:
                                info = number + '、' + bac_name
                            run = table.cell(row,col).paragraphs[0]
                            run = run.add_run(info)
                            run.font.color.rgb = RGBColor(255,255,255)
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        Set_cell_border(table.cell(row, col), bottom={"sz": 24, "val": "single", "color": "#FFD965"})
                        table.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    elif row == 1:
                        if col == 0: 
                            table.cell(row, col).width = Cm(2.8)
                            info = '临床意义'
                            run = table.cell(row,col).paragraphs[0]
                            run = run.add_run(info)
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            run.font.bold = True
                        elif col == 1: 
                            table.cell(row, col).width = Cm(14.7)
                            try:
                                info = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '临床意义'].iloc[0]
                            except IndexError:
                                info = 'NA'
                            table.cell(row,col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            run = table.cell(row,col).paragraphs[0]
                            run = run.add_run(info)
                            run.font.size = Pt(10.5)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        Set_cell_border(table.cell(row, col), bottom={"sz": 4, "val": "single", "color": "#000000"})
                        table.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    elif row == 2:
                        if col == 0: 
                            table.cell(row, col).width = Cm(2.8)
                            info = '常用药物'
                            run = table.cell(row,col).paragraphs[0]
                            run = run.add_run(info)
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            run.font.bold = True
                        elif col == 1: 
                            table.cell(row, col).width = Cm(14.7)
                            try:
                                info = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '常用药物'].iloc[0]
                            except IndexError:
                                info = 'NA'
                            table.cell(row,col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            run = table.cell(row,col).paragraphs[0]
                            run = run.add_run(info)
                            run.font.size = Pt(10.5)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        Set_cell_border(table.cell(row, col), bottom={"sz": 4, "val": "single", "color": "#000000"})
                        table.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.rows[0].height = Cm(1)
            table.rows[1].height = Cm(3.8)
            table.rows[2].height = Cm(6)
            table.cell(0,0).merge(table.cell(0,1))
            target = None
            for paragraph in doc.paragraphs:
                # print(paragraph.text)
                paragraph_text = paragraph.text
                if paragraph_text.endswith(Interpretation):
                    target = paragraph
                    break
            move_table_after(table, target)
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            if paragraph_text.endswith(Interpretation):
                for run in paragraph.runs:
                    run.text=run.text.replace(Interpretation,'')
                new_paragraph = target
# 定义表格的文字特征
def change_format(table,
    row:int,
    col:int
    ):
    run = table.cell(row,col).paragraphs[0]
    content = run.text
    run.text = ''
    run = run.add_run(content)
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')    
    if run.text == '检出':
        run.font.color.rgb = RGBColor(255,0,0)
        run = table.cell(row,col-1).paragraphs[0]
        content = run.text
        run.text = ''
        run = run.add_run(content)
        run.font.color.rgb = RGBColor(255,0,0)
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  
        run = table.cell(row,col-2).paragraphs[0]
        content = run.text
        run.text = ''
        run = run.add_run(content)
        run.font.color.rgb = RGBColor(255,0,0)
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')         
        if len(table.columns) != 6:
            run = table.cell(row,col-3).paragraphs[0]
            content = run.text
            run.text = ''
            run = run.add_run(content)
            run.font.color.rgb = RGBColor(255,0,0)
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑') 
def read_report_species(
    info_client,
    medical_DB
    ):
    dic_report = {}
    for row_index in range(info_client.shape[0]):
        if '通用引物' in info_client['备注'][row_index]:
            sample_code = info_client['样本编号'][row_index].split('-')[0]
            Sample_name = info_client['患者姓名'][row_index]
            Sample_barcode = info_client['barcode'][row_index]
            print(sample_code,Sample_name,Sample_barcode)
            # print(project_shortname)
            #这里需要根据有*号的数据分别添加进对应患者姓名及编号的的字典，然后进行阴性，正式，补充报告
            Sample_info = Sample_name+'_'+sample_code
            # print('Sample_info', Sample_info)
            dic_report.setdefault(Sample_info.replace(' ',''), {})
    # print('dic_report', dic_report)
    for number in range(Interpretation.shape[0]):#遍历每一行标有*号的#再次输出归属列的表头
        for number2 in range(Interpretation.shape[1]):
            if   r'***' in str(Interpretation.iloc[number,number2]):#符合不同标准的*号进行对应患者信息的正式报告和补充报告
                Name_id=('%s_%s'%(Interpretation.iloc[number].axes[0][number2].split('_')[3], Interpretation.iloc[number].axes[0][number2].split('_')[4])).replace('-s', '')
                Species = Interpretation.iloc[number].iloc[0]
                # print('阴性报告', Site, Species, Name_id)
                # dic_report[Sample_info]['阴性报告'] = [Species]
                if '阴性报告' in dic_report[Name_id.replace(' ','')].keys():
                    Species2 = dic_report[Name_id.replace(' ','')]['阴性报告']
                    Species2.append(Species)
                    dic_report[Name_id.replace(' ','')]['阴性报告']=Species2
                else:
                    dic_report[Name_id.replace(' ','')]['阴性报告'] = [Species]
            elif r'**' in str(Interpretation.iloc[number,number2]):#符合不同标准的*号进行对应患者信息的正式报告和补充报告
                Name_id=('%s_%s'%(Interpretation.iloc[number].axes[0][number2].split('_')[3], Interpretation.iloc[number].axes[0][number2].split('_')[4])).replace('-s', '')
                Species = Interpretation.iloc[number].iloc[0]
                # print('补充报告', Site, Species, Name_id)
                # dic_report[Sample_info]['补充报告'] = [Species]
                if '补充报告' in dic_report[Name_id.replace(' ','')].keys():
                    Species2 = dic_report[Name_id.replace(' ','')]['补充报告']
                    Species2.append(Species)
                    dic_report[Name_id.replace(' ','')]['补充报告']=Species2
                else:
                    dic_report[Name_id.replace(' ','')]['补充报告'] = [Species]
            elif r'*' in str(Interpretation.iloc[number,number2]):#符合不同标准的*号进行对应患者信息的正式报告和补充报告
                Name_id=('%s_%s'%(Interpretation.iloc[number].axes[0][number2].split('_')[3], Interpretation.iloc[number].axes[0][number2].split('_')[4])).replace('-s', '')
                Species = Interpretation.iloc[number].iloc[0]
                # print('正式报告', Site, Species, Name_id)
                # dic_report[Sample_info]['正式报告'] = [Species]
                if '正式报告' in dic_report[Name_id.replace(' ','')].keys():
                    Species2 = dic_report[Name_id.replace(' ','')]['正式报告']
                    Species2.append(Species)
                    dic_report[Name_id.replace(' ','')]['正式报告']=Species2
                else:
                    dic_report[Name_id.replace(' ','')]['正式报告'] = [Species]
    # print('dic_report', dic_report)#需要在这里进行不同报告物种的排序，以及同一物种多个菌的多序列降序排列
    for key,value in dic_report.items():
        # print(key,value)
        for keys,values in value.items():
            if '阴性报告' in value.keys():#阴性报告无需进行排序分类
                pass
            else:#此处进行正式报告和补充报告的物种分类及序列降排序
                if '正式报告' in value.keys():#阴性报告无需进行排序分类
                    if len(value['正式报告'])==1:#只有一个菌不用排序#大写替换
                        New_one=value['正式报告']
                        New_one=list(map(lambda x: x[0].upper() + x[1:].lower(), New_one))#此处调整拉丁文大写
                        dic_report[key]['正式报告']=New_one
                    else:
                        # print('正式报告长度', key, len(value['正式报告']))
                        # print('???????????????', key,keys,values)
                        all_kinds_list = microbial_classification(bacteria_list=value['正式报告'],medical_DB=medical_DB)
                        # print('all_kinds_list',all_kinds_list)
                        new_all_kinds_list=[]
                        for I in all_kinds_list:
                            # print('I'*100,I)
                            if len(I)==0:#分类后没有菌，不用进行排序
                                pass
                            elif len(I)==1:#分类后只有一个菌，不用进行排序
                                new_all_kinds_list.append(I[0])
                            else:
                                # print(I)#此处需要遍历寻找物种中名字
                                Sorting_sequence_number_list=[]
                                for II in I:
                                    Seqs=Sorting_sequence_number(Name_ID=key, Species_ID=II, Interpretation=Interpretation)
                                    # print('Seqs',Seqs)
                                    Sorting_sequence_number_list.append(Seqs)
                                Sorting_sequence_number_list.sort(key=lambda x:float(x[1]),reverse=True)###进行关键部分降序排列
                                Sorted_list_1=[i[0] for i in  Sorting_sequence_number_list]
                                # print('Sorting_sequence_number_list',Sorting_sequence_number_list)
                                # print('Sorted_list_1',Sorted_list_1)
                                #此处排序完成后进行字典键值替换
                                new_all_kinds_list.extend(Sorted_list_1)
                        # print('new_all_kinds_list',new_all_kinds_list)
                        new_all_kinds_list = list(map(lambda x: x[0].upper() + x[1:].lower(), new_all_kinds_list))#此处调整拉丁文大写
                        # print('new_all_kinds_list',new_all_kinds_list)
                        dic_report[key]['正式报告']=new_all_kinds_list

                if  '补充报告' in value.keys():#阴性报告无需进行排序分类
                    if len(value['补充报告'])==1:#只有一个菌不用排序
                        New_one=value['补充报告']
                        New_one=list(map(lambda x: x[0].upper() + x[1:].lower(), New_one))#此处调整拉丁文大写
                        dic_report[key]['补充报告']=New_one
                    else:
                        # print('补充报告长度', key, len(value['补充报告']))
                        # print('！！！！！！！！', key,keys,values)
                        all_kinds_list2 = microbial_classification(bacteria_list=value['补充报告'],medical_DB=medical_DB)
                        # print(all_kinds_list2)
                        new_all_kinds_list2=[]
                        for J in all_kinds_list2:
                            if len(J)==0:#分类后只有一个菌，不用进行排序
                                pass
                            elif len(J)==1:#分类后只有一个菌，不用进行排序
                                new_all_kinds_list2.append(J[0])
                            else:
                                # print(J)#有多个菌进行序列数量提取并排序
                                Sorting_sequence_number_list2=[]
                                for JJ in J:
                                    Seqs2=Sorting_sequence_number(Name_ID=key, Species_ID=JJ,Interpretation=Interpretation)
                                    # print('Seqs2',Seqs2)
                                    Sorting_sequence_number_list2.append(Seqs2)
                                Sorting_sequence_number_list2.sort(key=lambda x:float(x[1]),reverse=True)###进行关键部分降序排列
                                # print('Sorting_sequence_number_list2',Sorting_sequence_number_list2)
                                Sorted_list_2=[i[0] for i in Sorting_sequence_number_list2]
                                # print(dic_report[key]['补充报告'])
                                # print('Sorted_list_2',Sorted_list_2)
                                new_all_kinds_list2.extend(Sorted_list_2)
                        # print('new_all_kinds_list2',new_all_kinds_list2)
                        new_all_kinds_list2 = list(map(lambda x: x[0].upper() + x[1:].lower(), new_all_kinds_list2))#此处调整拉丁文大写
                        dic_report[key]['补充报告']=new_all_kinds_list2

    dic_report_true = {}#存放修改为模仿之前报告格式的报告内容#什么都没有的判断都为NA，有阴性报告标识符的，判断正式报告和补充报告为，正式报告，补充报告有或没有为NA,进行数值填充，随后供文件输出报告判断
    for key,value in dic_report.items():
        if len(value)==0:#不进行报告输出
            dic_report_true.setdefault(key, {})
            dic_report_true[key]['正式报告']='NA'
            dic_report_true[key]['补充报告']='NA'
        elif '阴性报告' in value.keys():
            dic_report_true.setdefault(key, {})
            dic_report_true[key]['正式报告']=','
            dic_report_true[key]['补充报告']='NA'
        elif '正式报告' in value.keys():
            if  '补充报告' in value.keys():
                dic_report_true.setdefault(key, {})
                dic_report_true[key]['正式报告']=','.join(value['正式报告'])
                dic_report_true[key]['补充报告']=','.join(value['补充报告'])
            else:
                dic_report_true.setdefault(key, {})
                dic_report_true[key]['正式报告']=','.join(value['正式报告'])
                dic_report_true[key]['补充报告']='NA'
        elif '补充报告' in value.keys():
            if '正式报告' not in value.keys():
                dic_report_true.setdefault(key, {})
                dic_report_true[key]['补充报告']=','.join(value['补充报告'])
                dic_report_true[key]['正式报告']='NA'
    # print('dic_report_true', dic_report_true)
    return dic_report_true

def Sorting_sequence_number(
    Name_ID,
    Species_ID,
    Interpretation
    ):
    Species_ID=Species_ID.lower()
    for number2 in range(Interpretation.shape[1]):
        Title=(Interpretation.iloc[number2].axes[0])
        for Title_range in Title:
            # print(Title_range,Name_ID,Species_ID)
            if Title_range.replace(' ','').endswith(Name_ID):
                if '_R_' in Title_range:
                    # print(Interpretation.loc[Interpretation['name']==Species_ID,Title_range])
                    try:
                        Reads_Number=int(Interpretation.loc[Interpretation['name']==Species_ID,Title_range].iloc[0].replace('*',''))
                        return [Species_ID,Reads_Number]
                    except:
                        Reads_Number=int(Interpretation.loc[Interpretation['name']==Species_ID,Title_range].iloc[0])
                        return [Species_ID,Reads_Number]
                    
                    
                    # print('Title_range',Title_range)
                    # print('Name_ID',Name_ID)
                    # print('Species_ID',Species_ID)
def Removestar(X):
    X.replace('*','')
    return(X)
####################################################
# 统一输入的英文名
# print(medical_DB)
result_file_name = args.result_excel.split("/")[-1]
complex_df = pd.read_excel(args.complex_excel)
name_df = pd.read_excel(args.name_excel)
medical_DB['种'] = medical_DB['种'].apply(Nor_col)
medical_DB['英文名称'] = medical_DB['英文名称'].apply(Nor)
# print(medical_DB)
info_client['样本编号'] = info_client['样本编号'].apply(Nor_col)
info_client['患者姓名'] = info_client['患者姓名'].apply(Nor_col)
info_client['代理商'] = info_client['代理商'].apply(Nor_col)
Interpretation['name'] = Interpretation['name'].apply(Nor)
complex_df['name'] = complex_df['name'].apply(Nor)
info_client['采样时间'] = info_client['采样时间'].map(lambda x: str(x).split(' ')[0])

###############################################
# MTB和NTM的耐药基因和对应药物
MTB_gene_dict = {'rpoB':['利福平'],'embB':['乙胺丁醇'],'pncA':['吡嗪酰胺'],'katG':['异烟肼'],'inhA':['异烟肼','乙硫异烟胺','丙硫异烟胺'],'gyrA':['氟喹诺酮类'],'gyrB':['氟喹诺酮类'],'rpsL':['链霉素'],'rrs':['链霉素','阿米卡星','卡那霉素','卷曲霉素'],'folC':['对氨基水杨酸'],'thyA':['对氨基水杨酸'],'alr':['环丝氨酸'],'Rv0678':['氯法齐明','贝达喹啉'],'rplC':['利奈唑胺']}
NTM_gene_dict = {'rrl':['克拉霉素','阿奇霉素'],'erm':['克拉霉素','阿奇霉素'],'rrs':['阿米卡星','卡那霉素','庆大霉素']}
# 药物分类
first_list = ['利福平','异烟肼','吡嗪酰胺','乙胺丁醇']
second_list = ['阿米卡星','氟喹诺酮类','乙硫异烟胺','丙硫异烟胺','对氨基水杨酸','链霉素','卡那霉素','卷曲霉素','环丝氨酸','氯法齐明','贝达喹啉','利奈唑胺']
NTM_list = ['克拉霉素','阿奇霉素','阿米卡星','卡那霉素','庆大霉素']
# 附录中的种
appendix_species = [ '结核分枝杆菌', '非洲分枝杆菌', '牛分枝杆菌', '山羊分枝杆菌', '田鼠分枝杆菌', '卡内蒂分枝杆菌', '鳍脚分枝杆菌', '獴分枝杆菌', '鸟分枝杆菌', '胞内分枝杆菌', '副胞内分枝杆菌', '堪萨斯分枝杆菌', '龟分枝杆菌', '脓肿分枝杆菌', '猿猴分枝杆菌', '溃疡分枝杆菌', '偶发分枝杆菌', '玛尔摩分枝杆菌', '海分枝杆菌', '马赛分枝杆菌', '戈登分枝杆菌', '副戈登分枝杆菌', '产粘液分枝杆菌']
# 对应物种的基因组ID
id_dict = {'NC_000962.3':'Mycobacterium tuberculosis','CU458896.1':'Mycobacteroides abscessus','NR_025584.1':'Mycobacterium avium','NR_042165.1':'Mycobacterium intracellulare','NR_121712.2':'Mycobacterium kansasii','NR_114659.1':'Mycobacteroides chelonae','NR_025311.1':'Mycolicibacterium smegmatis','NR_042912.1':'Mycolicibacterium fortuitum','NC_010397.1':'Mycobaccterium abscessus','NG_041979.1':'Mycobacterium avium','NR_076151.1':'Mycobacterium intracellulare','NZ_CP019883.1':'Mycobacterium kansasii','GU143889.1':'Mycobacteroides chelonae','AB011184.1':'Mycolicibacterium smegmatis','NZ_CP011269.1:c3552322-3549198':'Mycolicibacterium fortuitum','NZ_CP014955.1:2353195-2353716':'Mycobacteroides abscessus'}
AMR_df = pd.read_excel(args.result_excel, sheet_name='drug_report').fillna('NA')
# 查询模板的字典
doc_dic = {
            #'RD': ['RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位'],
            'HX': ['RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位DNA','RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位DNA+RNA','RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位RNA','RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位']
}
# 查询结果名称的字典
report_dic = {
            #'RD': ['RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位'],
            'HX': ['RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位DNA','RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位DNA+RNA','RDseq-呼吸感染症候群基因检测+耐药基因鉴定-定位RNA']
}
# 调用模板的字典
doc_list = ['seegene','boruilin','beijing']

# 生成附录的对应位置和附录表格的列宽
expect_text_list_dict = {'beijing':['常见细菌筛查范围', '常见真菌筛查范围','常见非典型病原体筛查范围','常见人体共生菌筛查范围'],'boruilin':['常见细菌筛查范围', '常见真菌筛查范围','常见非典型病原体筛查范围','常见人体共生菌筛查范围'],'beagle':['常见细菌筛查范围', '常见真菌筛查范围','其他病原微生物','常见人体共生菌'],'seegene':['常见细菌筛查范围', '常见真菌筛查范围','常见非典型病原体筛查范围','常见人体共生菌筛查范围']}
col_width_dic_list = [{0: Cm(1.1), 1: Cm(4.4), 2: Cm(2), 3: Cm(1.4), 4: Cm(1.1), 5: Cm(4.4), 6: Cm(2), 7: Cm(1.4)}, {0: Cm(2.1), 1: Cm(3.4), 2: Cm(2), 3: Cm(1.4), 4: Cm(2.1), 5: Cm(3.4), 6: Cm(2), 7: Cm(1.4)},{0: Cm(2.4), 1: Cm(3.25), 2: Cm(1.8), 3: Cm(1.45), 4: Cm(2.4), 5: Cm(3.25), 6: Cm(1.8), 7: Cm(1.45)},{0: Cm(1.4), 1: Cm(5.1), 2: Cm(2.4), 3: Cm(1.4), 4: Cm(5.1), 5: Cm(2.4)}]
# 导入正文部分的字典生成
dic_client = {}
#在此处收集需要收集每个人的报告结果、
dic_report=read_report_species(info_client=info_client, medical_DB=medical_DB)##此处需要修改标注的引物类型
# print('dic_report_true',dic_report)
for row_index in range(info_client.shape[0]):
    if '通用引物' in info_client['备注'][row_index]:
        Name_ID=info_client['患者姓名'][row_index].replace(' ','')+'_'+info_client['样本编号'][row_index].split('-')[0]
        formal_report=dic_report[Name_ID]['正式报告']
        supplementary_report=dic_report[Name_ID]['补充报告']
        if formal_report != 'NA' or supplementary_report != 'NA':
        # if info_client['正式报告结果'][row_index] != 'NA' or info_client['补充报告结果'][row_index] != 'NA':
        # for company in doc_list:
        #     if result_file_name.startswith(company):
        #         manufacturer = company
        # if info_client['代理商'][row_index].lower() in doc_list:
        #     manufacturer = info_client['代理商'][row_index].lower()
            if info_client['代理商'][row_index].lower() in doc_list:
                manufacturer = info_client['代理商'][row_index].lower()
            else:
                for company in doc_list:
                    if result_file_name.startswith(company):
                        manufacturer = company
            # print(manufacturer)
            dic_client.setdefault(info_client['样本编号'][row_index], {})
            dic_client[info_client['样本编号'][row_index]]['样本编号'] = info_client['样本编号'][row_index].split('-')[0]
            dic_client[info_client['样本编号'][row_index]]['代理商'] = manufacturer
            dic_client[info_client['样本编号'][row_index]]['姓名'] = info_client['患者姓名'][row_index]
            logging.info(f"开始处理{info_client['样本编号'][row_index]} {info_client['患者姓名'][row_index]}的信息")
            dic_client[info_client['样本编号'][row_index]]['性别'] = info_client['性别'][row_index]
            try:
                dic_client[info_client['样本编号'][row_index]]['年龄'] = int(info_client['年龄'][row_index])
            except ValueError:
                dic_client[info_client['样本编号'][row_index]]['年龄'] = info_client['年龄'][row_index]
            dic_client[info_client['样本编号'][row_index]]['送检单位'] = info_client['医院'][row_index]
            dic_client[info_client['样本编号'][row_index]]['科室'] = info_client['科室'][row_index]
            dic_client[info_client['样本编号'][row_index]]['检测项目'] = info_client['检测项目'][row_index].strip()
            dic_client[info_client['样本编号'][row_index]]['科室'] = info_client['科室'][row_index].strip()
            dic_client[info_client['样本编号'][row_index]]['降钙素原'] = info_client['降钙素原(PCT)'][row_index]
            dic_client[info_client['样本编号'][row_index]]['白细胞'] = info_client['白细胞(WBC)'][row_index]
            dic_client[info_client['样本编号'][row_index]]['反应蛋白'] = info_client['C-反应蛋白(CRP)'][row_index]
            dic_client[info_client['样本编号'][row_index]]['培养结果'] = info_client['培养结果'][row_index].strip()
            dic_client[info_client['样本编号'][row_index]]['中性粒细胞'] = str(info_client['中性粒细胞比率'][row_index]).strip()
            dic_client[info_client['样本编号'][row_index]]['淋巴细胞'] = str(info_client['淋巴细胞比率'][row_index]).strip()
            dic_client[info_client['样本编号'][row_index]]['样本类型'] = info_client['样本类型'][row_index].strip()
            dic_client[info_client['样本编号'][row_index]]['样本性状'] = info_client['样本性状'][row_index].strip()
            dic_client[info_client['样本编号'][row_index]]['样本颜色'] = info_client['样本颜色'][row_index].strip()
            dic_client[info_client['样本编号'][row_index]]['样本接收'] = info_client['收样异常情况'][row_index].strip()
            dic_client[info_client['样本编号'][row_index]]['临床诊断'] = info_client['临床诊断'][row_index]
            dic_client[info_client['样本编号'][row_index]]['抗感染用药史'] = info_client['抗感染用药史'][row_index].strip()
            dic_client[info_client['样本编号'][row_index]]['重点关注病原菌'] = info_client['重点关注病原菌'][row_index].strip()
            dic_client[info_client['样本编号'][row_index]]['检测微生物'] = formal_report
            dic_client[info_client['样本编号'][row_index]]['补充微生物'] = supplementary_report
            receivedate = str(info_client['收样时间'][row_index]).strip().split(' ')[0]
            reportdate = str(info_client['报告日期'][row_index]).split(' ')[0].strip()
            # print('reportdate'*100,reportdate)
            if receivedate != 'NA':
                dic_client[info_client['样本编号'][row_index]]['收样日期'] = receivedate.split('/')[0]+'.'+receivedate.split('/')[1]+'.'+receivedate.split('/')[2]
            else:
                dic_client[info_client['样本编号'][row_index]]['收样日期'] = 'NA'
            collectiondate = str(info_client['采样时间'][row_index]).strip()
            #print(collectiondate.split('/'))
            #print(reportdate.split('/'))
            if collectiondate != 'NA':
                try:
                    dic_client[info_client['样本编号'][row_index]]['采样日期'] = collectiondate.split('/')[0]+'.'+collectiondate.split('/')[1]+'.'+collectiondate.split('/')[2]
                except IndexError:
                    dic_client[info_client['样本编号'][row_index]]['采样日期'] = collectiondate.split('-')[0]+'.'+collectiondate.split('-')[1]+'.'+collectiondate.split('-')[2]
            if reportdate != 'NA':
                try:
                    dic_client[info_client['样本编号'][row_index]]['报告日期'] = reportdate.split('/')[0]+'.'+reportdate.split('/')[1]+'.'+reportdate.split('/')[2]
                except IndexError:
                    dic_client[info_client['样本编号'][row_index]]['报告日期'] = reportdate.split('-')[0]+'.'+reportdate.split('-')[1]+'.'+reportdate.split('-')[2]
            else:
                dic_client[info_client['样本编号'][row_index]]['采样日期'] = 'NA'
            try:
                exdate = str(info_client['上机日期'][row_index]).split(' ')[0]
                dic_client[info_client['样本编号'][row_index]]['上机日期'] = exdate.split('-')[0]+'.'+exdate.split('-')[1]+'.'+exdate.split('-')[2]
            except IndexError:
                exdate = str(info_client['上机日期'][row_index])
                dic_client[info_client['样本编号'][row_index]]['上机日期'] = exdate.split('/')[0]+'.'+exdate.split('/')[1]+'.'+exdate.split('/')[2]
            handle_df = info_client[info_client['样本编号'] == info_client['样本编号'][row_index]]
            pat_name = handle_df['患者姓名'].iloc[0]
            # print(handle_df['样本编号'].iloc[0])
            pat_sample_id_list=[]
            if '-s' in handle_df['样本编号'].iloc[0]:
                pat_sample_id_list.append(handle_df['样本编号'].iloc[0])
            else:
                pat_sample_id_list.append(handle_df['样本编号'].iloc[0]+'-s')
            pat_sample_id = pat_sample_id_list[0]
            # print(pat_name)
            # print(pat_sample_id)
            info_client_info2 = info_client[(info_client['患者姓名'] == pat_name) & (info_client['样本编号'] == pat_sample_id) & (info_client['备注'].str.contains('结核耐药'))]#寻找结核耐药barcode
            #print(info_client_info2)
            #print(info_client)
            colname = str(info_client_info2.loc[info_client['样本编号'] ==  pat_sample_id, 'barcode'].iloc[0]) + '_' + str(project_shorthand(info_client['样本编号'][row_index])) + '_' + str(info_client['患者姓名'][row_index]) + '_' + str(info_client['样本编号'][row_index])+'-s'
            Delete_short_name = '_' + str(project_shorthand(info_client['样本编号'][row_index]))
            colname2 = colname.replace(Delete_short_name,'')
            # print(colname)
            #print(picture_dtat_df.at[8, colname])
            try:
                dic_client[info_client['样本编号'][row_index]]['平均长度'] = picture_dtat_df.at[8, colname]
            except:
                dic_client[info_client['样本编号'][row_index]]['平均长度'] = picture_dtat_df.at[8, colname2]
            shortname = project_shorthand(info_client['样本编号'][row_index])
            #print(shortname)
            if type(shortname) != int:
                AMR_stat_colname = colname + '_count'
                #print(AMR_stat_colname)
                file_name = AMR_stat_colname.replace(" ", ".")
                #print(file_name)
            else:
                AMR_stat_colname = colname + '_count'
                AMR_stat_colname = AMR_stat_colname.replace('_RD','')
                file_name = AMR_stat_colname.replace(" ", ".")
            dic_client[info_client['样本编号'][row_index]]['测序深度图片文件名'] = file_name + '.png'
            #make_picture_depth(AMR_stat_colname=AMR_stat_colname,AMR_stat_df=args.result_excel,file_name=file_name)
            if formal_report != ',' or formal_report != '，' and '耐药基因' in info_client['检测项目'][row_index]:
                #dic_client[info_client['样本编号'][row_index]]['图片文件名'] = file_name + '.png'
                #print(AMR_stat_df)
                dic_client[info_client['样本编号'][row_index]]['测序深度'] = str(int(AMR_stat_df[AMR_stat_colname][11])) + '×'
                dic_client[info_client['样本编号'][row_index]]['覆盖率'] = str(float(format(AMR_stat_df[AMR_stat_colname][12], '.2f'))) + '%'
                make_picture_depth(AMR_stat_colname=AMR_stat_colname,AMR_stat_df=args.result_excel,file_name=file_name)#输入表中患者带有shortname，作图和识别也要带有shortname
            else:
                #dic_client[info_client['样本编号'][row_index]]['图片文件名'] = file_name + '.png'
                dic_client[info_client['样本编号'][row_index]]['测序深度'] = '--'
                dic_client[info_client['样本编号'][row_index]]['覆盖率'] = '--'
            if manufacturer == 'seegene' or "boruilin" or 'beijing':
                #make_picture_depth(AMR_stat_colname=AMR_stat_colname,AMR_stat_df=args.result_excel,file_name=file_name)
                #print(medical_DB)
                table_info = table2_context(sample_code=info_client['样本编号'][row_index], info_client=info_client, Interpretation=Interpretation, medical_DB=medical_DB, formal_report=formal_report, supplementary_report=supplementary_report)
                # print(medical_DB)
                result_table = table_info[0]
                dic_client[info_client['样本编号'][row_index]]['表1信息'] = result_table[0]
                dic_client[info_client['样本编号'][row_index]]['表2信息'] = result_table[1]
                dic_client[info_client['样本编号'][row_index]]['表3信息'] = result_table[2]
                dic_client[info_client['样本编号'][row_index]]['表4信息'] = result_table[3]
                dic_client[info_client['样本编号'][row_index]]['表5信息'] = result_table[4]
                dic_client[info_client['样本编号'][row_index]]['表6信息'] = result_table[5]
                dic_client[info_client['样本编号'][row_index]]['表10信息'] = table_info[1]
                table_info_TB = table_context_TB(sample_code=info_client['样本编号'][row_index], info_client=info_client, medical_DB=medical_DB, complex_df=complex_df, id_dict=id_dict, AMR_df=AMR_df, AMR_stat_colname=AMR_stat_colname, Interpretation=Interpretation,supplementary_report=supplementary_report,formal_report=formal_report)
                dic_client[info_client['样本编号'][row_index]]['表8信息'] = table_info_TB[0]
                dic_client[info_client['样本编号'][row_index]]['表9信息'] = table_info_TB[1]
                appendix_info = appendix_drugs(table2_list=table_info_TB[0],table3_list=table_info_TB[1],first_list=first_list,second_list=second_list,NTM_list=NTM_list)
                dic_client[info_client['样本编号'][row_index]]['一线检出'] = appendix_info[0]
                dic_client[info_client['样本编号'][row_index]]['一线未检出'] = appendix_info[1]
                dic_client[info_client['样本编号'][row_index]]['二线检出'] = appendix_info[2]
                dic_client[info_client['样本编号'][row_index]]['二线未检出1'] = appendix_info[3]
                dic_client[info_client['样本编号'][row_index]]['二线未检出2'] = appendix_info[4]
                dic_client[info_client['样本编号'][row_index]]['检出'] = appendix_info[5]
                dic_client[info_client['样本编号'][row_index]]['未检出'] = appendix_info[6]
                if type(drug_resistance_df) != int:
                    dic_client[info_client['样本编号'][row_index]]['表7信息'] = table7_make(sample_code=info_client['样本编号'][row_index], info_client=info_client, drug_resistance_df=drug_resistance_df)
                else:
                    dic_client[info_client['样本编号'][row_index]]['表7信息'] = [{'基因': '--', '药物': '--'}]
                dic_client[info_client['样本编号'][row_index]]['example'] = clinical(sample_code=info_client['样本编号'][row_index],info_client=info_client,medical_DB=medical_DB,manufacturer=manufacturer, formal_report=formal_report, supplementary_report=supplementary_report)
                #length_colname = str(info_client['barcode'][row_index]) + '_' + str(project_shorthand(info_client['样本编号'][row_index])) + '_' + str(dic_client[info_client['样本编号'][row_index]]['姓名']) + '_' + str(info_client['样本编号'][row_index])
                #print(length_colname)
                dic_client[info_client['样本编号'][row_index]]['测序长度图片文件名'] = colname2 + '_length.png'#测序长度名称不带有shortname表中长度不带有shortname
                dic_client[info_client['样本编号'][row_index]]['测序长度图片文件名单独'] = colname2 + '_length_single.png'#测序长度名称不带有shortname
                make_picture_length(length_colname=colname2,picture_dtat_df=picture_dtat_df)#长度使用没有shortname作图
                if formal_report != ',' and formal_report != '，' and formal_report != 'NA':
                    dic_client[info_client['样本编号'][row_index]]['注释'] = '注：常用药物为临床常规药物，且无法覆盖药敏结果，具体用药请结合临床药敏结果或医院耐药监测数据酌情用药。'
                else:
                    dic_client[info_client['样本编号'][row_index]]['注释'] = '本次样本中未检出疑似致病菌，结果仅对本次送检的样本负责，请临床根据患者症状结合其他检测结果进行综合判断。'
                if supplementary_report != ',' and supplementary_report != '，' and supplementary_report != 'NA':
                    dic_client[info_client['样本编号'][row_index]]['说明'] = '疑似微生物种解释说明'
logging.info("所有信息处理完成！")
#print(dic_client)
# print(reportdate)
try:
    filename_date = reportdate.split('/')[0]+reportdate.split('/')[1]+reportdate.split('/')[2]
except:
    filename_date = reportdate.split('-')[0]+reportdate.split('-')[1]+reportdate.split('-')[2]

# if __name__=='__main__':
#     p = Pool(args.processes_number)     
#     for sample_code,value in dic_client.items():
#         p.apply_async(make_word_report,args=(sample_code,dic_client,report_dic,doc_dic,medical_DB,value,complex_df,appendix_species,filename_date))
#     p.close()
#     p.join()
# logging.info(f"所有报告生成成功！")

for sample_code,value in dic_client.items():
    # print(value['姓名'])
    # print(value)
    Name_ID=value['姓名'].replace(' ','')+'_'+value['样本编号']
    #print('sample_code',sample_code)
    formal_report=dic_report[Name_ID]['正式报告']
    supplementary_report=dic_report[Name_ID]['补充报告']
    make_word_report(sample_code=sample_code,dic_client=dic_client,report_dic=report_dic,doc_dic=doc_dic,medical_DB=medical_DB,value=value,complex_df=complex_df,appendix_species=appendix_species,filename_date=filename_date, formal_report=formal_report, supplementary_report=supplementary_report)
logging.info(f"开始汇总结果！")
for sample_code,value in dic_client.items():
    #print('sample_code',sample_code)
    manufacturer = dic_client[sample_code]['代理商']
    summary_excel_path = os.path.join(args.summary_excel,manufacturer,'汇总.xlsx')
    #print('summary_excel_path',summary_excel_path)
    all_summary_df = pd.read_excel(summary_excel_path)
    all_info_list = []
    if all_summary_df.shape[0] != 0:
        for row_index in range(all_summary_df.shape[0]):
            line_info = str(all_summary_df['样本编号'][row_index]) + '_' + str(all_summary_df['患者姓名'][row_index]) + '_' + str(all_summary_df['basecalling data'][row_index]) + '_' + str(all_summary_df['Reports'][row_index])
            #print(line_info)
            all_info_list.append(line_info)
    all_info_dict = dict(zip(all_info_list,all_info_list))
    data_summary_excel_path = os.path.join(args.summary_excel,manufacturer,filename_date,'汇总.xlsx')
    if data_summary_excel_path and os.path.exists(data_summary_excel_path):
        data_summary_df = pd.read_excel(data_summary_excel_path)
    else:
        data_summary_df = pd.DataFrame(columns=all_summary_df.columns.tolist())
    
    need_info = info_client.query('样本编号 == @sample_code').iloc[0]
    all_list = []
    try:
        all_list = formal_report.split(',')    
    except AttributeError:
        all_list = []
    try:
        all_list.extend(supplementary_report.split(','))
    except AttributeError:
        all_list = all_list
    for micro in all_list:
        # print(micro)
        compare_bac_name = Nor(micro)
        data_info_list = []
        if data_summary_df.shape[0] != 0:
            for row_index in range(data_summary_df.shape[0]):
                line_info = str(data_summary_df['样本编号'][row_index]) + '_' + str(data_summary_df['患者姓名'][row_index]) + '_' + str(data_summary_df['basecalling data'][row_index]) + '_' + data_summary_df['Reports'][row_index]
                data_info_list.append(line_info)
        data_info_dict = dict(zip(data_info_list,data_info_list))
        Resisit_MT=['结核耐药-突变基因','结核耐药-潜在耐药药物','结核耐药-核酸突变结果','结核耐药-氨基酸突变','结核耐药-突变率',
                    '非结核耐药-突变基因','非结核耐药-潜在耐药药物','非结核耐药-核酸突变结果','非结核耐药-氨基酸突变','非结核耐药-突变率']
        Relist=['Reports','报告结果','报告结果','报告等级','物种','致病性描述','Reads','样本描述（解读人员手动添加）','备 注','采样时间.1']#存放不要的结果
        if len(micro) != 0 and micro != 'NA':
            info_list = all_summary_df.columns.tolist()[:-10]
            new_line_dict = {}
            for colname in info_list:
                if colname not in Relist:
                    new_line_dict[colname] = need_info[colname]
            new_line_dict['Reports'] = micro
            try:
                new_line_dict['报告结果'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]
            except IndexError:
                new_line_dict['报告结果'] = 'NA'
            try:
                new_line_dict['物种'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '分类'].iloc[0]
            except IndexError:
                new_line_dict['物种'] = 'NA'
            if micro in need_info['正式报告结果']:
                new_line_dict['报告等级'] = 'A'
            else:
                new_line_dict['报告等级'] = 'B'
            try:
                new_line_dict['致病性描述'] = medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '备注'].iloc[0]
            except IndexError:
                new_line_dict['致病性描述'] = 'NA'
            column_R = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_R_' + project_shorthand(sample_code) + '_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
            try:
                new_line_dict['Reads'] = Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0]
            except KeyError:
                column_R = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_R_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
                new_line_dict['Reads'] = Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0]
            except IndexError:
                new_line_dict['Reads'] = 'NA'
            new_line_dict['样本描述（解读人员手动添加）'] = ' '
            new_line_dict['制作人'] = ' '
            new_line_dict['审核人'] = ' '
            new_line_dict['备 注'] = ' '
            # print('表8信息',value['表8信息'])
            #print(value['表9信息'])
            MTgene=[];MTdrug=[];MTdescript=[];MTProtein=[];MTPrecent=[]
            for MTlist in value['表8信息']:
                #print(MTlist)
                for key8,value8 in MTlist.items():
                    # print(value['姓名'])
                    #print('key8,value8',key8,value8)
                    if key8=='基因':
                        MTgene.append(value8)
                    if key8=='药品':
                        MTdrug.append(value8)
                    if key8=='突变描述':
                        MTdescript.append(value8)
                    if key8=='氨基酸突变':
                        MTProtein.append(value8)
                    if key8=='突变率':
                        MTPrecent.append(value8)
            #print('MTgene,MTdrug,MTdescript,MTProtein,Precent',MTgene,MTdrug,MTdescript,MTProtein,MTPrecent)
            new_line_dict['结核耐药-突变基因'] = MTgene
            new_line_dict['结核耐药-潜在耐药药物'] = MTdrug
            new_line_dict['结核耐药-核酸突变结果'] = MTdescript
            new_line_dict['结核耐药-氨基酸突变'] = MTProtein
            new_line_dict['结核耐药-突变率'] = MTPrecent
            NMTgene=[];NMTdrug=[];NMTdescript=[];NMTProtein=[];NMTPrecent=[]
            for NMTlist in value['表9信息']:
                #print(MTlist)
                for key9,value9 in NMTlist.items():
                    # print(value['姓名'])
                    #print('key8,value8',key8,value8)
                    if key9=='基因':
                        NMTgene.append(value9)
                    if key9=='药品':
                        NMTdrug.append(value9)
                    if key9=='突变描述':
                        NMTdescript.append(value9)
                    if key9=='氨基酸突变':
                        NMTProtein.append(value9)
                    if key9=='突变率':
                        NMTPrecent.append(value9)
            new_line_dict['非结核耐药-突变基因'] = NMTgene
            new_line_dict['非结核耐药-潜在耐药药物'] = NMTdrug
            new_line_dict['非结核耐药-核酸突变结果'] = NMTdescript
            new_line_dict['非结核耐药-氨基酸突变'] = NMTProtein
            new_line_dict['非结核耐药-突变率'] = NMTPrecent
            new_df = (new_line_dict)
            new_line_info = str(new_line_dict['样本编号']) + '_' + str(new_line_dict['患者姓名']) + '_' + str(new_line_dict['basecalling data']) + '_' + new_line_dict['Reports']
            # print(data_info_dict)
            # print(all_info_dict)
            # print(new_line_info)
            if new_line_info not in data_info_dict:
                data_summary_df = data_summary_df.append(new_df, ignore_index=True)
            if new_line_info not in all_info_dict:
                all_summary_df = all_summary_df.append(new_df, ignore_index=True)
    data_summary_df.to_excel(data_summary_excel_path, index=False)
    all_summary_df.to_excel(summary_excel_path, index=False)
logging.info(f"汇总完成！")
