#!/usr/bin/env python

import argparse
import pandas as pd
import os
import docx
import logging
import re
from multiprocessing import Pool
from docx.shared import RGBColor
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docxtpl import DocxTemplate, RichText
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
import datetime

# 打印出运行的时间
time1 = '运行时间：' + str(datetime.datetime.now())
print(time1)

# 设定监控日志输出文件名和内容形式
logging.basicConfig(format='%(asctime)s - %(message)s', filename='/home/MicroReports/workspace/TNP-seq/OUTPUT/运行信息.txt', filemode='a', level=logging.INFO)

# 参数的导入与处理
parser = argparse.ArgumentParser()
parser.add_argument('-i', "--result_excel", required=True, help="the excel file with the result selected")
parser.add_argument('-b', "--database", type=str, default='/home/MicroReports/pipelines/TNP-seq/病原数据库/思可愈数据库-TNP-Seq病原菌测序项目2021.05.18.xlsx',help="database provided by the Ministry of Medicine")
parser.add_argument('-w', "--word_template_folder", type=str, default='/home/MicroReports/pipelines/TNP-seq/报告模板/word/',help="folder where all word report templates are located")
parser.add_argument('-e', "--excel_template_folder", type=str, default='/home/MicroReports/pipelines/TNP-seq/报告模板/excel/',help="folder where all excel report templates are located")
parser.add_argument('-n', "--processes_number", type=int, default=5,help="并行进程数目")
parser.add_argument('-c', "--complex_excel", type=str, default='/home/MicroReports/pipelines/TNP-seq/mycobacterium_tuberculosis_complex.xlsx',help="结核分支杆菌复合群包含微生物表格")
parser.add_argument('-s', "--summary_excel", type=str, default='/home/MicroReports/workspace/TNP-seq/OUTPUT/',help="summary documents before processing")
parser.add_argument('-o', "--output_dir", type=str, default='/home/MicroReports/workspace/TNP-seq/OUTPUT/',help="supplement sample result")
args = parser.parse_args()
info_client = pd.read_excel(args.result_excel).fillna('NA')
Interpretation = pd.read_excel(args.result_excel, sheet_name='species_report').fillna('NA')
sheet=pd.read_excel(args.result_excel,sheet_name=None)
if 'resistance_report' in list(sheet.keys()):
    drug_resistance_df = pd.read_excel(args.result_excel, sheet_name='resistance_report').fillna('NA')
    if drug_resistance_df.shape[0] == 0:
        drug_resistance_df = 0
else:
    drug_resistance_df = 0
# print(drug_resistance_df)
picture_dtat_df = pd.read_excel(args.result_excel, sheet_name='length_report').fillna('NA')
medical_DB = pd.read_excel(args.database).fillna('NA')
result_file_name = args.result_excel.split("/")[-1].lower()
complex_df = pd.read_excel(args.complex_excel)


####################################################
# 定义函数
# 判断是否为结核分枝杆菌复合群
def change_bacteria_list(bacteria_list: list):
    new_bacteria_list = []
    for bacteria in bacteria_list:
        compare_bacteria: str = Nor(bacteria)
        find_complex_df = complex_df[complex_df['name'] == compare_bacteria]
        if find_complex_df.shape[0] == 0:
            new_bacteria_list.append(bacteria)
        else:
            new_bacteria_list.append('Mycobacterium tuberculosis complex')
    return new_bacteria_list


# 标准化输入内容名称所用函数(变小写)
def Nor(x: str
    ) -> str:
    first: str= x.strip()
    standardized_string: str = " ".join(first.split())
    standardized_string: str = standardized_string.lower()
    return standardized_string


# 标准化输入列名称所用函数
def Nor_col(x: str
    ) -> str:
    first: str= str(x).strip()
    standardized_string: str = " ".join(first.split())
    return standardized_string


def make_picture(length_colname: str,
    picture_dtat_df: pd.DataFrame,
    ) -> None:
    R_out = length_colname + '.r'
    rscript=f'''#! /path/to/Rscript
library(openxlsx)
library(ggplot2)
data<- read.xlsx('{args.result_excel}', sheet='length_report')
names(data)[names(data) == '{length_colname}'] <- 'Frequency'
data <- data[c(1:8),]
row.names(data) <- as.list(data)$length
read_length_hist <- ggplot(data, mapping=aes(x=rownames(data), y=Frequency)) +
geom_bar(stat="identity", fill="#E1EFF9", colour="#E1EFF9") +
scale_x_discrete(limits=factor(rownames(data))) +
labs(x="length(bp)", y="ratio(%)") +
theme(panel.grid=element_blank(), panel.background=element_rect(color="black", fill="transparent")) + 
theme(axis.text =element_text(size=7))
ggsave(file="{length_colname}.png",read_length_hist, width = 6, height = 3)
    '''
    out = open(R_out,'w')
    out.write(rscript)
    out.close()
    cmd = "Rscript " + R_out
    os.system(cmd)
    os.remove(R_out)


def move_picture(doc,
    png_name: str
    ) -> None:
    table = doc.add_table(rows=1, cols=3)
    cell = table.cell(0, 1)
    ph =cell.paragraphs[0]
    run = ph.add_run()
    run.add_picture(png_name)
    target = None
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith('序列长度统计'):
            # print(paragraph_text)
            target = paragraph
            break
    move_table_after(table, target)
    os.remove(png_name)
    return doc


# 查询检测项目的简称
def project_shorthand(sample_code: str) -> str:
    project_name = info_client.loc[info_client['样本编号'] == sample_code, '检测项目'].iloc[0].rstrip()
    # print(project_name)
    if '呼吸' in project_name:
        hand: str = 'HX'
    elif '血液' in project_name:
        hand: str = 'XY'
    elif '神经' in project_name:
        hand: str = 'SJ'
    elif '胸腹' in project_name:
        hand: str = 'XF'
    elif '泌尿' in project_name:
        hand: str = 'MN'
    elif '消化' in project_name:
        hand: str = 'XH'
    elif '创口' in project_name:
        hand: str = 'CK'
    elif '眼科' in project_name:
        hand: str = 'YB'
    else:
        name = info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]
        logging.info(f'{name}的检测项目填写错误导致报告生成失败！')
    # print(hand)
    return hand


# 在数据库中搜索表格信息
def find_info(
    result_list: list,
    sample_code: str,
    Interpretation: pd.DataFrame,
    info_client: pd.DataFrame,
    medical_DB: pd.DataFrame,
    all_bac: list,
    mic_dict: dict,
    supplementary_results:int
) -> list:
    bac_list: list = []
    sample_result_list: list = []
    column_P = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_P_' + project_shorthand(sample_code) + '_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
    column_R = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_R_' + project_shorthand(sample_code) + '_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
    for bac_name in result_list:
        # print(bac_name)
        # print(column_P)
        compare_bac_name: str = Nor(bac_name)
        dic_bac: dict ={}
        try:
            dic_bac['中文名'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]
        except IndexError:
            dic_bac['中文名'] = 'NA'
        try:
            dic_bac['分类'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '分类'].iloc[0]
            try:
                dic_bac['分类顺序'] = mic_dict[dic_bac['分类']]
            except KeyError:
                dic_bac['分类顺序'] = 0   
        except IndexError:
            dic_bac['分类'] = 'NA'
            dic_bac['分类顺序'] = 0
        if dic_bac['分类'] not in ['DNA病毒','RNA病毒','病毒','真菌','细菌']:
            dic_bac['分类'] = '其他病原'
            try:
                dic_bac['分类顺序'] = mic_dict[dic_bac['分类']]
            except KeyError:
                dic_bac['分类顺序'] = 0   
        try:
            # print(Interpretation.columns.tolist())
            # print(Interpretation[column_P])
            # print(Interpretation.loc[Interpretation['name'] == compare_bac_name, column_P])
            dic_bac['相对丰度'] = '%.2f'%float(Interpretation.loc[Interpretation['name'] == compare_bac_name, column_P].iloc[0])
        except KeyError:
            column_P = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_P_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
            dic_bac['相对丰度'] = '%.2f'%float(Interpretation.loc[Interpretation['name'] == compare_bac_name, column_P].iloc[0])
        except IndexError:
            dic_bac['相对丰度'] = 'NA'
        try:
            dic_bac['序列数'] = int(Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0])
        except KeyError:
            column_R = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_R_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
            dic_bac['序列数'] = int(Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0])
        except IndexError:
            dic_bac['序列数'] = 'NA'
        report_name = bac_name.replace("[","")
        dic_bac['微生物'] = report_name.replace("]","")
        try:
            dic_bac['备注'] = medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '备注'].iloc[0]
        except IndexError:
            dic_bac['备注'] = 'NA'
        try:
            bac_list.append(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '分类'].iloc[0])
        except IndexError:
            pass
        sample_result_list.append(dic_bac)
    if supplementary_results == 1:
        return sample_result_list
    bac_list = list(set(bac_list))
    if 'DNA病毒' in bac_list:
        bac_list.append('DNA病毒')
    if '病毒' in bac_list or 'RNA病毒' in bac_list:
        bac_list.append('RNA病毒')
    for non_bac in all_bac:
        if non_bac not in bac_list:
            dic_bac = {}
            dic_bac['分类'] = non_bac
            dic_bac['分类顺序'] = mic_dict[non_bac]
            dic_bac['相对丰度'] = '--'
            dic_bac['序列数'] = '--'
            dic_bac['微生物'] = '--'
            dic_bac['备注'] = '--'
            dic_bac['中文名'] = '--'
            sample_result_list.append(dic_bac)
    sample_result_list: list = sorted(sample_result_list, key = lambda x:x['分类顺序'])
    return sample_result_list
    

# 分类正式报告结果
def microbial_classification(bacteria_list: list,
    medical_DB: pd.DataFrame) -> list:
    result1_list = []
    result2_list = []
    result3_list = []
    result4_list = []
    result5_list = []
    result6_list = []
    result7_list = []
    for microbial in bacteria_list:
        compare_bac_name: str = Nor(microbial)
        try:
            kingdom = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '分类'].iloc[0]
            genus = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '属'].iloc[0]
        except IndexError:
            kingdom = 'NA'
            genus = 'NA'
        if genus == '分枝杆菌属':
            result1_list.append(microbial)
        elif kingdom == '细菌':
            result2_list.append(microbial)
        elif kingdom == '真菌':
            result3_list.append(microbial)
        elif kingdom == 'DNA病毒':
            result4_list.append(microbial)
        elif kingdom == '病毒' or kingdom == 'RNA病毒':
            result5_list.append(microbial)
        elif kingdom == '古菌' or kingdom == '其他病原':
            result6_list.append(microbial)
        elif kingdom == '寄生虫':
            result7_list.append(microbial)
        else:
            logging.info(f"{microbial}没有在数据库中找到，导致报告中未显示！")
    return [result1_list,result2_list,result3_list,result4_list,result5_list,result6_list,result7_list]


# 表格信息的生成1
def table_context(sample_code: str,
    info_client: pd.DataFrame,
    Interpretation: pd.DataFrame,
    medical_DB: pd.DataFrame) -> list:
    # print(sample_code)
    if 'RNA' in info_client.loc[info_client['样本编号'] == sample_code, '检测项目'].iloc[0] and 'DNA' not in info_client.loc[info_client['样本编号'] == sample_code, '检测项目'].iloc[0]:
        all_bac: list = ['RNA病毒']
        mic_dict: dict = {'RNA病毒':1}
    elif 'RNA' in info_client.loc[info_client['样本编号'] == sample_code, '检测项目'].iloc[0]:
        all_bac: list = ['细菌', '真菌', 'DNA病毒', 'RNA病毒','其他病原']
        mic_dict: dict = {'细菌':1, '真菌':2, 'DNA病毒':3, 'RNA病毒':4, '其他病原':5}
    else:
        all_bac: list = ['细菌', '真菌', 'DNA病毒', '其他病原']
        mic_dict: dict = {'细菌':1, '真菌':2, 'DNA病毒':3, '其他病原':4}
    bacteria_list: list = []
    if info_client.loc[info_client['样本编号'] == sample_code, '正式报告结果'].iloc[0] != 'NA':
        bacteria_list: list = info_client.loc[info_client['样本编号'] == sample_code, '正式报告结果'].iloc[0].split(',')
    # print(bacteria_list)
    # bacteria_list = change_bacteria_list(bacteria_list)
    supplementary_results = 0
    table1_list = []
    table1_list = find_info(result_list=bacteria_list, sample_code=sample_code, Interpretation=Interpretation, info_client=info_client, medical_DB=medical_DB, all_bac=all_bac, mic_dict=mic_dict, supplementary_results=supplementary_results)
    table2_list = []
    if info_client.loc[info_client['样本编号'] == sample_code, '补充报告结果'].iloc[0] != 'NA':
        supplementary_results = 1
        bacteria_list: list = info_client.loc[info_client['样本编号'] == sample_code, '补充报告结果'].iloc[0].split(',')
        # bacteria_list = change_bacteria_list(bacteria_list)
        table2_list: list = find_info(result_list=bacteria_list, sample_code=sample_code, Interpretation=Interpretation, info_client=info_client, medical_DB=medical_DB, all_bac=all_bac, mic_dict=mic_dict, supplementary_results=supplementary_results)
    else:
        table2_list = [{'中文名': '--', '分类': '--', '分类顺序': 1, '相对丰度': '--', '序列数': '--', '微生物': '--', '备注': '--'}]
    return [table1_list,table2_list]


# 表格信息的生成2
def table2_context(sample_code: str,
    info_client: pd.DataFrame,
    Interpretation: pd.DataFrame,
    medical_DB: pd.DataFrame) -> list:
    # print(sample_code)
    all_bac_list: list = [['细菌'],['细菌'],['真菌'],['DNA病毒'],['RNA病毒'],['其他病原'],['寄生虫']]
    mic_dict_list: dict = [{'细菌':1}, {'细菌':1}, {'真菌':1}, {'DNA病毒':1,}, {'RNA病毒':1},{'其他病原':1},{'寄生虫':1}]
    bacteria_list: list = []
    if info_client.loc[info_client['样本编号'] == sample_code, '正式报告结果'].iloc[0] != 'NA':
        bacteria_list: list = info_client.loc[info_client['样本编号'] == sample_code, '正式报告结果'].iloc[0].split(',')
    # print(bacteria_list)
    # bacteria_list = change_bacteria_list(bacteria_list)
    all_kinds_list = microbial_classification(bacteria_list=bacteria_list,medical_DB=medical_DB)
    supplementary_results = 0
    table1_list = []
    for index, result_list in enumerate(all_kinds_list):
        # print(result_list)
        table1_list.append(find_info(result_list=result_list, sample_code=sample_code, Interpretation=Interpretation, info_client=info_client, medical_DB=medical_DB, all_bac=all_bac_list[index], mic_dict=mic_dict_list[index], supplementary_results=supplementary_results))
        # print(table1_list)
    table2_list = []
    if info_client.loc[info_client['样本编号'] == sample_code, '补充报告结果'].iloc[0] != 'NA':
        supplementary_results += 1
        bacteria_list: list = info_client.loc[info_client['样本编号'] == sample_code, '补充报告结果'].iloc[0].split(',')
        # bacteria_list = change_bacteria_list(bacteria_list)
        table2_list: list = find_info(result_list=bacteria_list, sample_code=sample_code, Interpretation=Interpretation, info_client=info_client, medical_DB=medical_DB, all_bac=all_bac_list[0], mic_dict=mic_dict_list[0], supplementary_results=supplementary_results)
    else:
        table2_list = [{'中文名': '--', '分类': '--', '分类顺序': 1, '相对丰度': '--', '序列数': '--', '微生物': '--', '备注': '--'}]
    return [table1_list,table2_list]


# 表格信息的生成2
def table7_make(sample_code: str,
    info_client: pd.DataFrame,
    drug_resistance_df: pd.DataFrame
    ) -> list:
    handle_df = info_client[info_client['样本编号'] == sample_code]
    pat_name = handle_df['患者姓名'].iloc[0]
    drug_resistance_info = info_client[(info_client['患者姓名'] == pat_name) & (info_client['备注'].str.contains('普通耐药'))]
    # print(drug_resistance_info)
    table_list = []
    if drug_resistance_info.shape[0] != 0:
        compare_sample_code = drug_resistance_info['样本编号'].iloc[0]
        # print(compare_sample_code)
        drug_resistance_colname = drug_resistance_info.loc[drug_resistance_info['样本编号'] == compare_sample_code, 'barcode'].iloc[0] + '_' + project_shorthand(compare_sample_code) + '_' + str(drug_resistance_info.loc[drug_resistance_info['样本编号'] == compare_sample_code, '患者姓名'].iloc[0]) + '_' + compare_sample_code + '_num'
        # print(drug_resistance_colname)
        # print(drug_resistance_df)
        try:
            drug_resistance_result = drug_resistance_df[drug_resistance_colname]
        except KeyError:
            drug_resistance_colname = drug_resistance_info.loc[drug_resistance_info['样本编号'] == compare_sample_code, 'barcode'].iloc[0] + '_' + str(drug_resistance_info.loc[drug_resistance_info['样本编号'] == compare_sample_code, '患者姓名'].iloc[0]) + '_' + compare_sample_code + '_num'
            drug_resistance_result = drug_resistance_df[drug_resistance_colname]
        gene_list = drug_resistance_df['gene'].tolist()
        for gene in gene_list:
            number = drug_resistance_df.query('gene == @gene').iloc[0,:][drug_resistance_colname]
            if int(number) != 0:
                dic_bac ={}
                dic_bac['基因'] = gene
                dic_bac['药物'] = drug_resistance_df.query('gene == @gene').iloc[0,:]['drug']
                table_list.append(dic_bac)
    if len(table_list) == 0:
        dic_bac ={}
        dic_bac['基因'] = '--'
        dic_bac['药物'] = '--'
        table_list.append(dic_bac)
    return table_list


# 临床解析的生成
def clinical(sample_code: str,
    info_client: pd.DataFrame,
    medical_DB: pd.DataFrame,
    manufacturer: str):
    rt = RichText('')
    if  manufacturer == 'beagle':
        bacteria_list: list =  info_client.loc[info_client['样本编号'] == sample_code, '正式报告结果'].iloc[0].split(',')
    elif manufacturer == 'seegene' or 'boruilin' or 'beijing':
        if info_client.loc[info_client['样本编号'] == sample_code, '补充报告结果'].iloc[0] != 'NA':
            bacteria_list: list =  info_client.loc[info_client['样本编号'] == sample_code, '补充报告结果'].iloc[0].split(',')
        else:
            return rt
    # bacteria_list = change_bacteria_list(bacteria_list)
    dic_number: dict = {1:'一', 2:'二', 3:'三', 4:'四', 5:'五', 6:'六', 7:'七', 8:'八', 9:'九', 10:'十', 11:'十一', 12:'十二', 13:'十三', 14:'十四', 15:'十五', 16:'十六', 17:'十七', 18:'十八', 19:'十九', 20:'二十', 21:'二十一', 22:'二十二', 23:'二十三', 24:'二十四', 25:'二十五', 26:'二十六', 27:'二十七', 28:'二十八', 29:'二十九', 30:'三十'}
    drug_indication_value: int = 0
    for index, bac_name in enumerate(bacteria_list):
        compare_bac_name: str = Nor(bac_name)
        try:
            new_name = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]
            if str(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]) != 'NA':
                chinese_name: str = str(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0])
            else:
                chinese_name = ''
            rt.add(dic_number[index+1]+' '+chinese_name, bold=True)
            rt.add('(', bold=True)
            report_name: str = bac_name.replace("[","")
            report_name: str = report_name.replace("]","")
            rt.add(report_name, italic=True, bold=True)
            rt.add(')\n', bold=True)
            if  manufacturer == 'beagle':
                try:
                    if str(medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '常用药物'].iloc[0]) != 'NA':
                        rt.add('    '+'1 临床意义', bold=True)
                        rt.add('\n'+'    '+medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '临床意义'].iloc[0] )
                        rt.add('\n'+'    '+'2 常用药物', bold=True)
                        rt.add('\n'+'    '+str(medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '常用药物'].iloc[0])+'\n'  )
                        drug_indication_value += 1
                    elif str(medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '临床意义'].iloc[0]) != 'NA':
                        rt.add('    '+'临床意义', bold=True)
                        rt.add('\n'+'    '+medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '临床意义'].iloc[0]+'\n' )
                except IndexError:
                    logging.info(f"正式结果中的{bac_name}在数据库中的检测项目信息有问题，影响临床解析的生成")
            elif manufacturer == 'seegene' or 'boruilin' or 'beijing':
                try:
                    rt.add('    '+medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '临床意义'].iloc[0]+'\n')
                except IndexError:
                    logging.info(f"正式结果中的{bac_name}在数据库中的检测项目信息有问题，影响临床解析的生成")
        except IndexError:
            logging.info(f"正式结果中的{bac_name}在数据库中没有找到，影响临床解析的生成")
    if drug_indication_value != 0 and manufacturer == 'beagle':
        rt.add('\n注：常用药物为临床常规药物，且无法覆盖药敏结果，具体用药请结合临床药敏结果或医院耐药监测数据酌情用药。', bold=True)
    return rt


# 参考文献的生成
def reference(sample_code: str):
    bacteria_list: list =  info_client.loc[info_client['样本编号'] == sample_code, '正式报告结果'].iloc[0].split(',')
    dic_number: dict = {1:'一', 2:'二', 3:'三', 4:'四', 5:'五', 6:'六', 7:'七', 8:'八', 9:'九'}
    index_str: int = len(bacteria_list) + 1
    rt = RichText(dic_number[index_str]+'、'+'参考文献\n')
    all_reference_list: list = []
    reference_list: list = []
    for bac_name in bacteria_list:
        compare_bac_name: str = Nor(bac_name)
        try:
            new_reference: str = medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '参考文献'].iloc[0]
            for literature in new_reference.split("\n"):
                all_reference_list.append(literature[2:])
            reference_list: list = list(set(all_reference_list))
        except IndexError:
            pass
    for index, val in enumerate(reference_list):
        rt.add(str(index+1)+'.'+str(val)+'\n')
    return rt


# seegene报告修改表格
def form_modification(doc,
    dic_client: dict,
    sample_code: str):
    table = doc.tables[3]
    if len(doc.tables[3].rows) != 2:
        add_dict = {'number_1':[],'number_2':[],'number_3':[],'number_4':[],'number_5':[],'number_6':[],'number_7':[]}
        for line_info in dic_client[sample_code]['表9信息']:
            if line_info['分类'] == '细菌':
                add_dict['number_2'].append(line_info)
            elif line_info['分类'] == '真菌':
                add_dict['number_3'].append(line_info)
            elif line_info['分类'] == 'DNA病毒':
                add_dict['number_4'].append(line_info)
            elif line_info['分类'] == '病毒' or line_info['分类'] == 'RNA病毒':
                add_dict['number_5'].append(line_info)
            elif line_info['分类'] == '古菌' or line_info['分类'] == '其他病原':
                add_dict['number_6'].append(line_info)
            elif line_info['分类'] == '寄生虫':
                add_dict['number_7'].append(line_info)
        number = 1
        while number < 8:
            table_info_name = f'表{number}信息'
            table_info = dic_client[sample_code][table_info_name]
            change_color_table = doc.tables[3+number]
            if table_info[0]['微生物'] != '--':
                for i,line in enumerate(table_info):
                    row = i + 2
                    col = len(change_color_table.columns)
                    for col_number in range(col):
                        run = change_color_table.cell(row,col_number).paragraphs[0]
                        content = run.text
                        run.text = ''
                        run = run.add_run(content)
                        run.font.color.rgb = RGBColor(255,0,0)
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        if col_number == 1:
                            run.italic = True              
            else:
                run = change_color_table.cell(2,1).paragraphs[0]
                run.text = ''
                run = run.add_run('--')
                run.font.color.rgb = RGBColor(0,0,0)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            add_key = f'number_{number}'
            table_info.extend(add_dict[add_key])
            result_list = []
            if len(table_info) != 1 or table_info[0]['微生物'] != '--':
                add_number = 0
                for input_dict in table_info:
                    if input_dict['备注'] == '人体共生菌':
                        add_number += 1
                    elif input_dict['中文名'] != 'NA' and input_dict['中文名'] != '--':
                        if input_dict in add_dict[add_key]:
                            result_list.append(input_dict['中文名'] + '（补充报告部分）')
                        else:
                            result_list.append(input_dict['中文名'])
                    elif input_dict['中文名'] == 'NA' and input_dict['微生物'] != '--':
                        if input_dict in add_dict[add_key]:
                            result_list.append(input_dict['微生物'] + '（补充报告部分）')
                        else:
                            result_list.append(input_dict['微生物'])
                if len(result_list) != 0:
                        result_info = "，".join(result_list)
                        run = table.cell(number,1).paragraphs[0]
                        run.text = ''
                        run = run.add_run(result_info)
                        run.font.color.rgb = RGBColor(255,0,0)
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                if add_number != 0:
                    p = table.cell(number,1).paragraphs[0]
                    if p.text != '未检出疑似病原体':
                        run = p.add_run('，疑似微生态菌群')
                    else:
                        p.text = ''
                        run = p.add_run('疑似微生态菌群')
                    run.font.color.rgb = RGBColor(255,0,0)
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')    
            number += 1
        table_info =  dic_client[sample_code]['表8信息']
        if table_info[0]['基因'] != '--':
            for infp_dict in table_info:
                result_list.append(infp_dict['基因'])
            result_info = ",".join(result_list)
            run = table.cell(number,1).paragraphs[0]
            run.text = ''
            run = run.add_run(result_info)
            run.font.color.rgb = RGBColor(255,0,0)
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    else:
        table_info =  dic_client[sample_code]['表4信息']
        change_color_table = doc.tables[4]
        if table_info[0]['微生物'] != '--':
            for i,line in enumerate(table_info):
                row = i + 2
                col = len(change_color_table.columns)
                for col_number in range(col):
                    run = change_color_table.cell(row,col_number).paragraphs[0]
                    content = run.text
                    run.text = ''
                    run = run.add_run(content)
                    run.font.color.rgb = RGBColor(255,0,0)
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    if col_number == 1:
                        run.italic = True     
        else:
            run = change_color_table.cell(2,1).paragraphs[0]
            run.text = ''
            run = run.add_run('--')
            run.font.color.rgb = RGBColor(0,0,0)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
        result_list = []
        if len(table_info) != 1 or table_info[0]['微生物'] != '--':
            add_number = 0
            for input_dict in table_info:
                if input_dict['中文名'] != 'NA' and input_dict['中文名'] != '--':
                    result_list.append(input_dict['中文名'])
                elif input_dict['中文名'] == 'NA' and input_dict['微生物'] != '--':
                    result_list.append(input_dict['微生物'])
                # print(result_list)
            if len(result_list) != 0:
                result_info = "，".join(result_list)
                run = table.cell(number,1).paragraphs[0]
                run.text = ''
                run = run.add_run(result_info)
                # print(run.text)
                run.font.color.rgb = RGBColor(255,0,0)
                run.font.size = Pt(11)
                run.font.name = 'Arial'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
def color_change3(doc,
    ) -> None:
    number = 8
    while number < 9:
        table_info_name = f'表{number}信息'
        print(table_info_name)
        table_info = dic_client[sample_code][table_info_name]
        print(table_info)
        change_color_table = doc.tables[2+number]
        if table_info[0]['基因'] != '--':
            for i,line in enumerate(table_info):
                row = i+2
                col = len(change_color_table.columns)
                #print(change_color_table(row,col).text)
                for col_number in range(col):
                    run = change_color_table.cell(row,col_number).paragraphs[0]
                    content = run.text
                    run.text = ''
                    run = run.add_run(content)
                    run.font.color.rgb = RGBColor(255,0,0)
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    if col_number == 0:
                        run.italic = True
        else:
            run = change_color_table.cell(2,1).paragraphs[0]
            run.text = ''
            run = run.add_run('--')
            run.font.color.rgb = RGBColor(0,0,0)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')						
        number += 1
def form_modification2(doc,
    dic_client: dict,
    sample_code: str):
    table = doc.tables[3]
    if len(doc.tables[3].rows) != 2:
        add_dict = {'number_1':[],'number_2':[],'number_3':[],'number_4':[],'number_5':[],'number_6':[]}
        for line_info in dic_client[sample_code]['表9信息']:
            if line_info['分类'] == '细菌':
                add_dict['number_2'].append(line_info)
            elif line_info['分类'] == '真菌':
                add_dict['number_3'].append(line_info)
            elif line_info['分类'] == 'DNA病毒':
                add_dict['number_4'].append(line_info)
            # elif line_info['分类'] == '病毒' or line_info['分类'] == 'RNA病毒':
            #     add_dict['number_5'].append(line_info)
            elif line_info['分类'] == '古菌' or line_info['分类'] == '其他病原':
                add_dict['number_5'].append(line_info)
            elif line_info['分类'] == '寄生虫':
                add_dict['number_6'].append(line_info)
        #print('add_dict',add_dict)
        number = 1
        while number < 7:
            if number <5:
                table_info_name = f'表{number}信息'
                #print('table_info_name',table_info_name)
                table_info = dic_client[sample_code][table_info_name]
                change_color_table = doc.tables[3+number]
                #print('doc.tables[3+number]',3+number)
                #print('table_info[0][微生物]',table_info)
                if table_info[0]['微生物'] != '--':
                    for i,line in enumerate(table_info):
                        #print(i,line)
                        row = i + 2
                        col = len(change_color_table.columns)
                        for col_number in range(col):
                            run = change_color_table.cell(row,col_number).paragraphs[0]
                            content = run.text
                            # print(content)
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            if col_number == 1:
                                run.italic = True
                else:
                    run = change_color_table.cell(2,1).paragraphs[0]
                    run.text = ''
                    run = run.add_run('--')
                    run.font.color.rgb = RGBColor(0,0,0)
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                add_key = f'number_{number}'
                table_info.extend(add_dict[add_key])
                result_list = []
                if len(table_info) != 1 or table_info[0]['微生物'] != '--':
                    add_number = 0
                    for input_dict in table_info:
                        if input_dict['备注'] == '人体共生菌':
                            add_number += 1
                        elif input_dict['中文名'] != 'NA' and input_dict['中文名'] != '--':
                            if input_dict in add_dict[add_key]:
                                result_list.append(input_dict['中文名'] + '（补充报告部分）')
                            else:
                                result_list.append(input_dict['中文名'])
                        elif input_dict['中文名'] == 'NA' and input_dict['微生物'] != '--':
                            if input_dict in add_dict[add_key]:
                                result_list.append(input_dict['微生物'] + '（补充报告部分）')
                            else:
                                result_list.append(input_dict['微生物'])
                    if len(result_list) != 0:
                            result_info = "，".join(result_list)
                            run = table.cell(number,1).paragraphs[0]
                            run.text = ''
                            run = run.add_run(result_info)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(11)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    if add_number != 0:
                        p = table.cell(number,1).paragraphs[0]
                        if p.text != '未检出疑似病原体':
                            run = p.add_run('，疑似微生态菌群')
                        else:
                            p.text = ''
                            run = p.add_run('疑似微生态菌群')
                        run.font.color.rgb = RGBColor(255,0,0)
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')    
                number += 1
            elif number == 5:
                pass
                number += 1
            else:
                table_info_name = f'表{number}信息'
                # print('table_info_name',table_info_name)
                table_info = dic_client[sample_code][table_info_name]
                change_color_table = doc.tables[2+number]
                # print('doc.tables[3+number]',3+number)
                if table_info[0]['微生物'] != '--':
                    for i,line in enumerate(table_info):
                        row = i + 2
                        col = len(change_color_table.columns)
                        for col_number in range(col):
                            run = change_color_table.cell(row,col_number).paragraphs[0]
                            content = run.text
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            if col_number == 1:
                                run.italic = True              
                else:
                    run = change_color_table.cell(2,1).paragraphs[0]
                    run.text = ''
                    run = run.add_run('--')
                    run.font.color.rgb = RGBColor(0,0,0)
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                add_key = f'number_{number-1}'
                table_info.extend(add_dict[add_key])
                result_list = []
                if len(table_info) != 1 or table_info[0]['微生物'] != '--':
                    add_number = 0
                    for input_dict in table_info:
                        if input_dict['备注'] == '人体共生菌':
                            add_number += 1
                        elif input_dict['中文名'] != 'NA' and input_dict['中文名'] != '--':
                            if input_dict in add_dict[add_key]:
                                result_list.append(input_dict['中文名'] + '（补充报告部分）')
                            else:
                                result_list.append(input_dict['中文名'])
                        elif input_dict['中文名'] == 'NA' and input_dict['微生物'] != '--':
                            if input_dict in add_dict[add_key]:
                                result_list.append(input_dict['微生物'] + '（补充报告部分）')
                            else:
                                result_list.append(input_dict['微生物'])
                    if len(result_list) != 0:
                            result_info = "，".join(result_list)
                            run = table.cell(number-1,1).paragraphs[0]
                            run.text = ''
                            run = run.add_run(result_info)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(11)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    if add_number != 0:
                        p = table.cell(number-1,1).paragraphs[0]
                        if p.text != '未检出疑似病原体':
                            run = p.add_run('，疑似微生态菌群')
                        else:
                            p.text = ''
                            run = p.add_run('疑似微生态菌群')
                        run.font.color.rgb = RGBColor(255,0,0)
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')    
                number += 1             
        table_info =  dic_client[sample_code]['表8信息']
        if table_info[0]['基因'] != '--':
            for infp_dict in table_info:
                result_list.append(infp_dict['基因'])
            result_info = ",".join(result_list)
            run = table.cell(number,1).paragraphs[0]
            run.text = ''
            run = run.add_run(result_info)
            run.font.color.rgb = RGBColor(255,0,0)
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    else:
        table_info =  dic_client[sample_code]['表4信息']
        change_color_table = doc.tables[4]
        if table_info[0]['微生物'] != '--':
            for i,line in enumerate(table_info):
                row = i + 2
                col = len(change_color_table.columns)
                for col_number in range(col):
                    run = change_color_table.cell(row,col_number).paragraphs[0]
                    content = run.text
                    run.text = ''
                    run = run.add_run(content)
                    run.font.color.rgb = RGBColor(255,0,0)
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    if col_number == 1:
                        run.italic = True     
        else:
            run = change_color_table.cell(2,1).paragraphs[0]
            run.text = ''
            run = run.add_run('--')
            run.font.color.rgb = RGBColor(0,0,0)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
        result_list = []
        if len(table_info) != 1 or table_info[0]['微生物'] != '--':
            add_number = 0
            for input_dict in table_info:
                if input_dict['中文名'] != 'NA' and input_dict['中文名'] != '--':
                    result_list.append(input_dict['中文名'])
                elif input_dict['中文名'] == 'NA' and input_dict['微生物'] != '--':
                    result_list.append(input_dict['微生物'])
                # print(result_list)
            if len(result_list) != 0:
                result_info = "，".join(result_list)
                run = table.cell(number,1).paragraphs[0]
                run.text = ''
                run = run.add_run(result_info)
                # print(run.text)
                run.font.color.rgb = RGBColor(255,0,0)
                run.font.size = Pt(11)
                run.font.name = 'Arial'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
# 检出微生物添加到附录中
def add_micro(df: pd.DataFrame,
    chinese_name: str,
    genus_name: str,
    micro_type: str,
    pathogenicity_info: str,
    col_numbers: int,
    classification: str
) -> pd.DataFrame:
    # print('df',df)
    appendix_list = df.iloc[:, 1].apply(Nor).tolist()
    # print(appendix_list)
    # print(Nor(chinese_name))
    # print(Nor(genus_name))
    if (Nor(chinese_name) in appendix_list) and (Nor(genus_name) in appendix_list):
        micro_row = df[df.iloc[:, 1].apply(Nor)==chinese_name].index.tolist()[0]
        genus_row = df[df.iloc[:, 1].apply(Nor)==genus_name].index.tolist()[0]
        df.loc[micro_row, '结果'] = '检出'
        df.loc[genus_row, '结果'] = '检出'
    elif (Nor(genus_name) in appendix_list):
        genus_row = df[df.iloc[:, 1].apply(Nor)==genus_name].index.tolist()[0]
        df1 = df.loc[:genus_row]
        df2 = df.loc[(genus_row+1):]          
        if col_numbers == 3:
            df3 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:['  '+chinese_name],df.columns.tolist()[2]:['检出']})
        else:
            df3 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:['  '+chinese_name],df.columns.tolist()[2]:[pathogenicity_info],df.columns.tolist()[3]:['检出']}) 
        df = df1.append(df3, ignore_index = True).append(df2, ignore_index = True)
        df.loc[genus_row, '结果'] = '检出'
    elif '病毒' not in classification:
        if col_numbers == 3:
            df4 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:[genus_name],df.columns.tolist()[2]:['检出']})
            df5 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:['  '+chinese_name],df.columns.tolist()[2]:['检出']})            
        else:
            df4 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:[genus_name],df.columns.tolist()[2]:[pathogenicity_info],df.columns.tolist()[3]:['检出']})
            df5 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:['  '+chinese_name],df.columns.tolist()[2]:[pathogenicity_info],df.columns.tolist()[3]:['检出']})
        df = df.append(df4, ignore_index = True).append(df5, ignore_index = True)
    if '病毒' in classification:
        if (Nor(chinese_name) in appendix_list):
            micro_row = df[df.iloc[:, 1].apply(Nor)==chinese_name].index.tolist()[0]
            df.loc[micro_row, '结果'] = '检出'
        else:
            if col_numbers == 3:
                df6 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:[chinese_name],df.columns.tolist()[2]:['检出']}) 
            else:
                df6 = pd.DataFrame({df.columns.tolist()[0]:[micro_type],df.columns.tolist()[1]:[chinese_name],df.columns.tolist()[2]:[pathogenicity_info],df.columns.tolist()[3]:['检出']})
            df = df.append(df6, ignore_index = True)
    return df


# 添加表格框线
def Set_cell_border(cell: _Cell, **kwargs):
    """
    设置单元格边框函数
    使用方法:
    Set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    传入参数有cell, 即单元格；top指上边框；bottom指下边框；start指左边框；end指右边框。
    "sz"指线的粗细程度；"val"指线型，比如单线，虚线等；"color"指颜色，颜色编码可百度；
    "space"指间隔，一般不设置，设置的值大于0会导致线错开；"shadow"指边框阴影
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


# 将表格添加到特定的字后面
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


# 定义表格的文字特征
def change_format(table,
    row:int,
    col:int
    ):
    run = table.cell(row,col).paragraphs[0]
    content = run.text
    run.text = ''
    run = run.add_run(content)
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')    
    if run.text == '检出':
        run.font.color.rgb = RGBColor(255,0,0)
        run = table.cell(row,col-1).paragraphs[0]
        content = run.text
        run.text = ''
        run = run.add_run(content)
        run.font.color.rgb = RGBColor(255,0,0)
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  
        run = table.cell(row,col-2).paragraphs[0]
        content = run.text
        run.text = ''
        run = run.add_run(content)
        run.font.color.rgb = RGBColor(255,0,0)
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')         
        if len(table.columns) != 6:
            run = table.cell(row,col-3).paragraphs[0]
            content = run.text
            run.text = ''
            run = run.add_run(content)
            run.font.color.rgb = RGBColor(255,0,0)
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑') 



# 将excel中的数据框转化为docx中的表格
def change_type(df: pd.DataFrame,
    table,
    col_width_dic: dict
    ):
    for col in list(range(len(table.columns))):
        for row in list(range(len(table.rows))):
            # print(col)
            # print(row)
            table.cell(row, col).width = col_width_dic[col]
            if row == 0:
                try:
                    table.cell(row,col).text = list(df.columns)[col]
                    # print(table.cell(row,col).text)
                    change_format(table,row,col)
                    run = table.cell(row,col).paragraphs[0]
                    run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except IndexError:
                    cell_col_index = col - int(len(table.columns)/2)
                    table.cell(row,col).text = list(df.columns)[cell_col_index]
                    change_format(table,row,col)
                    run = table.cell(row,col).paragraphs[0]
                    run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # print(table.cell(row,col).text)
            else:
                cell_col_index = col
                cell_row_index = row - 1
                try:
                    table.cell(row,col).text = df.iloc[cell_row_index, cell_col_index]
                    run = table.cell(row,col).paragraphs[0]
                    # print(table.cell(row,col).text)
                    change_format(table,row,col)
                    if cell_col_index == 1 and run.text != '--' :
                        run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except IndexError:
                    cell_col_index -= int(len(table.columns)/2)
                    cell_row_index = row + len(table.rows) - 2
                    table.cell(row,col).text = df.iloc[cell_row_index, cell_col_index]
                    # print(table.cell(row,col).text)
                    change_format(table,row,col)
                    run = table.cell(row,col).paragraphs[0]
                    if cell_col_index == 1 and run.text != '--' :
                        run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            Set_cell_border(table.cell(row,col),
            top={"sz": 12, "val": "single", "color": "#F2F2F2"},
            bottom={"sz": 12, "val": "single", "color": "#F2F2F2"},
            start={"sz": 12, "val": "single", "color": "#F2F2F2"},
            end={"sz": 12, "val": "single", "color": "#F2F2F2"})
    return table


# word中改变首行背景
def tabBgColor(table,cols,colorStr):
    shading_list = locals()
    for i in range(cols):
        shading_list['shading_elm_'+str(i)] = parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'),bgColor = colorStr))
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_'+str(i)])


# 解读数量确认并添加
def Interpretation_addition(sample_code: str,
    dic_client: dict,
    doc,
    insert_info: str):
    dic_number: dict = {1:'一', 2:'二', 3:'三', 4:'四', 5:'五', 6:'六', 7:'七', 8:'八', 9:'九'}
    bac_list: list = dic_client[sample_code]['检测微生物'].split(',')
    Interpretation_list = []
    for i,bac in enumerate(bac_list):
        Interpretation_list.append(f'解读{dic_number[i + 1]}')
    for Interpretation in Interpretation_list:
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            if paragraph_text.strip() == insert_info.strip():
                prior_paragraph = paragraph.insert_paragraph_before(Interpretation)
    return Interpretation_list


# seegene模板添加解析表格
def parse_table_add(sample_code: str,
    dic_client: dict,
    doc,
    medical_DB: pd.DataFrame,
    Interpretation_list: list):
    dic_number: dict = {1:'一', 2:'二', 3:'三', 4:'四', 5:'五', 6:'六', 7:'七', 8:'八', 9:'九'}
    bac_list: list = dic_client[sample_code]['检测微生物'].split(',')
    # bac_list = change_bacteria_list(bac_list)
    for bac_name in reversed(bac_list):
        #print('bac_name',bac_name)
        if len(bac_name) != 0:
            number = dic_number[bac_list.index(bac_name) + 1]
            #print('number',number)
            compare_bac_name = Nor(bac_name)
            table = doc.add_table(rows=3, cols=2)
            tabBgColor(table, 2, '#E1EFF9')
            Interpretation = Interpretation_list[bac_list.index(bac_name)]
            # print('Interpretation',Interpretation)
            for col in list(range(len(table.columns))):
                for row in list(range(len(table.rows))):
                    if row == 0:
                        if col == 0: 
                            table.cell(row, col).width = Cm(2.8)
                            try:
                                if medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0] != 'NA':
                                    info = number + '、' + medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0].strip() + '(' + bac_name + ')'
                                else:
                                    info = number + '、' + bac_name
                            except IndexError:
                                info = number + '、' + bac_name
                            run = table.cell(row,col).paragraphs[0]
                            run = run.add_run(info)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        Set_cell_border(table.cell(row, col), bottom={"sz": 24, "val": "single", "color": "#1F3864"})
                        table.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    elif row == 1:
                        if col == 0: 
                            table.cell(row, col).width = Cm(2.8)
                            info = '临床意义'
                            run = table.cell(row,col).paragraphs[0]
                            run = run.add_run(info)
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            run.font.bold = True
                        elif col == 1: 
                            table.cell(row, col).width = Cm(14.7)
                            try:
                                info = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '临床意义'].iloc[0]
                            except IndexError:
                                info = 'NA'
                            table.cell(row,col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            run = table.cell(row,col).paragraphs[0]
                            run = run.add_run(info)
                            run.font.size = Pt(10.5)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        Set_cell_border(table.cell(row, col), bottom={"sz": 4, "val": "single", "color": "#1F3864"})
                        table.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    elif row == 2:
                        if col == 0: 
                            table.cell(row, col).width = Cm(2.8)
                            info = '常用药物'
                            run = table.cell(row,col).paragraphs[0]
                            run = run.add_run(info)
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            run.font.bold = True
                        elif col == 1: 
                            table.cell(row, col).width = Cm(14.7)
                            try:
                                info = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '常用药物'].iloc[0]
                            except IndexError:
                                info = 'NA'
                            table.cell(row,col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            run = table.cell(row,col).paragraphs[0]
                            run = run.add_run(info)
                            run.font.size = Pt(10.5)
                            run.font.name = 'Times New Roman'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        Set_cell_border(table.cell(row, col), bottom={"sz": 4, "val": "single", "color": "#1F3864"})
                        table.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.rows[0].height = Cm(1)
            table.rows[1].height = Cm(3.8)
            table.rows[2].height = Cm(6)
            table.cell(0,0).merge(table.cell(0,1))
            target = None
            for paragraph in doc.paragraphs:
                #print(paragraph.text)
                paragraph_text = paragraph.text
                if paragraph_text.endswith(Interpretation):
                    target = paragraph
                    break
            #print(table, target)
            move_table_after(table, target)
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            if paragraph_text.endswith(Interpretation):
                for run in paragraph.runs:
                    run.text=run.text.replace(Interpretation,'')
                new_paragraph = target



def make_word_report(sample_code: str,
    dic_client: dict,
    report_dic: dict,
    doc_dic: dict,
    medical_DB: pd.DataFrame
    ) -> None:
    # print(sample_code)
    # print(str(dic_client[sample_code]['姓名']))
    project_shortname = project_shorthand(sample_code)
    suffix = dic_client[sample_code]['检测项目'].split('测')[-1]
    # print(suffix)
    if dic_client[sample_code]['代理商'] == 'seegene':
        if 'DNA' in suffix and 'RNA' in suffix:
            df_name = report_dic[project_shortname][1] + '.xlsx'
            open_name = doc_dic[project_shortname][1] + '.docx'
        elif 'RNA' in suffix:
            df_name = report_dic[project_shortname][2] + '.xlsx'
            open_name = doc_dic[project_shortname][2] + '.docx'
        else:
            df_name = report_dic[project_shortname][0] + '.xlsx'
            open_name = doc_dic[project_shortname][0] + '.docx'
    elif dic_client[sample_code]['代理商'] == 'beagle':
        if 'DNA' in suffix and 'RNA' in suffix:
            df_name = report_dic[project_shortname][1] + '.xlsx'
            open_name = doc_dic[project_shortname][-1] + '.docx'
        elif 'RNA' in suffix:
            df_name = report_dic[project_shortname][2] + '.xlsx'
            open_name = doc_dic[project_shortname][-1] + '.docx'
        else:
            df_name = report_dic[project_shortname][0] + '.xlsx'
            open_name = doc_dic[project_shortname][-1] + '.docx'
    elif dic_client[sample_code]['代理商'] == 'boruilin':
        if 'DNA' in suffix and 'RNA' in suffix:
            df_name = report_dic[project_shortname][1] + '.xlsx'
            open_name = doc_dic[project_shortname][1] + '.docx'
        elif 'RNA' in suffix:
            df_name = report_dic[project_shortname][2] + '.xlsx'
            open_name = doc_dic[project_shortname][2] + '.docx'
        else:
            df_name = report_dic[project_shortname][0] + '.xlsx'
            open_name = doc_dic[project_shortname][0] + '.docx'
    elif dic_client[sample_code]['代理商'] == 'beijing':
        if 'DNA' in suffix and 'RNA' in suffix:
            df_name = report_dic[project_shortname][1] + '.xlsx'
            open_name = doc_dic[project_shortname][1] + '.docx'
        elif 'RNA' in suffix:
            df_name = report_dic[project_shortname][2] + '.xlsx'
            open_name = doc_dic[project_shortname][2] + '.docx'
        else:
            df_name = report_dic[project_shortname][0] + '.xlsx'
            open_name = doc_dic[project_shortname][0] + '.docx'
    # print('df_name',df_name)
    # print('open_name',open_name)
    result_report_name = filename_date + '_'+ sample_code + '_' + str(dic_client[sample_code]['姓名'])+'_Seq&Treat病原微生物基因检测'+ suffix + '.docx'
    # print(result_report_name)
    manufacturer = dic_client[sample_code]['代理商']
    # print(manufacturer)
    save_path = os.path.join(args.output_dir, manufacturer,filename_date)
    # print(save_path)
    if save_path and not os.path.exists(save_path):
        os.makedirs(save_path)
    open_docx_path = os.path.join(args.word_template_folder,manufacturer,open_name)
    # print('open_docx_path ',open_docx_path)
    if open_docx_path and os.path.exists(open_docx_path):
        doc = DocxTemplate(open_docx_path)
    else:
        logging.info(f"生成{str(dic_client[sample_code]['姓名'])}时发现，检测项目填写有误！")
        return
    sample_result_name_list = Nor(dic_client[sample_code]['检测微生物']).split(",")
    if Nor(dic_client[sample_code]['补充微生物']) != 'na':
        supply_result_name_list = Nor(dic_client[sample_code]['补充微生物']).split(",")
        sample_result_name_list.extend(supply_result_name_list)
    # sample_result_name_list = change_bacteria_list(sample_result_name_list)
    if 'na' in sample_result_name_list:
        sample_result_name_list.remove('na')
    excel_reader = pd.ExcelFile(os.path.join(args.excel_template_folder,df_name))
    sheet_name_list = excel_reader.sheet_names
    # print('sheet_name_list',sheet_name_list)
    df1 = excel_reader.parse(sheet_name=sheet_name_list[0])
    # print('df1',df1)#读取细菌的excel
    try:
        df2 = excel_reader.parse(sheet_name=sheet_name_list[1])
        # print('df2',df2)#读取真菌
    except IndexError:
        df2 = []   
    try:
        df3 = excel_reader.parse(sheet_name=sheet_name_list[2])
        # print('df3',df3)#读取其他类型病原微生物
    except IndexError:
        df3 = []   
    try:
        df4 = excel_reader.parse(sheet_name=sheet_name_list[3])
        # print('df41',df4)#读取人体共生菌
    except IndexError:
        df4 = []
        # print('df42',df4)
    logging.info(f"{result_report_name}开始生成")
    for result_name in sample_result_name_list:
        print('result_name',result_name)
        try:
            chinese_name = medical_DB.loc[(medical_DB['英文名称']== Nor(result_name)) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '种'].iloc[0]
            print('chinese_name',chinese_name)
            if chinese_name != 'NA':
                genus_name = medical_DB.loc[(medical_DB['英文名称']== Nor(result_name)) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '属'].iloc[0]
                # print('genus_name',genus_name)
                if genus_name != 'NA' or '病毒' in chinese_name:
                    micro_type = medical_DB.loc[(medical_DB['英文名称']== Nor(result_name)) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '类型'].iloc[0]
                    # print('micro_type',micro_type)
                    classification = medical_DB.loc[(medical_DB['英文名称']== Nor(result_name)) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '分类'].iloc[0]
                    print('classification',classification)
                    pathogenicity_info = medical_DB.loc[(medical_DB['英文名称']== Nor(result_name)) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '备注'].iloc[0]
                    # print('pathogenicity_info',pathogenicity_info)
                    if pathogenicity_info == '人体共生菌':
                        col_numbers = 3
                        # print('col_numbers',col_numbers)
                        # print('df4',df4)
                        df4 = add_micro(df4, chinese_name, genus_name, micro_type, pathogenicity_info, col_numbers, classification)
                    elif classification == sheet_name_list[0]:
                        col_numbers = 4
                        df1 = add_micro(df1, chinese_name, genus_name, micro_type, pathogenicity_info, col_numbers, classification)
                        print('df1',df1)
                    elif classification == sheet_name_list[1]:
                        col_numbers = 4
                        df2 = add_micro(df2, chinese_name, genus_name, micro_type, pathogenicity_info, col_numbers, classification)              
                    else:
                        col_numbers = 4
                        df3 = add_micro(df3, chinese_name, genus_name, micro_type, pathogenicity_info, col_numbers, classification)
                else:
                    logging.info(f"生成{result_report_name}时发现，数据库中{chinese_name}没有属名,请手动添加")
            else:
                logging.info(f"生成{result_report_name}时发现，数据库中{result_name}没有中文名,请手动添加")
        except IndexError:
            if result_name != 'NA':
                print('NA   result_name',result_name)
                logging.info(f"生成{result_report_name}时发现，数据库中不存在{result_name},或数据库中{result_name}检测项目信息有问题")
    df1['致病性'] = df1['致病性'].map(lambda x: x.replace("条件致病菌","条件致病"))
    if type(df2) != list:
        df2['致病性'] = df2['致病性'].map(lambda x: x.replace("条件致病菌","条件致病")) 
    if type(df3) != list:
        df3['致病性'] = df3['致病性'].map(lambda x: x.replace("条件致病菌","条件致病"))
    if type(df4) == list:
        if type(df3) == list:
            if type(df2) == list:
                appendix_list = [df1]
            else:
                appendix_list = [df1,df2]
        else:
            appendix_list = [df1,df2,df3]
    else:
        appendix_list = [df1,df2,df3,df4]
    # print(df1)
    use_col_width_dic_list = col_width_dic_list
    use_expect_text_list_dict = expect_text_list_dict
    if 'RNA' in suffix and 'DNA' not in suffix:
        if manufacturer == 'seegene':
            use_expect_text_list_dict = {'seegene':['RNA常见病毒筛查范围']}
            use_col_width_dic_list = [{0: Cm(2.0), 1: Cm(3.5), 2: Cm(2), 3: Cm(1.4), 4: Cm(2.0), 5: Cm(3.5), 6: Cm(2), 7: Cm(1.4)}]
        if manufacturer == 'boruilin':
            use_expect_text_list_dict = {'boruilin':['RNA常见病毒筛查范围']}
            use_col_width_dic_list = [{0: Cm(2.0), 1: Cm(3.5), 2: Cm(2), 3: Cm(1.4), 4: Cm(2.0), 5: Cm(3.5), 6: Cm(2), 7: Cm(1.4)}]
        if manufacturer == 'beijing':
            use_expect_text_list_dict = {'beijing':['RNA常见病毒筛查范围']}
            use_col_width_dic_list = [{0: Cm(2.0), 1: Cm(3.5), 2: Cm(2), 3: Cm(1.4), 4: Cm(2.0), 5: Cm(3.5), 6: Cm(2), 7: Cm(1.4)}]
    for index,df in enumerate(appendix_list):
        while (df.shape[0]) % 2 != 0 or df.iloc[int((df.shape[0])/2),1][0].isspace():
            df.loc[df.shape[0]]=['--'] * df.shape[1]
        row_num = int((df.shape[0])/2) + 1
        col_num = 8
        if index == 3:
            col_num = 6
        table = doc.add_table(rows=row_num, cols=col_num)
        col_width_dic = use_col_width_dic_list[index]
        table = change_type(df,table,col_width_dic)
        # print(result_report_name)
        # print(table)
        target = None
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            # print(paragraph_text)
            #print('use_expect_text_list_dict[manufacturer][index]',use_expect_text_list_dict[manufacturer][index])
            if paragraph_text.endswith(use_expect_text_list_dict[manufacturer][index]):
                # print(paragraph.text)
                target = paragraph
                break
        move_table_after(table, target)
    # print(result_report_name)
    doc.render(dic_client[sample_code])
    result_file = os.path.join(save_path, result_report_name)
    # print('result_file',result_file)
    doc.save(result_file)
    doc = DocxTemplate(result_file)
    # for paragraph in doc.paragraphs:
    #     print(paragraph.text)
    if dic_client[sample_code]['代理商'] == 'seegene' or 'boruilin' or 'beijing':
        if 'DNA' in suffix and 'RNA' in suffix:
            form_modification(doc=doc,dic_client=dic_client,sample_code=sample_code)
        elif 'RNA' in suffix:
            form_modification(doc=doc,dic_client=dic_client,sample_code=sample_code)
        else:
            form_modification2(doc=doc,dic_client=dic_client,sample_code=sample_code)
        # form_modification2(doc=doc,dic_client=dic_client,sample_code=sample_code)
        if dic_client[sample_code]['检测微生物'] != ',' and dic_client[sample_code]['检测微生物'] != '，' and dic_client[sample_code]['检测微生物'] != 'NA':
            Interpretation_list = Interpretation_addition(sample_code=sample_code,dic_client=dic_client,doc=doc,insert_info=dic_client[sample_code]['注释'])
            # print('dic_client[sample_code][注释]',dic_client[sample_code]['注释'])
            parse_table_add(sample_code=sample_code,dic_client=dic_client,doc=doc,medical_DB=medical_DB,Interpretation_list=Interpretation_list)#添加医学解读部分出问题
        doc = move_picture(doc=doc,png_name=dic_client[sample_code]['图片文件名'])
    # print(result_report_name)
    # doc.save(result_report_name)
    # print(dic_client[sample_code])
    # doc.render(dic_client[sample_code])
    # for paragraph in doc.paragraphs:
    #     print(paragraph.text)
    # result_file = os.path.join(save_path, result_report_name)
    # print(dic_client[sample_code]['姓名'])
    color_change3(doc=doc)
    doc.save(result_file)
    if result_file and os.path.exists(result_file):
        logging.info(f"{result_report_name}生成成功！")


def make_excel_dict(df:pd.Series,
    excel_list: list):
    line_info = str(df['样本编号']) + '_' + str(df['患者姓名']) + '_' + str(df['basecalling data']) + '_' + df['Reports']
    # print(df['样本编号'])
    # print(df['患者姓名'])
    # print(df['basecalling data'])
    # print(df['Reports'])
    excel_list.append(line_info)
    return excel_list
    
####################################################
# 统一输入的英文名
medical_DB['种'] = medical_DB['种'].apply(Nor_col)
medical_DB['英文名称'] = medical_DB['英文名称'].apply(Nor)
info_client['样本编号'] = info_client['样本编号'].apply(Nor_col)
info_client['患者姓名'] = info_client['患者姓名'].apply(Nor_col)
info_client['代理商'] = info_client['代理商'].apply(Nor_col)
Interpretation['name'] = Interpretation['name'].apply(Nor)
complex_df['name'] = complex_df['name'].apply(Nor)
info_client['采样时间'] = info_client['采样时间'].map(lambda x: str(x).split(' ')[0])
# 查询模板的字典
doc_dic = {'XY': ['Seq&Treat血液系统病原微生物检测DNA','Seq&Treat血液系统病原微生物检测DNA+RNA','Seq&Treat血液系统病原微生物检测RNA','Seq&Treat血液系统病原微生物检测'],
            'CK': ['Seq&Treat创口系统病原微生物检测DNA','Seq&Treat创口系统病原微生物检测DNA+RNA','Seq&Treat创口系统病原微生物检测RNA','Seq&Treat创口系统病原微生物检测'],
            'HX': ['Seq&Treat呼吸系统病原微生物检测DNA','Seq&Treat呼吸系统病原微生物检测DNA+RNA','Seq&Treat呼吸系统病原微生物检测RNA','Seq&Treat呼吸系统病原微生物检测'],
            'MN': ['Seq&Treat泌尿生殖系统病原微生物检测DNA','Seq&Treat泌尿生殖系统病原微生物检测DNA+RNA','Seq&Treat泌尿生殖系统病原微生物检测RNA','Seq&Treat泌尿生殖系统病原微生物检测'],
            'XH': ['Seq&Treat消化系统病原微生物检测DNA','Seq&Treat消化系统病原微生物检测DNA+RNA','Seq&Treat消化系统病原微生物检测RNA','Seq&Treat消化系统病原微生物检测'],
            'YB': ['Seq&Treat眼科系统病原微生物检测DNA','Seq&Treat眼科系统病原微生物检测DNA+RNA','Seq&Treat眼科系统病原微生物检测RNA','Seq&Treat眼科系统病原微生物检测'],
            'SJ': ['Seq&Treat神经系统病原微生物检测DNA','Seq&Treat神经系统病原微生物检测DNA+RNA','Seq&Treat神经系统病原微生物检测RNA','Seq&Treat神经系统病原微生物检测'],
            'XF': ['Seq&Treat胸腹腔系统病原微生物检测DNA','Seq&Treat胸腹腔系统病原微生物检测DNA+RNA','Seq&Treat胸腹腔系统病原微生物检测RNA','Seq&Treat胸腹腔系统病原微生物检测']
}
# 查询结果名称的字典
report_dic = {'XY': ['Seq&Treat血液系统病原微生物检测DNA','Seq&Treat血液系统病原微生物检测DNA+RNA','Seq&Treat血液系统病原微生物检测RNA'],
            'CK': ['Seq&Treat创口系统病原微生物检测DNA','Seq&Treat创口系统病原微生物检测DNA+RNA','Seq&Treat创口系统病原微生物检测RNA'],
            'HX': ['Seq&Treat呼吸系统病原微生物检测DNA','Seq&Treat呼吸系统病原微生物检测DNA+RNA','Seq&Treat呼吸系统病原微生物检测RNA'],
            'MN': ['Seq&Treat泌尿生殖系统病原微生物检测DNA','Seq&Treat泌尿生殖系统病原微生物检测DNA+RNA','Seq&Treat泌尿生殖系统病原微生物检测RNA'],
            'XH': ['Seq&Treat消化系统病原微生物检测DNA','Seq&Treat消化系统病原微生物检测DNA+RNA','Seq&Treat消化系统病原微生物检测RNA'],
            'YB': ['Seq&Treat眼科系统病原微生物检测DNA','Seq&Treat眼科系统病原微生物检测DNA+RNA','Seq&Treat眼科系统病原微生物检测RNA'],
            'SJ': ['Seq&Treat神经系统病原微生物检测DNA','Seq&Treat神经系统病原微生物检测DNA+RNA','Seq&Treat神经系统病原微生物检测RNA'],
            'XF': ['Seq&Treat胸腹腔系统病原微生物检测DNA','Seq&Treat胸腹腔系统病原微生物检测DNA+RNA','Seq&Treat胸腹腔系统病原微生物检测RNA']
}
# 调用模板的字典
doc_list = ['seegene','beagle','boruilin','beijing']
# 生成附录的对应位置和附录表格的列宽
expect_text_list_dict = {'beijing':['常见细菌筛查范围', '常见真菌筛查范围','常见其他病原微生物筛查范围','常见人体共生菌筛查范围'],'boruilin':['常见细菌筛查范围', '常见真菌筛查范围','常见其他病原微生物筛查范围','常见人体共生菌筛查范围'],'beagle':['常见细菌筛查范围', '常见真菌筛查范围','其他病原微生物','常见人体共生菌'],'seegene':['常见细菌筛查范围', '常见真菌筛查范围','常见其他病原微生物筛查范围','常见人体共生菌筛查范围']}
col_width_dic_list = [{0: Cm(1.1), 1: Cm(4.4), 2: Cm(2), 3: Cm(1.4), 4: Cm(1.1), 5: Cm(4.4), 6: Cm(2), 7: Cm(1.4)}, {0: Cm(2.1), 1: Cm(3.4), 2: Cm(2), 3: Cm(1.4), 4: Cm(2.1), 5: Cm(3.4), 6: Cm(2), 7: Cm(1.4)},{0: Cm(2.4), 1: Cm(3.25), 2: Cm(1.8), 3: Cm(1.45), 4: Cm(2.4), 5: Cm(3.25), 6: Cm(1.8), 7: Cm(1.45)},{0: Cm(1.4), 1: Cm(5.1), 2: Cm(2.4), 3: Cm(1.4), 4: Cm(5.1), 5: Cm(2.4)}]
# 导入正文部分的字典生成
dic_client = {}
for row_index in range(info_client.shape[0]):
    if info_client['正式报告结果'][row_index] != 'NA' or info_client['补充报告结果'][row_index] != 'NA':
        # for company in doc_list:
        #     if result_file_name.startswith(company):
        #         manufacturer = company
        # if info_client['代理商'][row_index].lower() in doc_list:
        #     manufacturer = info_client['代理商'][row_index].lower()
        if info_client['代理商'][row_index].lower() in doc_list:
            manufacturer = info_client['代理商'][row_index].lower()
        else:
            for company in doc_list:
                if result_file_name.startswith(company):
                    manufacturer = company
        dic_client.setdefault(info_client['样本编号'][row_index], {})
        dic_client[info_client['样本编号'][row_index]]['样本编号'] = info_client['样本编号'][row_index].split('-')[0]
        dic_client[info_client['样本编号'][row_index]]['代理商'] = manufacturer
        dic_client[info_client['样本编号'][row_index]]['姓名'] = info_client['患者姓名'][row_index]
        print(dic_client[info_client['样本编号'][row_index]]['姓名'])
        logging.info(f"开始处理{info_client['样本编号'][row_index]} {info_client['患者姓名'][row_index]}的信息")
        dic_client[info_client['样本编号'][row_index]]['性别'] = info_client['性别'][row_index]
        try:
            dic_client[info_client['样本编号'][row_index]]['年龄'] = int(info_client['年龄'][row_index])
        except ValueError:
            dic_client[info_client['样本编号'][row_index]]['年龄'] = info_client['年龄'][row_index]
        dic_client[info_client['样本编号'][row_index]]['送检单位'] = info_client['医院'][row_index]
        dic_client[info_client['样本编号'][row_index]]['科室'] = info_client['科室'][row_index]
        dic_client[info_client['样本编号'][row_index]]['检测项目'] = info_client['检测项目'][row_index].strip()
        dic_client[info_client['样本编号'][row_index]]['科室'] = info_client['科室'][row_index].strip()
        # if dic_client[info_client['样本编号'][row_index]]['降钙素原'] != 'NA':
        #     dic_client[info_client['样本编号'][row_index]]['降钙素原'] = str(info_client['降钙素原(PCT)'][row_index]) + 'ng/mL'
        # else:
        #     dic_client[info_client['样本编号'][row_index]]['降钙素原'] = info_client['降钙素原(PCT)'][row_index]
        # if dic_client[info_client['样本编号'][row_index]]['白细胞'] != 'NA':
        #     dic_client[info_client['样本编号'][row_index]]['白细胞'] = str(info_client['白细胞(WBC)'][row_index]) + '×10-9/L'
        # else:
        #     dic_client[info_client['样本编号'][row_index]]['白细胞'] = info_client['白细胞(WBC)'][row_index]
        # if dic_client[info_client['样本编号'][row_index]]['降钙素原'] != 'NA':
        #     dic_client[info_client['样本编号'][row_index]]['降钙素原'] = str(info_client['降钙素原(PCT)'][row_index]) + 'ng/mL'
        # else:
        #     dic_client[info_client['样本编号'][row_index]]['降钙素原'] = info_client['降钙素原(PCT)'][row_index]
        dic_client[info_client['样本编号'][row_index]]['降钙素原'] = info_client['降钙素原(PCT)'][row_index]
        dic_client[info_client['样本编号'][row_index]]['白细胞'] = info_client['白细胞(WBC)'][row_index]
        dic_client[info_client['样本编号'][row_index]]['反应蛋白'] = info_client['C-反应蛋白(CRP)'][row_index]
        dic_client[info_client['样本编号'][row_index]]['培养结果'] = info_client['培养结果'][row_index].strip()
        if 'DNA+RNA' in dic_client[info_client['样本编号'][row_index]]['检测项目']:
            dic_client[info_client['样本编号'][row_index]]['核酸类型'] = 'DNA+RNA'
        elif 'RNA' in dic_client[info_client['样本编号'][row_index]]['检测项目']:
            dic_client[info_client['样本编号'][row_index]]['核酸类型'] = 'RNA'
        else:
            dic_client[info_client['样本编号'][row_index]]['核酸类型'] = 'DNA'
        dic_client[info_client['样本编号'][row_index]]['样本类型'] = info_client['样本类型'][row_index]
        if '血' in dic_client[info_client['样本编号'][row_index]]['样本类型']:
            dic_client[info_client['样本编号'][row_index]]['样本类型'] = '血液'
        dic_client[info_client['样本编号'][row_index]]['临床诊断'] = info_client['临床诊断'][row_index]
        dic_client[info_client['样本编号'][row_index]]['抗感染用药史'] = info_client['抗感染用药史'][row_index].strip()
        dic_client[info_client['样本编号'][row_index]]['重点关注病原菌'] = info_client['重点关注病原菌'][row_index].strip()
        dic_client[info_client['样本编号'][row_index]]['检测微生物'] = info_client['正式报告结果'][row_index]
        dic_client[info_client['样本编号'][row_index]]['补充微生物'] = info_client['补充报告结果'][row_index]
        receivedate = str(info_client['收样时间'][row_index]).strip()
        if receivedate != 'NA':
            dic_client[info_client['样本编号'][row_index]]['收样日期'] = receivedate.split('/')[0]+'年'+receivedate.split('/')[1]+'月'+receivedate.split('/')[2]+'日'
        else:
            dic_client[info_client['样本编号'][row_index]]['收样日期'] = 'NA'
        collectiondate = str(info_client['采样时间'][row_index]).strip()
        if collectiondate != 'NA':
            try:
                dic_client[info_client['样本编号'][row_index]]['采样日期'] = collectiondate.split('/')[0]+'年'+collectiondate.split('/')[1]+'月'+collectiondate.split('/')[2]+'日'
            except IndexError:
                dic_client[info_client['样本编号'][row_index]]['采样日期'] = collectiondate.split('-')[0]+'年'+collectiondate.split('-')[1]+'月'+collectiondate.split('-')[2]+'日'
        else:
            dic_client[info_client['样本编号'][row_index]]['采样日期'] = 'NA'
        try:
            exdate = str(info_client['上机日期'][row_index]).split(' ')[0]
            dic_client[info_client['样本编号'][row_index]]['实验日期'] = exdate.split('-')[0]+'年'+exdate.split('-')[1]+'月'+exdate.split('-')[2]+'日'
        except IndexError:
            exdate = str(info_client['上机日期'][row_index])
            dic_client[info_client['样本编号'][row_index]]['实验日期'] = exdate.split('/')[0]+'年'+exdate.split('/')[1]+'月'+exdate.split('/')[2]+'日'
        reportdate = str(info_client['报告日期'][row_index]).split(' ')[0]
        # print(reportdate)
        dic_client[info_client['样本编号'][row_index]]['报告日期'] = reportdate.split('/')[0]+'年'+reportdate.split('/')[1]+'月'+reportdate.split('/')[2]+'日'
        colname = info_client['barcode'][row_index] + '_' + info_client['患者姓名'][row_index] + '_' + info_client['样本编号'][row_index]
        dic_client[info_client['样本编号'][row_index]]['平均长度'] = picture_dtat_df.at[8, colname]
        if manufacturer == 'seegene' or 'boruilin' or 'beijing':
            table_info = table2_context(sample_code=info_client['样本编号'][row_index], info_client=info_client, Interpretation=Interpretation, medical_DB=medical_DB)
            result_table = table_info[0]
            dic_client[info_client['样本编号'][row_index]]['表1信息'] = result_table[0]
            dic_client[info_client['样本编号'][row_index]]['表2信息'] = result_table[1]
            dic_client[info_client['样本编号'][row_index]]['表3信息'] = result_table[2]
            dic_client[info_client['样本编号'][row_index]]['表4信息'] = result_table[3]
            dic_client[info_client['样本编号'][row_index]]['表5信息'] = result_table[4]
            dic_client[info_client['样本编号'][row_index]]['表6信息'] = result_table[5]
            dic_client[info_client['样本编号'][row_index]]['表7信息'] = result_table[6]
            dic_client[info_client['样本编号'][row_index]]['表9信息'] = table_info[1]
            if type(drug_resistance_df) != int:
                dic_client[info_client['样本编号'][row_index]]['表8信息'] = table7_make(sample_code=info_client['样本编号'][row_index], info_client=info_client, drug_resistance_df=drug_resistance_df)
            else:
                dic_client[info_client['样本编号'][row_index]]['表8信息'] = [{'基因': '--', '药物': '--'}]
            dic_client[info_client['样本编号'][row_index]]['example'] = clinical(sample_code=info_client['样本编号'][row_index],info_client=info_client,medical_DB=medical_DB,manufacturer=manufacturer)
            length_colname = str(info_client['barcode'][row_index]) + '_'  + dic_client[info_client['样本编号'][row_index]]['姓名'] + '_' + str(info_client['样本编号'][row_index])
            dic_client[info_client['样本编号'][row_index]]['图片文件名'] = length_colname + '.png'
            make_picture(length_colname=length_colname,picture_dtat_df=picture_dtat_df)
            if info_client['正式报告结果'][row_index] != ',' and info_client['正式报告结果'][row_index] != '，' and info_client['正式报告结果'][row_index] != 'NA':
                dic_client[info_client['样本编号'][row_index]]['注释'] = '注：常用药物为临床常规药物，且无法覆盖药敏结果，具体用药请结合临床药敏结果或医院耐药监测数据酌情用药。'
            else:
                dic_client[info_client['样本编号'][row_index]]['注释'] = '本次样本中未检出疑似致病菌，结果仅对本次送检的样本负责，请临床根据患者症状结合其他检测结果进行综合判断。'
            if info_client['补充报告结果'][row_index] != ',' and info_client['补充报告结果'][row_index] != '，' and info_client['补充报告结果'][row_index] != 'NA':
                dic_client[info_client['样本编号'][row_index]]['说明'] = '疑似微生物种解释说明'
        elif manufacturer == 'beagle':
            table_info = table_context(sample_code=info_client['样本编号'][row_index], info_client=info_client, Interpretation=Interpretation, medical_DB=medical_DB)
            dic_client[info_client['样本编号'][row_index]]['表1信息'] = table_info[0]
            dic_client[info_client['样本编号'][row_index]]['表2信息'] = table_info[1]
            if info_client.loc[info_client['样本编号'] == info_client['样本编号'][row_index], '正式报告结果'].iloc[0] != 'NA':
                dic_client[info_client['样本编号'][row_index]]['example'] = dic_client[info_client['样本编号'][row_index]]['example'] = clinical(sample_code=info_client['样本编号'][row_index],info_client=info_client,medical_DB=medical_DB,manufacturer=manufacturer)
                dic_client[info_client['样本编号'][row_index]]['reference'] = reference(info_client['样本编号'][row_index])
# for key,value in dic_client.items():
#     print(key,value)
logging.info("所有信息处理完成！")
# # 输出文件夹的创建
filename_date = reportdate.split('/')[0]+reportdate.split('/')[1]+reportdate.split('/')[2]
# # if __name__=='__main__':
# #     p = Pool(args.processes_number)     
# #     for sample_code in dic_client.keys():
# #         p.apply_async(make_word_report, args=(sample_code,dic_client,report_dic,doc_dic,medical_DB))
# #     p.close()
# #     p.join()
# # logging.info(f"所有报告生成成功！")
for sample_code,value in dic_client.items():
    # print(value)
    make_word_report(sample_code=sample_code,dic_client=dic_client,report_dic=report_dic,doc_dic=doc_dic,medical_DB=medical_DB)
logging.info(f"开始汇总结果！")
for sample_code in dic_client:
    manufacturer = dic_client[sample_code]['代理商']
    summary_excel_path = os.path.join(args.summary_excel,manufacturer,'汇总.xlsx')
    all_summary_df = pd.read_excel(summary_excel_path)
    all_info_list = []
    if all_summary_df.shape[0] != 0:
        for row_index in range(all_summary_df.shape[0]):
            line_info = str(all_summary_df['样本编号'][row_index]) + '_' + str(all_summary_df['患者姓名'][row_index]) + '_' + str(all_summary_df['basecalling data'][row_index]) + '_' + all_summary_df['Reports'][row_index]
            all_info_list.append(line_info)
    all_info_dict = dict(zip(all_info_list,all_info_list))
    data_summary_excel_path = os.path.join(args.summary_excel,manufacturer,filename_date,'汇总.xlsx')
    if data_summary_excel_path and os.path.exists(data_summary_excel_path):
        data_summary_df = pd.read_excel(data_summary_excel_path)
    else:
        data_summary_df = pd.DataFrame(columns=all_summary_df.columns.tolist())
    
    need_info = info_client.query('样本编号 == @sample_code').iloc[0]
    all_list = []
    try:
        all_list = need_info['正式报告结果'].split(',')    
    except AttributeError:
        all_list = []
    try:
        all_list.extend(need_info['补充报告结果'].split(','))
    except AttributeError:
        all_list = all_list
    for micro in all_list:
        # print(micro)
        compare_bac_name = Nor(micro)
        data_info_list = []
        if data_summary_df.shape[0] != 0:
            for row_index in range(data_summary_df.shape[0]):
                line_info = str(data_summary_df['样本编号'][row_index]) + '_' + str(data_summary_df['患者姓名'][row_index]) + '_' + str(data_summary_df['basecalling data'][row_index]) + '_' + data_summary_df['Reports'][row_index]
                data_info_list.append(line_info)
        data_info_dict = dict(zip(data_info_list,data_info_list))
        if len(micro) != 0 and micro != 'NA':
            info_list = all_summary_df.columns.tolist()[:-10]
            new_line_dict = {}
            for colname in info_list:   
                new_line_dict[colname] = need_info[colname]
            new_line_dict['Reports'] = micro
            try:
                new_line_dict['报告结果'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]
            except IndexError:
                new_line_dict['报告结果'] = 'NA'
            try:
                new_line_dict['物种'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '分类'].iloc[0]
            except IndexError:
                new_line_dict['物种'] = 'NA'
            if micro in need_info['正式报告结果']:
                new_line_dict['报告等级'] = 'A'
            else:
                new_line_dict['报告等级'] = 'B'
            try:
                new_line_dict['致病性描述'] = medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '备注'].iloc[0]
            except IndexError:
                new_line_dict['致病性描述'] = 'NA'
            column_R = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_R_' + project_shorthand(sample_code) + '_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
            try:
                new_line_dict['Reads'] = Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0]
            except KeyError:
                column_R = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_R_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
                new_line_dict['Reads'] = Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0]
            except IndexError:
                new_line_dict['Reads'] = 'NA'
            new_line_dict['样本描述（解读人员手动添加）'] = ' '
            new_line_dict['制作人'] = ' '
            new_line_dict['审核人'] = ' '
            new_line_dict['备 注'] = ' '
            new_df = (new_line_dict)
            new_line_info = str(new_line_dict['样本编号']) + '_' + str(new_line_dict['患者姓名']) + '_' + str(new_line_dict['basecalling data']) + '_' + new_line_dict['Reports']
            # print(data_info_dict)
            # print(all_info_dict)
            # print(new_line_info)
            if new_line_info not in data_info_dict:
                data_summary_df = data_summary_df.append(new_df, ignore_index=True)
            if new_line_info not in all_info_dict:
                all_summary_df = all_summary_df.append(new_df, ignore_index=True)
    data_summary_df.to_excel(data_summary_excel_path, index=False)
    all_summary_df.to_excel(summary_excel_path, index=False)
logging.info(f"汇总完成！")
