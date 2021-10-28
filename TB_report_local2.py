#!/usr/bin/env python

import argparse
import pandas as pd
import os
import docx
import logging
import datetime
import math
from docxtpl import DocxTemplate, RichText
from docx.shared import RGBColor
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Cm

# 打印出运行的时间
time1 = '运行时间：' + str(datetime.datetime.now())
# print(time1)

# 设定监控日志输出文件名和内容形式
logging.basicConfig(format='%(asctime)s - %(message)s', filename='/mnt/c/Users/luping/Desktop/报告流程/TB/OUTPUT/运行信息.txt', filemode='a', level=logging.INFO)

# 参数的导入与处理
parser = argparse.ArgumentParser()
parser.add_argument('-i', "--result_excel", required=True, help="the excel file with the result selected")
parser.add_argument('-b', "--database", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/TB/思可愈数据库-TNP-Seq病原菌测序项目2021.05.18.xlsx',help="database provided by the Ministry of Medicine")
# parser.add_argument('-w', "--word_template_folder", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/TB/分枝杆菌(MTB&NTM)及其耐药基因检测报告.docx',help="folder where all word report templates are located")
# parser.add_argument('-e', "--excel_template_folder", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/TB/报告模板/excel/',help="folder where all excel report templates are located")
# parser.add_argument('-n', "--processes_number", type=int, default=5,help="并行进程数目")
parser.add_argument('-c', "--complex_excel", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/TB/mycobacterium_tuberculosis_complex.xlsx',help="结核分枝杆菌复合群包含微生物表格")
parser.add_argument('-s', "--summary_excel", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/TB/OUTPUT/',help="summary documents before processing")
parser.add_argument('-m', "--name_excel", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/TB/药品中英文对照表.xlsx',help="检测药品中英文对照表")
parser.add_argument('-o', "--output_dir", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/TB/OUTPUT/',help="supplement sample result")
parser.add_argument('-B', "--barcode_picutre", type=str, default='/mnt/c/Users/luping/Desktop/报告流程/RD/barcode/',help="条形码图片所在")
args = parser.parse_args()


####################################################
# 定义函数
# 标准化输入微生物名称所用函数
def change_bacteria_list(bacteria_list: list):
    new_bacteria_list = []
    for bacteria in bacteria_list:
        compare_bacteria: str = Nor(bacteria)
        find_complex_df = complex_df[complex_df['name'] == compare_bacteria]
        if find_complex_df.shape[0] == 0:
            new_bacteria_list.append(bacteria)
        else:
            new_bacteria_list.append('Mycobacterium tuberculosis complex')
    return new_bacteria_list


def Nor(x: str
    ) -> str:
    first: str= x.strip()
    standardized_string: str = " ".join(first.split())
    standardized_string: str = standardized_string.lower()
    return standardized_string


# 标准化输入列名称所用函数
def Nor_col(x: str
    ) -> str:
    first: str= str(x).strip()
    standardized_string: str = " ".join(first.split())
    return standardized_string


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def remove_item(old_list: list,
    remove_element: str):
    new_list = []
    for element in old_list:
        if element != remove_element:
            new_list.append(element)
    return new_list


def table_1_make(sample_code: str,
    medical_DB: pd.DataFrame,
    complex_df: pd.DataFrame,
    all_bac: list,
    mic_dict: dict,
    result_list: list,
    Interpretation: pd.DataFrame
    ) -> list:
    sample_result_list = []
    bac_list = []
    column_R = info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0] + '_R_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
    # print(column_R)
    for bac_name in result_list:
        # print(bac_name)
        if len(bac_name) != 0:
            compare_bac_name = Nor(bac_name)
            dic_bac: dict ={}
            find_complex_df = complex_df[complex_df['name'] == compare_bac_name]
            if find_complex_df.shape[0] == 0:
                dic_bac['分类'] = all_bac[1]
            else:
                dic_bac['分类'] = all_bac[0]
            dic_bac['检测结果'] = '阳性'
            try:
                dic_bac['中文名'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]
                dic_bac['微生物'] = bac_name
                dic_bac['分类顺序'] = mic_dict[dic_bac['分类']]
            except IndexError:
                logging.info(f"{bac_name}在数据库中没有找到")
            try:
                dic_bac['序列数'] = Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0]
            except KeyError:
                logging.info(f"{bac_name}没有找到相应的序列数")
            bac_list.append(dic_bac['分类'])
            sample_result_list.append(dic_bac)
    bac_list = list(set(bac_list))
    for non_bac in all_bac:
        if non_bac not in bac_list:
            dic_bac = {}
            dic_bac['分类'] = non_bac
            dic_bac['分类顺序'] = mic_dict[dic_bac['分类']]
            dic_bac['检测结果'] = '阴性'
            dic_bac['中文名'] = '/'
            dic_bac['微生物'] = '/'
            dic_bac['序列数'] = '/'
            sample_result_list.append(dic_bac)
    sample_result_list: list = sorted(sample_result_list, key = lambda x:x['分类顺序'])
    return sample_result_list


def table_2_3_make(sample_code: str,
    complex_df: pd.DataFrame,
    AMR_df: pd.DataFrame,
    AMR_stat_colname: str,
    id_dict: dict
    ) -> list:
    AMR_stat_colname1 = AMR_stat_colname.replace('count','depth')
    # print(AMR_stat_colname1)
    number = 0
    for depth in AMR_df[AMR_stat_colname1].tolist():
        # print(depth)
        if str(depth).startswith(r'*'):
            # print(1)
            number += 1
            break
    if number != 0:
        # print(AMR_df)
        handle_df = AMR_df[AMR_df[AMR_stat_colname1].str.contains('\*', na=False)]
        table2_number = 0
        table3_number = 0
        table2_list = []
        table3_list = []
        handle_df.apply(table_2_3_make_dict, complex_df=complex_df,AMR_stat_colname=AMR_stat_colname,id_dict=id_dict,table2_list=table2_list,table3_list=table3_list, axis=1)
        for table_line in table2_list:
            table2_number += 1
            table_line['序号'] = table2_number
        for table_line in table3_list:
            table3_number += 1
            table_line['序号'] = table3_number 
        if len(table2_list) == 0:
            table2_list = [{'突变率': '-', '基因': '-', '突变描述': '-', '药品': '-', '氨基酸突变': '-', '序号': '-'}]
        if len(table3_list) == 0:
            table3_list = [{'突变率': '-', '基因': '-', '突变描述': '-', '药品': '-', '氨基酸突变': '-', '序号': '-'}]
    else:
        table2_list = [{'突变率': '-', '基因': '-', '突变描述': '-', '药品': '-', '氨基酸突变': '-', '序号': '-'}]
        table3_list = [{'突变率': '-', '基因': '-', '突变描述': '-', '药品': '-', '氨基酸突变': '-', '序号': '-'}]
    return [table2_list,table3_list]


def table_4_make(sample_code: str,
    medical_DB: pd.DataFrame,
    result_list: list
    ) -> list:
    sample_result_list = []
    bac_list = []
    for bac_name in result_list:
        # print(bac_name)
        if len(bac_name) != 0:
            compare_bac_name = Nor(bac_name)
            dic_bac: dict ={}
            try:
                dic_bac['名称'] = str(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]) + '(' + bac_name + ')'
                dic_bac['临床意义'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '临床意义'].iloc[0]
                dic_bac['推荐用药'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '常用药物'].iloc[0]
            except IndexError:
                logging.info(f"{bac_name}在数据库中没有找到")
            dic_bac['行名1'] = '检 出 菌 种'
            dic_bac['行名2'] = '临 床 意 义'
            dic_bac['行名3'] = '推 荐 用 药'
            sample_result_list.append(dic_bac)
    if len(sample_result_list) == 0:
        dic_bac: dict ={}
        dic_bac['行名1'] = '检 出 菌 种'
        dic_bac['行名2'] = '临 床 意 义'
        dic_bac['行名3'] = '推 荐 用 药'
        dic_bac['名称'] = '-'
        dic_bac['临床意义'] = '-'
        dic_bac['推荐用药'] = '-'
        sample_result_list.append(dic_bac)
    return sample_result_list


def table_2_3_make_dict(AMR_df: pd.Series,
    complex_df: pd.DataFrame,
    AMR_stat_colname: str,
    id_dict: dict,
    table2_list: list,
    table3_list: list,
    ) -> list:
    new_dict = {}
    compare_bac_name = Nor(id_dict[AMR_df['chr']])
    find_complex_df = complex_df[complex_df['name'] == compare_bac_name]
    AMR_stat_colname2 = AMR_stat_colname.replace('count','vaf')
    new_dict['突变率'] = AMR_df[AMR_stat_colname2]
    new_dict['基因'] = AMR_df['gene_name']
    new_dict['突变描述'] = AMR_df['hgvs_c']
    new_dict['氨基酸突变'] = AMR_df['hgvs_p']
    if 'NA' in new_dict['氨基酸突变']:
        new_dict['氨基酸突变']= '-'
    # print('AMR_df',AMR_df['drug'])
    medicine_list = AMR_df['drug'].split(';')
    # print('medicine_list',medicine_list)
    medicine_name_list = []
    for medicine in medicine_list:
        medicine_name = name_df.loc[name_df['药品名称'] == medicine, '中文名'].iloc[0]
        medicine_name_list.append(medicine_name)
    if len(medicine_name_list) == 1:
        new_dict['药品'] = medicine_name_list[0]
    else:
        new_dict['药品'] = ';'.join(medicine_name_list)
    if find_complex_df.shape[0] == 0:
        table3_list.append(new_dict)
    else:
        table2_list.append(new_dict)
    if len(table2_list) == 0:
        table2_list = [{'突变率': '-', '基因': '-', '突变描述': '-', '药品': '-', '氨基酸突变': '-', '序号': '-'}]
    if len(table3_list) == 0:
        table3_list = [{'突变率': '-', '基因': '-', '突变描述': '-', '药品': '-', '氨基酸突变': '-', '序号': '-'}]


def table_context(sample_code: str,
    info_client: pd.DataFrame,
    medical_DB: pd.DataFrame,
    complex_df: pd.DataFrame,
    id_dict: dict,
    AMR_df: pd.DataFrame,
    AMR_stat_colname: str,
    Interpretation: pd.DataFrame
    ) -> list:
    handle_df = info_client.query('样本编号 == @sample_code').iloc[0,:]
    all_bac: list = ['结核分枝杆菌复合群（MTBC）','非结核分枝杆菌（NTM）']
    mic_dict: dict = {'结核分枝杆菌复合群（MTBC）':1,'非结核分枝杆菌（NTM）':2}
    result_list = handle_df['正式报告结果'].split(',')
    table1_list = table_1_make(sample_code=sample_code, medical_DB=medical_DB, complex_df=complex_df, all_bac=all_bac, mic_dict=mic_dict, result_list=result_list, Interpretation=Interpretation)
    table4_list = table_4_make(sample_code=sample_code,medical_DB=medical_DB,result_list=result_list)
    result_list = table_2_3_make(sample_code=sample_code,complex_df=complex_df,AMR_df=AMR_df,AMR_stat_colname=AMR_stat_colname,id_dict=id_dict)
    table2_list = result_list[0]
    table3_list = result_list[1]
    return [table1_list,table2_list,table3_list,table4_list]



def make_picture(AMR_stat_colname: str,
    AMR_stat_df: pd.DataFrame,
    file_name: str
    ) -> None:
    R_out = file_name + '.r'
    rscript=f'''#! /path/to/Rscript
library(openxlsx)
library(ggplot2)
data<- read.xlsx('{AMR_stat_df}', sheet='depth_report')
names(data)[names(data) == '{file_name}'] <- 'Frequency'
data <- data[c(1:11),]
row.names(data) <- as.list(data)$depth
read_length_hist <- ggplot(data, mapping=aes(x=rownames(data), y=Frequency)) +
geom_bar(stat="identity", fill="#13A8B0", colour="#13A8B0") +
scale_x_discrete(limits=factor(rownames(data))) +
labs(x="Depth", y="Count") +
theme(panel.grid=element_blank(), panel.background=element_rect(color="black", fill="transparent")) + 
theme(axis.text =element_text(size=7))
ggsave(file="{file_name}.png",read_length_hist, width = 6, height = 3)
    '''
    out = open(R_out,'w')
    out.write(rscript)
    out.close()
    cmd = "Rscript " + R_out
    os.system(cmd)
    os.remove(R_out)


def move_picture(doc,
    png_name: str
    ) -> None:
    table = doc.add_table(rows=1, cols=3)
    cell = table.cell(0, 1)
    ph =cell.paragraphs[0]
    run = ph.add_run()
    run.add_picture(png_name)
    target = None
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        if '测序深度统计' in paragraph_text:
            # print(paragraph_text)
            target = paragraph
            break
    move_table_after(table, target)
    os.remove(png_name)
    return doc


# 查询检测项目的简称
def project_shorthand(sample_code: str) -> str:
    project_name = info_client.loc[info_client['样本编号'] == sample_code, '检测项目'].iloc[0].rstrip()
    hand = 0
    if '呼吸' in project_name:
        hand: str = 'HX'
    elif '血液' in project_name:
        hand: str = 'XY'
    elif '神经' in project_name:
        hand: str = 'SJ'
    elif '胸腹' in project_name:
        hand: str = 'XF'
    elif '泌尿' in project_name:
        hand: str = 'MN'
    elif '消化' in project_name:
        hand: str = 'XH'
    elif '创口' in project_name:
        hand: str = 'CK'
    elif '眼科' in project_name:
        hand: str = 'YB'
    else:
        name = info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]
    return hand


def appendix_drugs(table2_list: list,
    table3_list: list,
    first_list: list,
    second_list: list,
    NTM_list: list
    ) -> list:
    table2_drug_list = []
    for table_line in table2_list:
        new_drug_list = table_line['药品'].split(';')
        table2_drug_list.extend(new_drug_list)
    first_find_list = []
    first_dont_find_list = []
    for first_drug in first_list:
        if first_drug in table2_drug_list:
            first_find_list.append(first_drug)
        else:
            first_dont_find_list.append(first_drug)
    if len(first_find_list) == 0:
        first_find_list = ['/']
    if len(first_dont_find_list) == 0:
        first_dont_find_list = ['/']
    first_find = '\n'.join(first_find_list)
    first_dont_find = '\n'.join(first_dont_find_list)
    rt1 = RichText('')
    rt2 = RichText('')
    rt1.add(first_find)
    rt2.add(first_dont_find)
    second_find_list = []
    second_dont_find_list = []
    for second_drug in second_list:
        if second_drug in table2_drug_list:
            second_find_list.append(second_drug)
        else:
            second_dont_find_list.append(second_drug)
    if len(second_find_list) == 0:
        second_find_list = ['/']
    second_find = '\n'.join(second_find_list)
    number = math.floor(len(second_dont_find_list)/2)
    second_dont_find_one_list = second_dont_find_list[0:number]
    second_dont_find_two_list = second_dont_find_list[number:]
    if len(second_dont_find_one_list) == 0:
        second_dont_find_one_list = ['/']
    if len(second_dont_find_two_list) == 0:
        second_dont_find_two_list = ['/']
    second_dont_find_one = '\n'.join(second_dont_find_one_list)
    second_dont_find_two = '\n'.join(second_dont_find_two_list)
    rt3 = RichText('')
    rt4 = RichText('')
    rt5 = RichText('')
    rt3.add(second_find)
    rt4.add(second_dont_find_one)
    rt5.add(second_dont_find_two)
    table3_drug_list = []
    for table_line in table3_list:
        new_drug_list = table_line['药品'].split(';')
        table3_drug_list.extend(new_drug_list)
    third_find_list = []
    third_dont_find_list = []
    for third_drug in NTM_list:
        if third_drug in table3_drug_list:
            third_find_list.append(third_drug)
        else:
            third_dont_find_list.append(third_drug)
    if len(third_find_list) == 0:
        third_find_list = ['/']
    if len(third_dont_find_list) == 0:
        third_dont_find_list = ['/']
    third_find = '\n'.join(third_find_list)
    third_dont_find = '\n'.join(third_dont_find_list)
    rt6 = RichText('')
    rt7 = RichText('')
    rt6.add(third_find)
    rt7.add(third_dont_find)
    return [rt1,rt2,rt3,rt4,rt5,rt6,rt7]


def add_medical_interpretation(AMR_stat_colname: str,
    sample_type: str,
    all_result_list: list,
    MTB_result_list: list,
    NTM_result_list: list,
    id_dict: dict,
    MTB_gene_dict: dict,
    NTM_gene_dict: dict,
    AMR_df: pd.DataFrame):
    MTB_gene_found_list = []
    MTB_mutation_gene_found_list = []
    MTB_gene_unfound_list = []
    NTM_gene_found_list = []
    NTM_mutation_gene_found_list = []
    NTM_gene_unfound_list = []
    MTB_drugs_found_list = []
    MTB_mutation_drugs_found_list = []
    MTB_drugs_unfound_list = []
    NTM_drugs_found_list = []
    NTM_mutation_drugs_found_list = []
    NTM_drugs_unfound_list = []
    rt = RichText('')
    rt.add(f'该{sample_type}中,')
    unfound_info = ''
    handle_df = AMR_df[AMR_df[AMR_stat_colname] != 0]
    number = 0
    for result in all_result_list:
        if result['微生物'] == 'Mycobacterium tuberculosis':
            number += 1
            MTB_df = handle_df[handle_df['chr'] == 'NC_000962.3']
            MTB_genes = MTB_df['gene_name'].tolist()
            for key, value in MTB_gene_dict.items():
                if key in MTB_genes:
                    MTB_gene_found_list.append(key)
                    MTB_drugs_found_list.extend(value)
                else:
                    MTB_gene_unfound_list.append(key)
                    MTB_drugs_unfound_list.extend(value)
            MTB_drugs_found_list = list(set(MTB_drugs_found_list))
            MTB_drugs_unfound_list = list(set(MTB_drugs_unfound_list))
            for line in MTB_result_list:
                if line['序号'] != '-':
                    MTB_mutation_gene_found_list.append(line['基因'])
                    MTB_mutation_drugs_found_list.append(line['药品'])
            new_MTB_drugs_unfound_list = MTB_drugs_unfound_list
            new_MTB_gene_unfound_list = MTB_gene_unfound_list
            for key, value in MTB_gene_dict.items():
                for drug in value:
                    if drug in MTB_drugs_unfound_list and drug in MTB_mutation_drugs_found_list:
                        new_MTB_drugs_unfound_list = remove_item(MTB_drugs_unfound_list,drug)
                        if len(value) == 1:
                            new_MTB_gene_unfound_list = remove_item(MTB_gene_unfound_list,key)
            if len(MTB_mutation_gene_found_list) != 0:
                mutation_gene_info = '、'.join(MTB_mutation_gene_found_list)
                mutation_drugs_info = '、'.join(MTB_mutation_drugs_found_list)
                rt.add(f'结核分枝杆菌检出{mutation_gene_info}基因耐药位点突变，表示可能会对{mutation_drugs_info}药物耐药。请结合突变率及分子DST与表型DST符合率来综合判断。\n')
            else:
                found_gene_info = '、'.join(MTB_gene_found_list)
                found_drugs_info = '、'.join(MTB_drugs_found_list)
                rt.add(f'结核分枝杆菌检出{found_gene_info}基因耐药位点未发生突变，表示对{found_drugs_info}药物敏感。请结合突变率及分子DST与表型DST符合率来综合判断。\n')
            if len(new_MTB_gene_unfound_list) != 0:
                unfound_gene_info = '、'.join(new_MTB_gene_unfound_list)
                unfound_drugs_info = '、'.join(new_MTB_drugs_unfound_list)
                unfound_info = unfound_info + f'由于该样本中结核分枝杆菌含量较低，{unfound_gene_info}基因未能检出，无法判断其是否突变，从而不能判断对{unfound_drugs_info}敏感/耐药。\n'
        elif result['微生物'] != '/' and result['分类'] == '非结核分枝杆菌（NTM）':
            chinese_name = result['中文名']
            number += 1
            NTM_id_list = []
            for key, value in NTM_result_list.items():
                if value == result['微生物']:
                    NTM_id_list.append(key)
            NTM_df = handle_df[handle_df['chr'].isin(NTM_id_list)]
            NTM_genes = NTM_df['gene_name'].tolist()
            tmp_NTM_gene_found_list = []
            tmp_NTM_mutation_gene_found_list = []
            tmp_NTM_gene_unfound_list = []
            tmp_NTM_drugs_found_list = []
            tmp_NTM_mutation_drugs_found_list = []
            tmp_NTM_drugs_unfound_list = []
            for key, value in NTM_gene_dict.items():
                if key in NTM_genes:
                    tmp_NTM_gene_found_list.append(key)
                    tmp_NTM_drugs_found_list.extend(value)
                else:
                    tmp_NTM_gene_unfound_list.append(key)
                    tmp_NTM_drugs_unfound_list.extend(value)
            for line in NTM_result_list:
                if line['序号'] != '-':
                    tmp_NTM_mutation_gene_found_list.append(line['基因'])
                    tmp_NTM_mutation_drugs_found_list.append(line['药品'])
            new_tmp_NTM_drugs_unfound_list = tmp_NTM_drugs_unfound_list
            new_tmp_NTM_gene_unfound_list = tmp_NTM_gene_unfound_list
            for key, value in NTM_gene_dict.items():
                for drug in value:
                    if drug in tmp_NTM_drugs_unfound_list and drug in tmp_NTM_mutation_drugs_found_list:
                        new_tmp_NTM_drugs_unfound_list = remove_item(tmp_NTM_drugs_unfound_list,drug)
                        if len(value) == 1:
                            new_tmp_NTM_gene_unfound_list = remove_item(tmp_NTM_gene_unfound_list,key)
            if len(tmp_NTM_mutation_gene_found_list) != 0:
                mutation_gene_info = '、'.join(tmp_NTM_mutation_gene_found_list)
                mutation_drugs_info = '、'.join(tmp_NTM_mutation_drugs_found_list)
                rt.add(f'{chinese_name}检出{mutation_gene_info}基因耐药位点突变，表示可能会对{mutation_drugs_info}药物耐药。请结合突变率及分子DST与表型DST符合率来综合判断。\n')
            else:
                found_gene_info = '、'.join(tmp_NTM_gene_found_list)
                found_drugs_info = '、'.join(tmp_NTM_drugs_found_list)
                rt.add(f'{chinese_name}检出{found_gene_info}基因耐药位点未发生突变，表示对{found_drugs_info}药物敏感。请结合突变率及分子DST与表型DST符合率来综合判断。\n')
            if len(new_tmp_NTM_gene_unfound_list) != 0:
                unfound_gene_info = '、'.join(new_tmp_NTM_gene_unfound_list)
                unfound_drugs_info = '、'.join(new_tmp_NTM_drugs_unfound_list)
                unfound_info = unfound_info + f'由于该样本中{chinese_name}含量较低，{unfound_gene_info}基因未能检出，无法判断其是否突变，从而不能判断对{unfound_drugs_info}敏感/耐药。\n'
            NTM_gene_found_list.extend(tmp_NTM_gene_found_list)
            NTM_mutation_gene_found_list.extend(tmp_NTM_mutation_gene_found_list)
            NTM_gene_unfound_list.extend(tmp_NTM_gene_unfound_list)
            NTM_drugs_found_list.extend(tmp_NTM_drugs_found_list)
            NTM_mutation_drugs_found_list.extend(tmp_NTM_mutation_drugs_found_list)
            NTM_drugs_unfound_list.extend(tmp_NTM_drugs_unfound_list)
    if number == 0:
        rt.add(f'未检出分枝杆菌，请结合临床综合判断是否为其他的感染。\n')
    else:
        rt.add(unfound_info)
    return [rt,MTB_gene_found_list,NTM_gene_found_list]
    


def appendix_color_change(doc,
    table_list: list,
    number: int
    ) -> None:
    if number == 1:
        table_number_list = [10,11]
    else:
        table_number_list = [12]
    for table_number in table_number_list:
        table = doc.tables[table_number]
        rownums = len(table.rows)
        for table_line in table_list:
            drug_list = table_line['药品'].split(';')
            for drug in drug_list:
                for x in range(rownums):
                    drug_name = table.cell(x,1).text
                    if drug_name == drug:
                        if table_line['基因'] == table.cell(x,0).text:
                            run = table.cell(x,2).paragraphs[0]
                            run.text = "检出"
                            content = run.text
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(10)
                            run.font.name = 'Arial'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def color_change(doc,
    table_number
    ) -> None:
    table = doc.tables[table_number+1]
    rownums = len(table.rows)
    for x in range(rownums):
        if table.cell(x,0).text == '结核分枝杆菌复合群（MTBC）' or table.cell(x,0).text == '非结核分枝杆菌（NTM）' :
            run = table.cell(x,table_number-3).paragraphs[0]
            content = run.text
            run.text = ''
            run = run.add_run(content)
            if content == '阳性':
                run.font.color.rgb = RGBColor(255,0,0)
            else:
                run.font.color.rgb = RGBColor(0,176,80)
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def color_change2(doc,
    sample_code: str,
    info_client: pd.DataFrame,
    medical_DB: pd.DataFrame,
    complex_df: pd.DataFrame,
    appendix_species: list
    ) -> None:
    handle_df = info_client.query('样本编号 == @sample_code').iloc[0,:]
    result_list = handle_df['正式报告结果'].split(',')
    table = doc.tables[7]
    rownums = len(table.rows)
    find_list = []
    for bac_name in result_list:
        if len(bac_name) != 0:
            compare_bac_name = Nor(bac_name)
            chinese_name = 0
            try:
                chinese_name = Nor_col(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0])
            except IndexError:
                logging.info(f"{handle_df['患者姓名']}的结果{bac_name}未找到中文名")
            # print(chinese_name)
            if type(chinese_name) == str:
                if chinese_name in appendix_species:
                    for x in range(rownums):
                        if table.cell(x,1).text == chinese_name:
                            run = table.cell(x,3).paragraphs[0]
                            run.text = '检出'
                            content = run.text
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(9)
                            run.font.name = 'Arial'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            run = table.cell(x,2).paragraphs[0]
                            content = run.text
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(9)
                            run.font.name = 'Arial'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            run = table.cell(x,1).paragraphs[0]
                            content = run.text
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(9)
                            run.font.name = 'Arial'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')                            
                    result_list.remove(bac_name)
                    find_list.append(chinese_name)
    if len(result_list) != 0:
        for bac_name in result_list:
            if len(bac_name) != 0:
                compare_bac_name = Nor_col(bac_name)
                find_complex_df = complex_df[complex_df['name'] == compare_bac_name]
                if find_complex_df.shape[0] == 0:
                    classification_name = '非结核分枝杆菌'
                else:
                    classification_name = '结核分枝杆菌复合群'
                chinese_name = 0
                try:
                    chinese_name = Nor_col(medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0])
                except IndexError:
                    logging.info(f"{handle_df['患者姓名']}的结果{bac_name}未找到中文名")
                if type(chinese_name) == str:
                    for x in range(rownums):
                        if table.cell(x,1).text != '结核分枝杆菌' and table.cell(x,1).text not in find_list and table.cell(x,1).text == classification_name:
                            table.cell(x,1).text = chinese_name
                            run = table.cell(x,3).paragraphs[0]
                            run.text = '检出'
                            content = run.text
                            run.text = ''
                            run = run.add_run(content)
                            run.font.color.rgb = RGBColor(255,0,0)
                            run.font.size = Pt(10)
                            run.font.name = 'Arial'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    result_list.remove(bac_name)
                    find_list.append(chinese_name)
#def move_picture_barcode(doc,
#    png_name: str,
#    barcode_picture_path):
#    table = doc.tables[0]
#    png_name=png_name.replace('-s','')
#    paragraph=table.cell(0,0).paragraphs[0]
#    run = paragraph.add_run()
#    # print('barcode_list[0]',barcode_list[0])
#    os.chdir(barcode_picture_path)
#    barcode_list=os.listdir('.')
#    try:
#        run.add_picture(barcode_picture_path+png_name+'.png',width=Cm(2.5),height=Cm(2.5))
#        print(barcode_list)
#    except:
#        for picture_name in barcode_list:
#            if str(png_name) in str(picture_name):
#                print(png_name,picture_name)
#                run.add_picture(barcode_picture_path+picture_name,width=Cm(2.5),height=Cm(2.5))
#    print(barcode_picture_path+png_name+'.png')
#    return doc
def move_picture_barcode(doc,
    png_name: str,
    barcode_picture_path):
    table = doc.tables[0]
    png_name=png_name.replace('-s','')
    paragraph=table.cell(1,0).paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(barcode_picture_path+png_name+'.png',width=Cm(2.1),height=Cm(2.1))
    print(barcode_picture_path+png_name+'.png')
    return doc


####################################################
barcode_picture_path = args.barcode_picutre
doc_list = ['seegene','mobai','boruilin','beijing']
# 对应物种的基因组ID
id_dict = {'NC_000962.3':'Mycobacterium tuberculosis','CU458896.1':'Mycobacteroides abscessus','NR_025584.1':'Mycobacterium avium','NR_042165.1':'Mycobacterium intracellulare','NR_121712.2':'Mycobacterium kansasii','NR_114659.1':'Mycobacteroides chelonae','NR_025311.1':'Mycolicibacterium smegmatis','NR_042912.1':'Mycolicibacterium fortuitum','NC_010397.1':'Mycobaccterium abscessus','NG_041979.1':'Mycobacterium avium','NR_076151.1':'Mycobacterium intracellulare','NZ_CP019883.1':'Mycobacterium kansasii','GU143889.1':'Mycobacteroides chelonae','AB011184.1':'Mycolicibacterium smegmatis','NZ_CP011269.1:c3552322-3549198':'Mycolicibacterium fortuitum','NZ_CP014955.1:2353195-2353716':'Mycobacteroides abscessus'}
# MTB和NTM的耐药基因和对应药物
MTB_gene_dict = {'rpoB':['利福平'],'embB':['乙胺丁醇'],'pncA':['吡嗪酰胺'],'katG':['异烟肼'],'inhA':['异烟肼','乙硫异烟胺','丙硫异烟胺'],'gyrA':['氟喹诺酮类'],'gyrB':['氟喹诺酮类'],'rpsL':['链霉素'],'rrs':['链霉素','阿米卡星','卡那霉素','卷曲霉素'],'folC':['对氨基水杨酸钠'],'thyA':['对氨基水杨酸钠'],'alr':['环丝氨酸'],'Rv0678':['氯法齐明','贝达喹啉'],'rplC':['利奈唑胺']}
NTM_gene_dict = {'rrl':['克拉霉素','阿奇霉素'],'erm':['克拉霉素','阿奇霉素'],'rrs':['阿米卡星','卡那霉素','庆大霉素']}
# 药物分类
first_list = ['利福平','异烟肼','吡嗪酰胺','乙胺丁醇']
second_list = ['阿米卡星','氟喹诺酮类','乙硫异烟胺','丙硫异烟胺','对氨基水杨酸钠','链霉素','卡那霉素','卷曲霉素','环丝氨酸','氯法齐明','贝达喹啉','利奈唑胺']
NTM_list = ['克拉霉素','阿奇霉素','阿米卡星','卡那霉素','庆大霉素']
# 附录中的种
appendix_species = [ '结核分枝杆菌', '非洲分枝杆菌', '牛分枝杆菌', '山羊分枝杆菌', '田鼠分枝杆菌', '卡内蒂分枝杆菌', '鳍脚分枝杆菌', '獴分枝杆菌', '鸟分枝杆菌', '胞内分枝杆菌', '副胞内分枝杆菌', '堪萨斯分枝杆菌', '龟分枝杆菌', '脓肿分枝杆菌', '猿猴分枝杆菌', '溃疡分枝杆菌', '偶发分枝杆菌', '玛尔摩分枝杆菌', '海分枝杆菌', '马赛分枝杆菌', '戈登分枝杆菌', '副戈登分枝杆菌', '产粘液分枝杆菌']
####################################################
info_client = pd.read_excel(args.result_excel).fillna('NA')
# Interpretation = pd.read_excel(args.result_excel, sheet_name='解读').fillna('NA')
AMR_df = pd.read_excel(args.result_excel, sheet_name='drug_report').fillna('NA')
# print(AMR_df.columns.values)
AMR_stat_df = pd.read_excel(args.result_excel, sheet_name='depth_report').fillna('NA')
Interpretation = pd.read_excel(args.result_excel, sheet_name='species_report').fillna('NA')
medical_DB = pd.read_excel(args.database).fillna('NA')
result_file_name = args.result_excel.split("/")[-1]
complex_df = pd.read_excel(args.complex_excel)
name_df = pd.read_excel(args.name_excel)
# 统一输入的英文名
medical_DB['种'] = medical_DB['种'].apply(Nor_col)
medical_DB['英文名称'] = medical_DB['英文名称'].apply(Nor)
info_client['样本编号'] = info_client['样本编号'].apply(Nor_col)
info_client['患者姓名'] = info_client['患者姓名'].apply(Nor_col)
Interpretation['name'] = Interpretation['name'].apply(Nor)
# quality_info['name'] = quality_info['name'].apply(Nor)
AMR_df['drug'] = AMR_df['drug'].apply(Nor)
complex_df['name'] = complex_df['name'].apply(Nor)
dic_client = {}
for row_index in range(info_client.shape[0]):
    if info_client['正式报告结果'][row_index] != 'NA' or info_client['补充报告结果'][row_index] != 'NA':
        dic_client.setdefault(info_client['样本编号'][row_index], {})
        dic_client[info_client['样本编号'][row_index]]['样本编号'] = info_client['样本编号'][row_index].split('-')[0]
        dic_client[info_client['样本编号'][row_index]]['姓名'] = info_client['患者姓名'][row_index]
        logging.info(f"开始处理{info_client['样本编号'][row_index]} {info_client['患者姓名'][row_index]}的信息")
        dic_client[info_client['样本编号'][row_index]]['性别'] = info_client['性别'][row_index]
        if info_client['代理商'][row_index].lower() in doc_list:
            manufacturer = info_client['代理商'][row_index].lower()
        else:
            for company in doc_list:
                if result_file_name.startswith(company):
                    manufacturer = company
        print(manufacturer)
        dic_client[info_client['样本编号'][row_index]]['代理商'] = manufacturer
        try:
            dic_client[info_client['样本编号'][row_index]]['年龄'] = int(info_client['年龄'][row_index])
        except ValueError:
            dic_client[info_client['样本编号'][row_index]]['年龄'] = info_client['年龄'][row_index]
        dic_client[info_client['样本编号'][row_index]]['临床检测'] = info_client['培养结果'][row_index]
        dic_client[info_client['样本编号'][row_index]]['临床诊断'] = info_client['临床诊断'][row_index]
        dic_client[info_client['样本编号'][row_index]]['临床用药'] = info_client['抗感染用药史'][row_index]
        dic_client[info_client['样本编号'][row_index]]['送检单位'] = info_client['医院'][row_index]
        dic_client[info_client['样本编号'][row_index]]['科室'] = info_client['科室'][row_index]
        dic_client[info_client['样本编号'][row_index]]['样本类型'] = info_client['样本类型'][row_index]
        if '血' in dic_client[info_client['样本编号'][row_index]]['样本类型']:
            dic_client[info_client['样本编号'][row_index]]['样本类型'] = '血液'
        try:
            receivedate = str(info_client['收样时间'][row_index]).split(' ')[0]
        except IndexError:
            dic_client[info_client['样本编号'][row_index]]['收样日期'] = ''
        dic_client[info_client['样本编号'][row_index]]['收样日期'] = receivedate.split('/')[0]+'年'+receivedate.split('/')[1]+'月'+receivedate.split('/')[2]+'日'
        try:
            exdate = str(info_client['上机日期'][row_index]).split(' ')[0]
            dic_client[info_client['样本编号'][row_index]]['实验日期'] = exdate.split('-')[0]+'年'+exdate.split('-')[1]+'月'+exdate.split('-')[2]+'日'
        except IndexError:
            exdate = str(info_client['上机日期'][row_index])
            dic_client[info_client['样本编号'][row_index]]['实验日期'] = exdate.split('/')[0]+'年'+exdate.split('/')[1]+'月'+exdate.split('/')[2]+'日'
        try:
            reportdate = str(info_client['报告日期'][row_index]).split(' ')[0]
        except IndexError:
            dic_client[info_client['样本编号'][row_index]]['报告日期'] = ''
        dic_client[info_client['样本编号'][row_index]]['报告日期'] = reportdate.split('/')[0]+'年'+reportdate.split('/')[1]+'月'+reportdate.split('/')[2]+'日'
        # print(reportdate)
        shortname = project_shorthand(info_client['样本编号'][row_index])
        if type(shortname) != int:
            AMR_stat_colname = str(info_client['barcode'][row_index]) + '_'  + project_shorthand(info_client['样本编号'][row_index]) + '_' + dic_client[info_client['样本编号'][row_index]]['姓名'] + '_' + str(info_client['样本编号'][row_index]) + '_count'
            file_name = AMR_stat_colname.replace(" ", ".")
        else:
            AMR_stat_colname = str(info_client['barcode'][row_index]) + '_'  + dic_client[info_client['样本编号'][row_index]]['姓名'] + '_' + str(info_client['样本编号'][row_index]) + '_count'
            file_name = AMR_stat_colname.replace(" ", ".")
        if info_client['正式报告结果'][row_index] != ',' and info_client['正式报告结果'][row_index] != '，' and '耐药基因' in info_client['检测项目'][row_index]:
            dic_client[info_client['样本编号'][row_index]]['图片文件名'] = file_name + '.png'
            dic_client[info_client['样本编号'][row_index]]['测序深度'] = str(int(AMR_stat_df[AMR_stat_colname][11])) + '×'
            dic_client[info_client['样本编号'][row_index]]['覆盖率'] = str(float(format(AMR_stat_df[AMR_stat_colname][12], '.2f'))) + '%'
            make_picture(AMR_stat_colname=AMR_stat_colname,AMR_stat_df=args.result_excel,file_name=file_name)
        else:
            dic_client[info_client['样本编号'][row_index]]['图片文件名'] = file_name + '.png'
            dic_client[info_client['样本编号'][row_index]]['测序深度'] = '--'
            dic_client[info_client['样本编号'][row_index]]['覆盖率'] = '--'
        table_info = table_context(sample_code=info_client['样本编号'][row_index], info_client=info_client, medical_DB=medical_DB, complex_df=complex_df, id_dict=id_dict, AMR_df=AMR_df, AMR_stat_colname=AMR_stat_colname, Interpretation=Interpretation)
        dic_client[info_client['样本编号'][row_index]]['表1信息'] = table_info[0]
        dic_client[info_client['样本编号'][row_index]]['表2信息'] = table_info[1]
        dic_client[info_client['样本编号'][row_index]]['表3信息'] = table_info[2]
        dic_client[info_client['样本编号'][row_index]]['表4信息'] = table_info[3]
        appendix_info = appendix_drugs(table2_list=table_info[1],table3_list=table_info[2],first_list=first_list,second_list=second_list,NTM_list=NTM_list)
        dic_client[info_client['样本编号'][row_index]]['一线检出'] = appendix_info[0]
        dic_client[info_client['样本编号'][row_index]]['一线未检出'] = appendix_info[1]
        dic_client[info_client['样本编号'][row_index]]['二线检出'] = appendix_info[2]
        dic_client[info_client['样本编号'][row_index]]['二线未检出1'] = appendix_info[3]
        dic_client[info_client['样本编号'][row_index]]['二线未检出2'] = appendix_info[4]
        dic_client[info_client['样本编号'][row_index]]['检出'] = appendix_info[5]
        dic_client[info_client['样本编号'][row_index]]['未检出'] = appendix_info[6]
        # result_info = add_medical_interpretation(AMR_stat_colname=AMR_stat_colname,sample_type=info_client['样本类型'][row_index],all_result_list=table_info[0],MTB_result_list=table_info[1],NTM_result_list=table_info[2],id_dict=id_dict,MTB_gene_dict=MTB_gene_dict,NTM_gene_dict=NTM_gene_dict,AMR_df=AMR_df)
        # dic_client[info_client['样本编号'][row_index]]['医学解读'] = result_info[0]
        # dic_client[info_client['样本编号'][row_index]]['MTB检出基因'] = result_info[0]
        # dic_client[info_client['样本编号'][row_index]]['NTM检出基因'] = result_info[0]
logging.info("所有信息处理完成！")
print(dic_client)
# 输出文件夹的创建
filename_date = reportdate.split('/')[0]+reportdate.split('/')[1]+reportdate.split('/')[2]
result_file_name = result_file_name.lower()
for company in doc_list:
    if result_file_name.startswith(company):
        manufacturer = company
try:
    save_path = os.path.join(args.output_dir, manufacturer,filename_date)
except NameError:
    logging.info(f"{result_file_name}文件名中没有正确的批次号,请修改！")
if save_path and not os.path.exists(save_path):
    os.makedirs(save_path)
for sample_code,value in dic_client.items():
    print(value['姓名'])
    print('value',value['表2信息'])
    print('value',value['表3信息'])
    project = info_client.loc[info_client['样本编号'] == sample_code, '检测项目'].iloc[0]
    agent = dic_client[sample_code]['代理商']
    # print(agent)
    save_path = os.path.join(args.output_dir, agent,filename_date)
    if save_path and not os.path.exists(save_path):
        os.makedirs(save_path)
    if '耐药基因' in project:
        table_number = 4
        result_report_name = filename_date + '_'+ dic_client[sample_code]['样本编号'] + '_' + str(dic_client[sample_code]['姓名'])+'_分枝杆菌(MTB&NTM)及其耐药基因检测报告.docx'
        logging.info(f"{result_report_name}开始生成")
        doc = DocxTemplate(f'/mnt/c/Users/luping/Desktop/报告流程/TB/{agent}/分枝杆菌(MTB&NTM)及其耐药基因检测报告.docx')
        # print(dic_client[sample_code])
        doc.render(dic_client[sample_code])
        png_name = dic_client[sample_code]['图片文件名']
        if os.path.exists(dic_client[sample_code]['图片文件名']):
            doc = move_picture(doc=doc,png_name=png_name)
        number = 1
        appendix_color_change(doc=doc,table_list=dic_client[sample_code]['表2信息'],number=number)
        number += 1
        appendix_color_change(doc=doc,table_list=dic_client[sample_code]['表3信息'],number=number)
        color_change(doc=doc,table_number=table_number)
        result_file = os.path.join(save_path, result_report_name)
        try:
            doc = move_picture_barcode(doc=doc, png_name=sample_code,barcode_picture_path=barcode_picture_path)
        except:
            print(value['姓名']+'的结果条形码图片未找到，请核对是否放入数据库中')
        doc.save(result_file)
        logging.info(f"{result_report_name}生成成功！")
    else:
        table_number = 4
        result_report_name = filename_date + '_'+ dic_client[sample_code]['样本编号'] + '_' + str(dic_client[sample_code]['姓名'])+'_TB_seq结核&非结核分枝杆菌鉴定.docx'
        logging.info(f"{result_report_name}开始生成")
        doc = DocxTemplate(f'/mnt/c/Users/luping/Desktop/报告流程/TB/{agent}/TB_seq结核&非结核分枝杆菌鉴定.docx')
        # print(dic_client[sample_code])
        doc.render(dic_client[sample_code])
        color_change(doc=doc,table_number=table_number)
        color_change2(doc=doc,sample_code=sample_code,info_client=info_client,medical_DB=medical_DB,complex_df=complex_df,appendix_species=appendix_species)
        result_file = os.path.join(save_path, result_report_name)
        try:
            doc = move_picture_barcode(doc=doc, png_name=sample_code,barcode_picture_path=barcode_picture_path)
        except:
            print(value['姓名']+'的结果条形码图片未找到，请核对是否放入数据库中')
        doc.save(result_file)
        logging.info(f"{result_report_name}生成成功！")
logging.info(f"开始汇总结果！")
for sample_code,value in dic_client.items():
    # print(value)
    # print('value',value['表2信息'][0]['基因'])
    # print('value',value['表3信息'][0]['基因'])
    manufacturer = dic_client[sample_code]['代理商']
    summary_excel_path = os.path.join(args.summary_excel,manufacturer,'汇总.xlsx')
    all_summary_df = pd.read_excel(summary_excel_path)
    all_info_list = []
    if all_summary_df.shape[0] != 0:
        for row_index in range(all_summary_df.shape[0]):
            line_info = str(all_summary_df['样本编号'][row_index]) + '_' + str(all_summary_df['患者姓名'][row_index]) + '_' + str(all_summary_df['basecalling data'][row_index]) + '_' + all_summary_df['Reports'][row_index]
            all_info_list.append(line_info)
    all_info_dict = dict(zip(all_info_list,all_info_list))
    data_summary_excel_path = os.path.join(args.summary_excel,manufacturer,filename_date,'汇总.xlsx')
    if data_summary_excel_path and os.path.exists(data_summary_excel_path):
        data_summary_df = pd.read_excel(data_summary_excel_path)
    else:
        data_summary_df = pd.DataFrame(columns=all_summary_df.columns.tolist())
    
    need_info = info_client.query('样本编号 == @sample_code').iloc[0]
    all_list = []
    try:
        all_list = need_info['正式报告结果'].split(',')    
    except AttributeError:
        all_list = []
    try:
        all_list.extend(need_info['补充报告结果'].split(','))
    except AttributeError:
        all_list = all_list
    for micro in all_list:
        #print(micro)
        compare_bac_name = Nor(micro)
        data_info_list = []
        if data_summary_df.shape[0] != 0:
            for row_index in range(data_summary_df.shape[0]):
                line_info = str(data_summary_df['样本编号'][row_index]) + '_' + str(data_summary_df['患者姓名'][row_index]) + '_' + str(data_summary_df['basecalling data'][row_index]) + '_' + data_summary_df['Reports'][row_index]
                data_info_list.append(line_info)
        data_info_dict = dict(zip(data_info_list,data_info_list))
        Resisit_MT=['结核耐药-突变基因','结核耐药-潜在耐药药物','结核耐药-核酸突变结果','结核耐药-氨基酸突变','结核耐药-突变率',
                    '非结核耐药-突变基因','非结核耐药-潜在耐药药物','非结核耐药-核酸突变结果','非结核耐药-氨基酸突变','非结核耐药-突变率']
        Relist=['Reports','报告结果','报告结果','报告等级','物种','致病性描述','Reads','样本描述（解读人员手动添加）','备 注','采样时间.1']#存放不要的结果
        if len(micro) != 0 and micro != 'NA':
            info_list = all_summary_df.columns.tolist()[:-10]
            new_line_dict = {}
            for colname in info_list: 
                if colname not in Relist:
                    new_line_dict[colname] = need_info[colname]
            #print('micro',micro)
            #print('new_line_dict[colname]',new_line_dict[colname])
            #print(type(new_line_dict[colname]))
            new_line_dict['Reports'] = micro
            new_line_dict['basecalling data'] = need_info['basecalling data']
            try:
                new_line_dict['报告结果'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '种'].iloc[0]
            except IndexError:
                new_line_dict['报告结果'] = 'NA'
            try:
                new_line_dict['物种'] = medical_DB.loc[medical_DB['英文名称'] == compare_bac_name, '分类'].iloc[0]
            except IndexError:
                new_line_dict['物种'] = 'NA'
            if micro in need_info['正式报告结果']:
                new_line_dict['报告等级'] = 'A'
            else:
                new_line_dict['报告等级'] = 'B'
            #print('sample_code',sample_code)
            #print('medical_DB.检测项目',medical_DB.检测项目)
            #print('project_shorthand(str(sample_code)',project_shorthand(sample_code))
            try:
                new_line_dict['致病性描述'] = medical_DB.loc[(medical_DB['英文名称']== compare_bac_name) & (medical_DB.检测项目.str.contains(project_shorthand(sample_code), regex=True)), '备注'].iloc[0]
            except:
                new_line_dict['致病性描述'] = 'NA'
            column_R = str(info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0]) + '_R_' + str(project_shorthand(sample_code)) + '_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
            try:
                new_line_dict['Reads'] = Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0]
            except KeyError:
                column_R = str(info_client.loc[info_client['样本编号'] == sample_code, 'barcode'].iloc[0]) + '_R_' + str(info_client.loc[info_client['样本编号'] == sample_code, '患者姓名'].iloc[0]) + '_' + sample_code
                new_line_dict['Reads'] = Interpretation.loc[Interpretation['name'] == compare_bac_name, column_R].iloc[0]
            except IndexError:
                new_line_dict['Reads'] = 'NA'
            new_line_dict['样本描述（解读人员手动添加）'] = ' '
            new_line_dict['制作人'] = ' '
            new_line_dict['审核人'] = ' '
            new_line_dict['备 注'] = ' '
            MTgene=[];MTdrug=[];MTdescript=[];MTProtein=[];MTPrecent=[]
            for MTlist in value['表2信息']:
                #print(MTlist)
                for key2,value2 in MTlist.items():
                    print(value['姓名'])
                    #print('key8,value8',key8,value8)
                    if key2=='基因':
                        MTgene.append(value2)
                    if key2=='药品':
                        MTdrug.append(value2)
                    if key2=='突变描述':
                        MTdescript.append(value2)
                    if key2=='氨基酸突变':
                        MTProtein.append(value2)
                    if key2=='突变率':
                        MTPrecent.append(value2)
            #print('MTgene,MTdrug,MTdescript,MTProtein,Precent',MTgene,MTdrug,MTdescript,MTProtein,MTPrecent)
            new_line_dict['结核耐药-突变基因'] = MTgene
            new_line_dict['结核耐药-潜在耐药药物'] = MTdrug
            new_line_dict['结核耐药-核酸突变结果'] = MTdescript
            new_line_dict['结核耐药-氨基酸突变'] = MTProtein
            new_line_dict['结核耐药-突变率'] = MTPrecent
            NMTgene=[];NMTdrug=[];NMTdescript=[];NMTProtein=[];NMTPrecent=[]
            for NMTlist in value['表3信息']:
                #print(MTlist)
                for key3,value3 in NMTlist.items():
                    print(value['姓名'])
                    #print('key8,value8',key8,value8)
                    if key3=='基因':
                        NMTgene.append(value3)
                    if key3=='药品':
                        NMTdrug.append(value3)
                    if key3=='突变描述':
                        NMTdescript.append(value3)
                    if key3=='氨基酸突变':
                        NMTProtein.append(value3)
                    if key3=='突变率':
                        NMTPrecent.append(value3)
            new_line_dict['非结核耐药-突变基因'] = NMTgene
            new_line_dict['非结核耐药-潜在耐药药物'] = NMTdrug
            new_line_dict['非结核耐药-核酸突变结果'] = NMTdescript
            new_line_dict['非结核耐药-氨基酸突变'] = NMTProtein
            new_line_dict['非结核耐药-突变率'] = NMTPrecent
            #print(new_line_dict)
            new_df = (new_line_dict)
            new_line_info = str(new_line_dict['样本编号']) + '_' + str(new_line_dict['患者姓名']) + '_' + str(new_line_dict['basecalling data']) + '_' + new_line_dict['Reports']
            # print(data_info_dict)
            # print(all_info_dict)
            # print(new_line_info)
            if new_line_info not in data_info_dict:
                data_summary_df = data_summary_df.append(new_df, ignore_index=True)
            if new_line_info not in all_info_dict:
                all_summary_df = all_summary_df.append(new_df, ignore_index=True)
    data_summary_df.to_excel(data_summary_excel_path, index=False)
    all_summary_df.to_excel(summary_excel_path, index=False)
logging.info(f"汇总完成！")
